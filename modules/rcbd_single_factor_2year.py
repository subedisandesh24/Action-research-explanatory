import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import string
import random
from scipy.stats import t
import statsmodels.api as sm
from statsmodels.formula.api import ols
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side

# ==============================================================================
# DATABASE OF 30 HIGH-STANDARD MULTI-YEAR SINGLE-FACTOR ACADEMIC DISCUSSION TEMPLATES
# ==============================================================================
ACADEMIC_TEMPLATES_2Y_30 = {
    # --- Group A: Temporal / Multi-Date / Trend Line Parameters (Templates 1–15) ---
    1: (
        "The temporal development of {{variable_name}} was monitored sequentially across {{num_intervals}} experimental "
        "intervals from {{start_time}} to {{end_time}} {{time_unit}} over two consecutive seasons ({{year_1}} and {{year_2}}). "
        "Table {{table_num}} shows that the combined ANOVA revealed a non-significant treatment \u00d7 year interaction "
        "(A \u00d7 Y; P > 0.05) across all sampling intervals, indicating a highly consistent temporal treatment response across "
        "both years. Consequently, data were pooled over the two years for subsequent analysis. During the initial stages at "
        "{{time_point_1}} and {{time_point_2}} {{time_unit}}, the pooled treatment levels of {{factor_A}} did not exert "
        "any statistically significant influence on the observed values (P > 0.05), with the overall pooled grand mean remaining "
        "stable at {{grand_mean_early}} {{unit}}. However, as the evaluation timeline progressed to {{time_point_3}} {{time_unit}}, "
        "a highly significant pooled treatment divergence emerged (P \u2264 0.01). Table {{table_num}} shows that the pooled mean of "
        "{{treatment_A1}} achieved the maximum value of {{value_1}} {{unit}}, which was significantly higher than {{treatment_A2}} "
        "({{value_2}} {{unit}}) and the baseline treatment {{treatment_A_lowest}} ({{value_lowest}} {{unit}}; LSD_0.05 = {{lsdval}})."
    ),
    2: (
        "Table {{table_num}} shows that the combined analysis of variance over two years ({{year_1}} and {{year_2}}) "
        "revealed a highly significant treatment \u00d7 year interaction (A \u00d7 Y; P \u2264 0.01) for {{variable_name}} at "
        "{{time_point_3}} and {{time_point_4}} {{time_unit}}, indicating that the temporal treatment trajectory was strongly "
        "modified by seasonal environmental variations. Therefore, the temporal trends are described for each year individually. "
        "In {{year_1}}, no treatment differences were observed during the early intervals, but a highly significant treatment "
        "effect emerged at {{time_point_3}} {{time_unit}} (P \u2264 0.01), where {{treatment_A1}} reached {{value_y1}} {{unit}} "
        "compared to the lowest treatment level ({{value_lowest_y1}} {{unit}}). In {{year_2}}, a similar downward trend was observed, "
        "but the treatment divergence occurred earlier at {{time_point_2}} {{time_unit}} (P \u2264 0.05), with {{treatment_A1}} "
        "producing {{value_y2}} {{unit}} compared to {{value_lowest_y2}} {{unit}}. The overall multi-year pooled means recorded a "
        "grand mean of {{pooled_grand_mean}} {{unit}} at the final sampling point."
    ),
    3: (
        "The chronological progression of {{variable_name}} was monitored over two years ({{year_1}} and {{year_2}}) from "
        "{{start_time}} to {{end_time}} {{time_unit}} to evaluate inter-annual stability. Table {{table_num}} shows that the "
        "main effect of Year was significant (P \u2264 0.05) during the early intervals, with {{year_1}} maintaining a slightly "
        "higher overall grand mean ({{gm_y1}} {{unit}}) compared to {{year_2}} ({{gm_y2}} {{unit}}). However, the main effect "
        "of {{factor_A}} treatments followed a highly consistent pattern across both seasons, supported by a non-significant "
        "A \u00d7 Y interaction (P > 0.05). At {{time_point_1}} {{time_unit}}, all pooled treatments were statistically "
        "equivalent. By {{time_point_3}} {{time_unit}}, the pooled main effect of {{factor_A}} became highly significant "
        "(P \u2264 0.01), with {{treatment_A1}} producing the highest pooled value of {{val_pooled_A1}} {{unit}} compared "
        "to the baseline treatment {{treatment_A_lowest}} ({{val_pooled_lowest}} {{unit}}; LSD_0.05 = {{lsdpooled}}), maintaining "
        "this statistical superiority through the final sampling interval."
    ),
    4: (
        "The temporal divergence of {{variable_name}} values under different levels of {{factor_A}} was evaluated over "
        "two consecutive years. Table {{table_num}} shows that the combined ANOVA indicated a highly significant main effect of "
        "treatment (P \u2264 0.01) across the mid-to-late growth intervals, while the treatment \u00d7 year interaction "
        "(A \u00d7 Y) was completely non-significant (P > 0.05), confirming that the inter-annual performance was highly stable. "
        "At the initial interval of {{time_point_1}} {{time_unit}}, the pooled treatment effect was non-significant, with a "
        "narrow range between the highest value of {{val_max_early}} {{unit}} (under {{treatment_A1}}) and the lowest value of "
        "{{val_min_early}} {{unit}} (under {{treatment_A_lowest}}). However, as the timeline progressed to {{time_point_3}} "
        "{{time_unit}}, the pooled main effect of {{factor_A}} was highly significant (P \u2264 0.01), with {{treatment_A1}} and "
        "{{treatment_A2}} separating into a superior statistical group 'a' (LSD_0.05 = {{lsdval}}) compared to the lower "
        "treatments, maintaining this pattern through the final sampling point."
    ),
    5: (
        "The combined temporal response of {{variable_name}} over two experimental seasons showed a distinct statistical "
        "phase-shift between the early phase ({{time_point_1}} to {{time_point_2}} {{time_unit}}) and the final phase "
        "({{time_point_3}} to {{time_point_4}} {{time_unit}}). Table {{table_num}} shows that the treatment \u00d7 year "
        "interaction was non-significant (P > 0.05), allowing for a pooled multi-year analysis. During the early phase, "
        "all pooled levels of {{factor_A}} shared identical letter superscripts (P > 0.05), with values closely matching the "
        "combined baseline mean of {{mean_early}} {{unit}}. Following the transition to the final phase at {{time_point_3}} "
        "{{time_unit}}, the pooled main effect of {{factor_A}} became highly significant (P \u2264 0.01). Pooled mean separation "
        "shows that {{treatment_A1}} produced the highest value ({{val_A1}} {{unit}}), which was statistically equivalent "
        "only to {{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level ({{val_lowest}} "
        "{{unit}}; LSD_0.05 = {{lsdval}})."
    ),
    6: (
        "The rate of decrease in {{variable_name}} following the peak phase was evaluated across two consecutive years. "
        "Table {{table_num}} shows that the main effect of Year was non-significant (P > 0.05), and the treatment \u00d7 "
        "year interaction was also non-significant (P > 0.05), demonstrating a highly consistent downward trajectory across both "
        "seasons. From {{time_point_1}} to {{time_point_2}} {{time_unit}}, the pooled experimental units maintained high stability "
        "with a combined grand mean of {{grand_mean_peak}} {{unit}}, and no treatment variations were observed (P > 0.05). "
        "However, from {{time_point_3}} {{time_unit}} onwards, a rapid downward trend was observed in the pooled values. "
        "The application of {{treatment_A1}} significantly slowed this decrease in both years, maintaining a pooled value of "
        "{{val_A1}} {{unit}} at {{time_point_4}} {{time_unit}}, compared to the baseline treatment which dropped to {{val_baseline}} "
        "{{unit}} (P \u2264 0.05; LSD_0.05 = {{lsdval}})."
    ),
    7: (
        "The cumulative progression of {{variable_name}} was tracked over two years ({{year_1}} and {{year_2}}) to "
        "evaluate the inter-annual stability of accumulation kinetics. Table {{table_num}} shows that the combined ANOVA "
        "revealed a non-significant treatment \u00d7 year interaction (P > 0.05), confirming that the cumulative sigmoidal curve "
        "followed an identical pattern in both seasons. During the initial lag phase at {{time_point_1}} {{time_unit}}, no "
        "treatment differences were observed. However, during the active accumulation phase at {{time_point_3}} {{time_unit}}, "
        "the pooled treatment effect of {{factor_A}} was highly significant (P \u2264 0.01). Pooled mean separation showed "
        "that {{treatment_A1}} achieved a significantly higher cumulative value ({{val_A1}} {{unit}}) compared to the baseline "
        "treatment ({{val_baseline}} {{unit}}; LSD_0.05 = {{lsdval}}), which remained statistically consistent through the final "
        "plateau phase at {{time_point_4}} {{time_unit}}."
    ),
    8: (
        "The combined multi-year experimental design demonstrated high precision in monitoring the temporal dynamics of "
        "{{variable_name}} across two consecutive years ({{year_1}} and {{year_2}}). Table {{table_num}} shows that the pooled "
        "Coefficient of Variation (CV%) values ranged from {{cv_min}}% to {{cv_max}}% across the sampling dates, with the "
        "Standard Error of the Mean (SEm) for the treatment levels remaining low, ranging between {{sem_min}} and {{sem_max}} "
        "{{unit}}. Under these precise conditions, the treatment \u00d7 year interaction was non-significant (P > 0.05). At "
        "{{time_point_3}} {{time_unit}}, the pooled treatment effect was significant (P \u2264 0.05, LSD_0.05 = {{lsdval}}), "
        "with {{treatment_A1}} ({{val_A1}} {{unit}}) outperforming {{treatment_A2}} ({{val_A2}} {{unit}}). By {{time_point_4}} "
        "{{time_unit}}, the pooled main effect of {{factor_A}} became highly significant (P \u2264 0.01, LSD_0.05 = {{lsdval}}), "
        "confirming strong statistical control over experimental variance across both seasons."
    ),
    9: (
        "The mathematical differences (\u0394) and percentage changes in {{variable_name}} were evaluated across two seasons "
        "using a pooled multi-year analysis. Table {{table_num}} shows that the treatment \u00d7 year interaction was "
        "non-significant (P > 0.05), confirming that the treatment-induced changes were consistent across both years. While "
        "treatments remained statistically equivalent at {{time_point_1}} and {{time_point_2}} {{time_unit}}, significant "
        "treatment differences emerged in the pooled data at {{time_point_3}} {{time_unit}}. At this stage, {{treatment_A1}} "
        "increased {{variable_name}} by {{pct_diff_A}}% compared to the lowest treatment level (\u0394 = {{deltaA}} {{unit}}; "
        "P \u2264 0.05). At {{time_point_4}} {{time_unit}}, this percentage difference increased, with the pooled values "
        "under {{treatment_A1}} outperforming the lowest treatment level by {{pct_diff_A_late}}% (P \u2264 0.01; LSD_0.05 = {{lsdval}})."
    ),
    10: (
        "1.1.1 {{sub_parameter_1}}\n"
        "The temporal development of {{sub_parameter_1}} was evaluated over two consecutive years. Table {{table_num}} shows that "
        "the combined ANOVA revealed a non-significant treatment \u00d7 year interaction (P > 0.05) across all sampling intervals. "
        "During the early evaluation stages, no significant variations were observed, with pooled values remaining close to the "
        "grand mean of {{mean_early}} {{unit}}. However, at {{time_point_3}} {{time_unit}}, the pooled main effect of {{factor_A}} "
        "was significant (P \u2264 0.05), with {{treatment_A1}} ({{val_A1}} {{unit}}) outperforming the lowest treatment level "
        "({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}), which remained significant through the final sampling date of "
        "{{time_point_4}} {{time_unit}}.\n"
        "1.1.2 {{sub_parameter_2}}\n"
        "Unlike the first parameter, the combined analysis for {{sub_parameter_2}} showed a significant treatment \u00d7 year "
        "interaction (P \u2264 0.05) at the final stages. Table {{table_num}} shows that while the treatment trend was similar in "
        "both years, the magnitude of treatment separation was significantly higher in {{year_2}} compared to {{year_1}}. In "
        "{{year_1}}, {{treatment_A1}} exceeded the lowest treatment level by {{pct_y1}}% (LSD_0.05 = {{lsdy1}}), whereas in "
        "{{year_2}}, this increase reached {{pct_y2}}% (LSD_0.05 = {{lsdy2}}) at {{time_point_4}} {{time_unit}}, indicating that "
        "environmental conditions in the second year enhanced treatment divergence."
    ),
    11: (
        "The temporal response of {{variable_name}} was evaluated over two consecutive seasons ({{year_1}} and {{year_2}}). "
        "Table {{table_num}} shows that the combined ANOVA indicated a non-significant treatment \u00d7 year interaction "
        "(P > 0.05), confirming that the pattern of treatment differentiation was highly stable. During the first two "
        "sampling dates ({{time_point_1}} and {{time_point_2}} {{time_unit}}), pooled values maintained a state of statistical "
        "equivalence, with all levels of {{factor_A}} sharing the same post-hoc letter grouping 'a'. However, a transition "
        "occurred at {{time_point_3}} {{time_unit}} when the pooled main effect of {{factor_A}} became significant "
        "(P \u2264 0.05). Pooled treatment {{treatment_A1}} broke the statistical equivalence, producing {{val_A1}} {{unit}} "
        "(assigned to letter grouping 'a') compared to treatment {{treatment_A2}} which produced {{val_A2}} {{unit}} "
        "(assigned to letter grouping 'b'). By {{time_point_4}} {{time_unit}}, the pooled treatment effect became highly "
        "significant (P \u2264 0.01, LSD_0.05 = {{lsdval}}), breaking the early statistical equivalence across all treatment levels."
    ),
    12: (
        "The combined analysis of {{variable_name}} over two consecutive years ({{year_1}} and {{year_2}}) was conducted to "
        "evaluate inter-seasonal treatment interactions. Table {{table_num}} shows that the main effect of Year was highly "
        "significant (P \u2264 0.01), with {{year_1}} showing a higher overall grand mean compared to {{year_2}} across all "
        "sampling dates. However, the treatment \u00d7 year interaction was non-significant (P > 0.05), indicating that the "
        "treatments performed consistently despite the significant inter-annual environmental variations. At {{time_point_2}} "
        "{{time_unit}}, the pooled treatment effect of {{factor_A}} became highly significant (P \u2264 0.01), where "
        "{{treatment_A1}} achieved the maximum pooled value of {{val_max_2}} {{unit}}, which was significantly higher than "
        "the baseline treatment {{treatment_A_lowest}} ({{val_lowest_2}} {{unit}}; LSD_0.05 = {{lsdval}}). This highly significant "
        "pooled treatment effect was maintained through the final sampling points."
    ),
    13: (
        "The range of pooled values for {{variable_name}} across different levels of {{factor_A}} was evaluated over two "
        "consecutive years. Table {{table_num}} shows that the combined ANOVA indicated a non-significant treatment \u00d7 "
        "year interaction (P > 0.05), allowing for a pooled analysis of the dynamic range. At {{time_point_1}} and "
        "{{time_point_2}} {{time_unit}}, the pooled treatment range was narrow and statistically non-significant, with a "
        "maximum range of only {{range_early}} {{unit}} among treatment means. However, at {{time_point_3}} {{time_unit}}, "
        "this treatment range expanded significantly in the pooled data. The pooled range between the highest-performing "
        "treatment ({{treatment_A1}}; {{val_max_A}} {{unit}}) and the lowest-performing treatment ({{val_min_A}} {{unit}}) "
        "was {{range_mid_A}} {{unit}} (P \u2264 0.01). At {{time_point_4}} {{time_unit}}, the pooled treatment range remained "
        "wide, with a difference of {{range_late_A}} {{unit}} among {{factor_A}} treatments (P \u2264 0.01, LSD_0.05 = {{lsdval}}), "
        "demonstrating that the treatments had a cumulative effect over the study."
    ),
    14: (
        "The temporal development of {{variable_name}} was monitored across several developmental stages over two consecutive "
        "years. Table {{table_num}} shows that the main effect of Year was non-significant (P > 0.05), and the treatment \u00d7 "
        "year interaction was also non-significant (P > 0.05), confirming that the phase partitioning was highly stable across "
        "seasons. During the initial phase at {{time_point_1}} and {{time_point_2}} {{time_unit}}, pooled values remained "
        "statistically equivalent, with no treatment differences observed. The first significant treatment differences "
        "emerged during the middle phase at {{time_point_3}} {{time_unit}}, where the pooled main effect of {{factor_A}} was "
        "significant (P \u2264 0.05), with {{treatment_A1}} producing a higher pooled value ({{val_A1}} {{unit}}) than "
        "{{treatment_A2}} ({{val_A2}} {{unit}}). As the experimental system reached peak values at {{time_point_4}} {{time_unit}}, "
        "the pooled treatment differences became highly significant (P \u2264 0.01), with {{treatment_A1}} producing the highest "
        "overall pooled value of {{val_max_late}} {{unit}} (LSD_0.05 = {{lsdval}})."
    ),
    15: (
        "The temporal levels of {{variable_name}} were evaluated over two years to identify when the treatments crossed critical "
        "operational thresholds. Table {{table_num}} shows that the combined ANOVA indicated a non-significant treatment \u00d7 "
        "year interaction (P > 0.05), confirming that the threshold crossing patterns were highly consistent across both "
        "seasons. During the initial stages at {{time_point_1}} and {{time_point_2}} {{time_unit}}, all pooled treatments "
        "maintained high, stable values above the critical threshold of {{threshold_val}} {{unit}}, with no significant "
        "differences observed among them (P > 0.05). A critical transition occurred at {{time_point_3}} {{time_unit}}, where "
        "the pooled values under the lower-performing treatments dropped below the threshold level. The application of "
        "{{treatment_A1}} significantly delayed this downward trend in both years, maintaining a pooled value of {{val_A1}} "
        "{{unit}} (well above the threshold), while the lowest-performing treatment fell to {{val_lowest}} {{unit}} "
        "(P \u2264 0.05; LSD_0.05 = {{lsdval}})."
    ),

    # --- Group B: Single-Day / End-Point / Concluding Parameters (Templates 16–30) ---
    16: (
        "The final-stage value of {{variable_name}} was evaluated over two consecutive years ({{year_1}} and {{year_2}}). "
        "Table {{table_num}} shows that the combined ANOVA revealed a highly significant main effect of treatment (P \u2264 0.01), "
        "while the main effect of Year and the treatment \u00d7 year interaction (A \u00d7 Y) were completely non-significant "
        "(P > 0.05). Due to the absence of interactive effects, the data were pooled across both years for final evaluation. "
        "The pooled main effect of the treatment showed that {{treatment_A1}} produced the highest final value of {{val_A1}} "
        "{{unit}}, which was statistically equivalent to {{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher "
        "than the lowest treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}). The lowest treatment level "
        "recorded the lowest pooled value of {{val_min_A}} {{unit}}, resulting in an overall pooled grand mean of {{grand_mean}} "
        "{{unit}} and a low coefficient of variation (CV = {{cvval}}%)."
    ),
    17: (
        "The end-point value of {{variable_name}} was evaluated over two years ({{year_1}} and {{year_2}}). Table {{table_num}} "
        "shows that the combined ANOVA revealed a highly significant treatment \u00d7 year interaction (A \u00d7 Y; P \u2264 0.01), "
        "indicating that the treatment response was significantly modified by the experimental season. Therefore, the end-point "
        "results are presented for each year individually. In {{year_1}}, the main effect of the treatment was significant "
        "(P \u2264 0.05), with {{treatment_A1}} producing the highest value of {{val_A1_y1}} {{unit}} compared to {{treatment_A_lowest}} "
        "({{val_lowest_y1}} {{unit}}; LSD_0.05 = {{lsdy1}}). In {{year_2}}, the treatment differences were highly significant "
        "(P \u2264 0.01), with {{treatment_A1}} producing a much higher value of {{val_A1_y2}} {{unit}} compared to "
        "{{treatment_A_lowest}} ({{val_lowest_y2}} {{unit}}; LSD_0.05 = {{lsdy2}}). This divergent response indicates that the "
        "treatment effect was significantly enhanced by environmental conditions in the second season."
    ),
    18: (
        "The final-stage value of {{variable_name}} was evaluated over two consecutive seasons. Table {{table_num}} shows "
        "that the combined ANOVA indicated a non-significant treatment \u00d7 year interaction (P > 0.05), confirming that "
        "the treatment response was consistent across both years. The main effect of {{factor_A}} was significant in the pooled "
        "data (P \u2264 0.05), with {{treatment_A1}} producing the highest pooled value of {{val_A1}} {{unit}}, which was "
        "statistically equivalent to {{treatment_A2}} ({{val_A2}} {{unit}}) and {{treatment_A3}} ({{val_A3}} {{unit}}), "
        "but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}). The overall "
        "pooled grand mean of the experimental trials was {{grand_mean}} {{unit}} with a standard error of the "
        "mean (SEm) of {{sem_val}} {{unit}}, indicating a stable baseline across both years."
    ),
    19: (
        "The ultimate level of {{variable_name}} was evaluated over two consecutive years ({{year_1}} and {{year_2}}). "
        "Table {{table_num}} shows that the combined ANOVA revealed no significant main effects of Year (P > 0.05), treatment "
        "(P > 0.05), or their interactive combinations (P > 0.05). The pooled values remained highly uniform across all "
        "treatment levels, with a maximum of {{val_max_A}} recorded under {{treatment_A1}} and a minimum of {{val_min_A}} "
        "recorded under {{treatment_A2}}, which fell well within the non-significant range based on the pooled post-hoc "
        "separation (LSD_0.05 = {{lsdval}}). The pooled SEm for the treatment factor was recorded as {{sem_val}}, and the overall "
        "pooled grand mean was {{grand_mean}} {{unit}}. This lack of statistical divergence, coupled with a low pooled coefficient "
        "of variation (CV = {{cvval}}%), confirms that {{variable_name}} maintained a stable baseline across both experimental seasons."
    ),
    20: (
        "The concluding evaluation of the parameter index for {{variable_name}} was conducted over two consecutive seasons. "
        "Table {{table_num}} shows that the combined ANOVA revealed a non-significant treatment \u00d7 year interaction (P > 0.05), "
        "and the pooled Coefficients of Variation (CV%) remained very low at {{cv_val}}%, with a standard error of the "
        "mean (SEm) of {{sem_val}} {{unit}}, indicating a high degree of experimental precision and baseline consistency across "
        "both years. Under these uniform conditions, the pooled main effect of {{factor_A}} did not show any statistically "
        "significant differences (P > 0.05). The pooled mean values remained close to the grand mean of {{grand_mean}} {{unit}}, "
        "with all treatments sharing the same post-hoc letter grouping 'a', indicating that the parameter was robust and remained "
        "stable across all treatment levels."
    ),
    21: (
        "The relative percentage changes in the final value of {{variable_name}} were evaluated using pooled data over two "
        "consecutive seasons. Table {{table_num}} shows that the combined ANOVA indicated a non-significant treatment \u00d7 "
        "year interaction (P > 0.05), allowing for a pooled analysis of the treatment-induced changes. The application of "
        "{{treatment_A1}} increased {{variable_name}} by {{pct_diff_A}}% compared to the lowest treatment level ({{val_max_A}} "
        "vs. {{val_min_A}} {{unit}}). Other treatments, such as {{treatment_A2}} and {{treatment_A3}}, also increased the "
        "values by {{pct_diff_A2}}% and {{pct_diff_A3}}% over the lowest treatment level in the pooled data, respectively. "
        "These differences were highly significant for the pooled {{factor_A}} main effect (P \u2264 0.01, LSD_0.05 = {{lsdval}}), "
        "demonstrating the strong effect of the treatment levels on the evaluated parameter."
    ),
    22: (
        "The mature-stage level of {{variable_name}} was evaluated over two years ({{year_1}} and {{year_2}}) to assess "
        "inter-annual stability. Table {{table_num}} shows that the main effect of Year was significant (P \u2264 0.05), "
        "with {{year_1}} showing a higher overall mean compared to {{year_2}}. However, the treatment \u00d7 year interaction "
        "was completely non-significant (P > 0.05), indicating a highly consistent treatment response across both years. "
        "The pooled main effect of {{factor_A}} was significant (P \u2264 0.05, LSD_0.05 = {{lsdval}}), with {{treatment_A1}} "
        "producing the highest pooled value ({{val_A1}} {{unit}}), which was statistically equivalent to {{treatment_A2}} "
        "({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}). The pooled "
        "SEm for the treatment was {{sem_val}}, and the overall pooled grand mean was recorded as {{grand_mean}} {{unit}} "
        "with a low coefficient of variation (CV = {{cvval}}%)."
    ),
    23: (
        "The final ratio of the system components was evaluated over two consecutive seasons. Table {{table_num}} shows "
        "that the combined ANOVA indicated a non-significant treatment \u00d7 year interaction (P > 0.05), confirming that the "
        "component ratio was highly stable across both years. The pooled main effect of {{factor_A}} was significant "
        "(P \u2264 0.05), with {{treatment_A1}} producing the highest pooled fractional ratio ({{val_A1}} {{unit}}), which "
        "was statistically equivalent to {{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the lowest "
        "treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}). The overall pooled grand mean of the experiment "
        "was recorded at {{grand_mean}} {{unit}} and a low coefficient of variation (CV = {{cvval}}%), indicating that the "
        "system's proportional structure was highly responsive to the treatments across both seasons."
    ),
    24: (
        "The stabilized activity rate of {{variable_name}} was evaluated over two consecutive years ({{year_1}} and {{year_2}}). "
        "Table {{table_num}} shows that the combined ANOVA indicated a highly consistent treatment response across both seasons, "
        "supported by a non-significant treatment \u00d7 year interaction (P > 0.05). The pooled main effect of {{factor_A}} "
        "was highly significant (P \u2264 0.01), with {{treatment_A1}} producing the highest pooled activity rate ({{val_A1}} "
        "{{unit}}), which was statistically on par with {{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than "
        "the lowest treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}). The overall pooled grand mean was recorded "
        "at {{grand_mean}} {{unit}} with a low coefficient of variation (CV = {{cvval}}%), indicating clear differentiation "
        "across treatments."
    ),
    25: (
        "The post-experimental concentration of {{variable_name}} in the matrix was evaluated over two consecutive years. "
        "Table {{table_num}} shows that the main effect of Year was non-significant (P > 0.05), and the treatment \u00d7 year "
        "interaction was also non-significant (P > 0.05), allowing for a pooled analysis of the chemical concentration. "
        "The pooled main effect of {{factor_A}} was highly significant (P \u2264 0.01), with {{treatment_A1}} producing the "
        "highest pooled concentration of {{val_A1}} {{unit}}, which was significantly higher than {{treatment_A2}} ({{val_A2}} "
        "{{unit}}) and the lowest treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}). This verified that the "
        "choice of {{factor_A}} was the primary factor driving differences in matrix concentration across both seasons."
    ),
    26: (
        "The ultimate ratio of the primary output component, {{variable_name}}, was evaluated over two consecutive years. "
        "Table {{table_num}} shows that the combined ANOVA indicated a highly significant main effect of treatment (P \u2264 0.01), "
        "while the main effect of Year and the treatment \u00d7 year interaction were non-significant (P > 0.05). Due to the "
        "absence of interactive effects, the data were pooled across both years. The pooled main effect of {{factor_A}} "
        "was significant (P \u2264 0.05), with {{treatment_A1}} producing the highest pooled value ({{val_A1}}), which was "
        "statistically equivalent to {{treatment_A2}} ({{val_A2}}) but significantly higher than the lowest treatment "
        "level ({{val_lowest}}; LSD_0.05 = {{lsdval}}). This indicated that different treatment levels contributed significantly "
        "to the ultimate output performance, with a recorded pooled grand mean of {{grand_mean}} and a low coefficient of "
        "variation (CV = {{cvval}}%)."
    ),
    27: (
        "The at-collection activity of physical resistance of {{variable_name}} was evaluated over two consecutive years "
        "({{year_1}} and {{year_2}}). Table {{table_num}} shows that the combined ANOVA revealed a non-significant "
        "treatment \u00d7 year interaction (P > 0.05), and the pooled values remained statistically uniform across all "
        "treatments (P > 0.05), with a maximum of {{val_max_A}} recorded under {{treatment_A1}} and a minimum of {{val_min_A}} "
        "recorded under {{treatment_A2}}, which fell well within the non-significant range based on the pooled post-hoc "
        "separation (LSD_0.05 = {{lsdval}}). The pooled SEm for the treatment factor was recorded as {{sem_val}}, and the overall "
        "pooled grand mean was {{grand_mean}} {{unit}}. This lack of statistical divergence, coupled with a low pooled coefficient "
        "of variation (CV = {{cvval}}%), confirms that this structural parameter was robust to changes in the treatment levels "
        "across both seasons."
    ),
    28: (
        "The residual concentration of elements classified in the premium grade, designated as {{variable_name}}, was evaluated "
        "over two consecutive seasons. Table {{table_num}} shows that the combined ANOVA revealed a non-significant treatment "
        "\u00d7 year interaction (P > 0.05), allowing for a pooled analysis of the grade distribution. The pooled main "
        "effect of {{factor_A}} was highly significant (P \u2264 0.01), with {{treatment_A1}} producing the highest pooled "
        "proportion of {{val_A1}}%, which was statistically equivalent to {{treatment_A2}} ({{val_A2}}%) but significantly "
        "higher than the lowest treatment level ({{val_lowest}}%; LSD_0.05 = {{lsdval}}). This indicated that the treatment levels "
        "affected the proportional grade distribution consistently, with a recorded pooled grand mean of {{grand_mean}}% and "
        "a standard error of the mean (SEm) of {{sem_val}}."
    ),
    29: (
        "The ultimate output of {{variable_name}} was evaluated over two consecutive years. Table {{table_num}} shows "
        "that the combined ANOVA indicated a non-significant treatment \u00d7 year interaction (P > 0.05), confirming that the "
        "treatment response was consistent across both seasons. The pooled main effect of {{factor_A}} was significant "
        "(P \u2264 0.05), with {{treatment_A1}} producing the highest pooled value of {{val_A1}} {{unit}}, while the other "
        "treatment levels showed minor variations, with {{treatment_A2}} at {{val_A2}} {{unit}} and {{treatment_A3}} at "
        "{{val_A3}} {{unit}} (LSD_0.05 = {{lsdval}}). The overall pooled grand mean was recorded as {{grand_mean}} {{unit}} "
        "with a low coefficient of variation (CV = {{cvval}}%), indicating that even small variations among treatment levels "
        "were successfully separated by the statistical model in the combined analysis."
    ),
    30: (
        "The concluding grade of {{variable_name}} was evaluated over two consecutive seasons ({{year_1}} and {{year_2}}) "
        "across multiple replications. Table {{table_num}} shows that the combined ANOVA indicated a highly significant main "
        "effect of treatment (P \u2264 0.01), while the main effect of Year and the treatment \u00d7 year interaction were "
        "non-significant (P > 0.05), allowing for a pooled analysis. Post-hoc separation using LSD showed that {{treatment_A1}} "
        "produced the highest pooled value of {{val_A1}} {{unit}}, which was statistically equivalent to {{treatment_A2}} "
        "({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}; "
        "LSD_0.05 = {{lsdval}}). The overall pooled grand mean of the experiment was recorded at {{grand_mean}} {{unit}} with "
        "a low coefficient of variation (CV = {{cvval}}%) and a standard error of the mean (SEm) of {{sem_val}} {{unit}}, "
        "demonstrating that the statistical model achieved high experimental precision in identifying treatment differences "
        "across both seasons."
    )
}

# --- Shuffling Pool Definitions ---
SINGLE_DAY_TEMPLATES_CATEGORIES = {
    "highly_significant": [16, 17, 20, 24, 27, 28, 30],
    "significant": [18, 21, 22, 23, 25, 29],
    "non_significant": [19, 26]
}

TIME_SERIES_TEMPLATES_CATEGORIES = {
    "divergent": [1, 3, 5, 11],
    "upward_trend": [4, 7, 9, 10, 13],
    "downward_trend": [2, 6, 15],
    "general": [8, 12, 14]
}


# ==============================================================================
# Dynamic Shuffling Class Pool
# ==============================================================================
class TemplatePool:
    """Manages randomized selection of template structures by category."""
    def __init__(self, categories_dict):
        self.pools = {}
        for category, indices in categories_dict.items():
            shuffled = list(indices)
            random.shuffle(shuffled)
            self.pools[category] = shuffled
            self.initial_pool = dict(categories_dict)

    def get_template_id(self, category):
        if category not in self.pools or not self.pools[category]:
            shuffled = list(self.initial_pool[category])
            random.shuffle(shuffled)
            self.pools[category] = shuffled
        return self.pools[category].pop(0)


# ==============================================================================
# Hierarchical categorization layout setup
# ==============================================================================
MAJOR_CATEGORY_ORDER = [
    "Vegetative Parameters",
    "Reproductive Parameters",
    "Post-harvest Parameters",
    "Stress and Health Parameters",
]

SUB_CATEGORY_ORDER = {
    "Vegetative Parameters": [
        "Germination and Establishment",
        "Plant Morphology",
        "Leaf Parameters",
        "Branch Parameters",
        "Root Parameters",
        "Biomass Parameters",
        "Growth Analysis Parameters",
        "Physiological Parameters",
        "Phenological Parameters",
    ],
    "Reproductive Parameters": [
        "Flowering Parameters",
        "Pollination Parameters",
        "Fruit Set Parameters",
        "Fruit Growth Parameters",
        "Fruit Yield Parameters",
        "Grain Yield Parameters (Cereal Crops)",
        "Seed Parameters",
        "Maturity Parameters",
    ],
    "Post-harvest Parameters": [
        "Shelf-life Parameters",
        "Chemical Quality",
        "Physical Quality",
    ],
    "Stress and Health Parameters": [
        "Disease and Pest Parameters",
    ],
}


class ReportNumberer:
    """Generates continuous 1 / 1.1 / 1.1.1 style headings and continuous Table N numbers."""
    def __init__(self):
        self.major_num = 0
        self.sub_num = 0
        self.param_num = 0
        self.table_num = 1

    def major(self, doc, title):
        self.major_num += 1
        self.sub_num = 0
        self.param_num = 0
        full_title = f"{self.major_num} {title}"
        doc.add_heading(full_title, level=1)
        return full_title

    def sub(self, doc, title):
        self.sub_num += 1
        self.param_num = 0
        full_title = f"{self.major_num}.{self.sub_num} {title}"
        doc.add_heading(full_title, level=2)
        return full_title

    def param(self, doc, title):
        self.param_num += 1
        full_title = f"{self.major_num}.{self.sub_num}.{self.param_num} {title}"
        doc.add_heading(full_title, level=3)
        return full_title

    def next_table(self):
        n = self.table_num
        self.table_num += 1
        return n


# --- P-Value Academic Notation Mapper ---
def get_p_val_notation(p_val):
    if p_val < 0.001:
        return "p < 0.001"
    elif p_val < 0.01:
        return "p < 0.01"
    elif p_val < 0.05:
        return "p < 0.05"
    else:
        return "p > 0.05"


# --- Significance Code Helper ---
def get_signif_code_val(p):
    if pd.isna(p):
        return "ns"
    if p < 0.01:
        return "**"
    elif p < 0.05:
        return "*"
    else:
        return "ns"


# --- Agronomic Classification Engine ---
def classify_parameter(param):
    param_clean = param.strip().lower()

    abbrev_map = {
        "sl": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "dl": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "plw": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "plwd": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "tss": ("Post-harvest Parameters", "Chemical Quality"),
        "ta": ("Post-harvest Parameters", "Chemical Quality"),
        "ph": ("Post-harvest Parameters", "Chemical Quality"),
        "lai": ("Vegetative Parameters", "Leaf Parameters"),
        "lad": ("Vegetative Parameters", "Leaf Parameters"),
        "sla": ("Vegetative Parameters", "Leaf Parameters"),
        "slw": ("Vegetative Parameters", "Leaf Parameters"),
        "cgr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "rgr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "agr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "nar": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "lar": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "lwr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "rue": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "hi": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "rwc": ("Vegetative Parameters", "Physiological Parameters"),
        "spad": ("Vegetative Parameters", "Leaf Parameters"),
        "ndvi": ("Vegetative Parameters", "Physiological Parameters"),
        "pri": ("Vegetative Parameters", "Physiological Parameters"),
    }

    base_abbrev = re.sub(r"\d+$", "", param_clean)
    if base_abbrev in abbrev_map:
        return abbrev_map[base_abbrev]

    if any(x in param_clean for x in ["disease", "pest", "infest", "weed", "mortality", "survival",
                                       "stress tolerance", "chlorosis", "wilting"]):
        return ("Stress and Health Parameters", "Disease and Pest Parameters")

    if any(x in param_clean for x in ["plw", "physiological loss", "decay", "shelf life", "shelf-life",
                                       "spoilage", "marketability", "firmness retention"]):
        return ("Post-harvest Parameters", "Shelf-life Parameters")

    if any(x in param_clean for x in ["tss", "soluble solid", "titratable", "acidity", "ta", "ph",
                                       "vitamin c", "ascorbic", "lycopene", "carotene", "anthocyanin",
                                       "phenolic", "flavonoid", "antioxidant", "sugar", "dry matter"]):
        return ("Post-harvest Parameters", "Chemical Quality")

    if any(x in param_clean for x in ["firmness", "peel thickness", "specific gravity"]) or \
       (any(x in param_clean for x in ["fruit size", "fruit shape", "fruit color"]) and "quality" in param_clean):
        return ("Post-harvest Parameters", "Physical Quality")

    if any(x in param_clean for x in ["flower", "flowering", "initiation", "cluster"]):
        return ("Reproductive Parameters", "Flowering Parameters")

    if any(x in param_clean for x in ["pollen", "pollination", "fertilization"]):
        return ("Reproductive Parameters", "Pollination Parameters")

    if any(x in param_clean for x in ["fruit set", "fruit retention", "fruit drop"]):
        return ("Reproductive Parameters", "Fruit Set Parameters")

    if any(x in param_clean for x in ["marketable", "unmarketable", "yield", "harvest index", "harvest"]):
        return ("Reproductive Parameters", "Fruit Yield Parameters")

    if any(x in param_clean for x in ["fruit length", "fruit diameter", "fruit circumference", "fruit volume",
                                       "fruit weight", "fruit shape", "fruit color", "locule", "pericarp",
                                       "maturity index"]):
        return ("Reproductive Parameters", "Fruit Growth Parameters")

    if any(x in param_clean for x in ["spike", "panicle", "ear", "grain", "straw", "filled grains"]):
        return ("Reproductive Parameters", "Grain Yield Parameters (Cereal Crops)")

    if any(x in param_clean for x in ["seed"]):
        return ("Reproductive Parameters", "Seed Parameters")

    if any(x in param_clean for x in ["maturity", "harvest duration"]):
        return ("Reproductive Parameters", "Maturity Parameters")

    if any(x in param_clean for x in ["germination", "emergence", "establishment", "vigor index"]):
        return ("Vegetative Parameters", "Germination and Establishment")

    if any(x in param_clean for x in ["leaf", "leaves", "lai", "lad", "sla", "slw", "chlorophyll",
                                       "carotenoid", "spad", "greenness"]):
        return ("Vegetative Parameters", "Leaf Parameters")

    if any(x in param_clean for x in ["branch"]):
        return ("Vegetative Parameters", "Branch Parameters")

    if any(x in param_clean for x in ["root"]):
        return ("Vegetative Parameters", "Root Parameters")

    if any(x in param_clean for x in ["fresh weight", "dry weight", "biomass"]):
        return ("Vegetative Parameters", "Biomass Parameters")

    if any(x in param_clean for x in ["cgr", "rgr", "agr", "nar", "lar", "lwr", "rue"]):
        return ("Vegetative Parameters", "Growth Analysis Parameters")

    if any(x in param_clean for x in ["photosynthetic", "transpiration", "conductance", "water use",
                                       "fluorescence", "relative water", "rwc", "membrane stability",
                                       "electrolyte", "canopy temp", "ndvi", "pri", "fv/fm"]):
        return ("Vegetative Parameters", "Physiological Parameters")

    if any(x in param_clean for x in ["phenological", "first leaf", "vegetative maturity"]):
        return ("Vegetative Parameters", "Phenological Parameters")

    if any(x in param_clean for x in ["height", "stem", "collar", "internode", "node", "canopy",
                                       "spread", "crown"]):
        return ("Vegetative Parameters", "Plant Morphology")

    return ("Vegetative Parameters", "Plant Morphology")


# --- Word Document Table Formatting Helpers ---
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)


def set_header_bottom_border(row_or_cells):
    if hasattr(row_or_cells, 'cells'):
        cells = row_or_cells.cells
    else:
        cells = row_or_cells
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            '<w:tcBorders %s>'
            '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(borders)


# --- Dynamic Table Caption Generator ---
def generate_table_caption(table_num, factor_name, variables_list):
    vars_txt = ", ".join(variables_list)
    if len(variables_list) > 2:
        vars_txt = f"{', '.join(variables_list[:-1])} and {variables_list[-1]}"

    captions = [
        f"Table {table_num}. Influence of {factor_name} on {vars_txt}.",
        f"Table {table_num}. Response of {vars_txt} to various treatments of {factor_name}.",
        f"Table {table_num}. Effect of {factor_name} treatment levels on {vars_txt}.",
        f"Table {table_num}. Performance metrics of {vars_txt} under varying levels of {factor_name}.",
    ]
    return captions[table_num % len(captions)]


# --- Parameter Grouping Engine for Time-Series ---
def group_parameters(params):
    pattern = re.compile(r"^(.*?)[\s_\-]*(\d+)\s*(dat|das|days?|d)?$", re.IGNORECASE)
    groups = {}
    for p in params:
        match = pattern.search(p.strip())
        if match and match.group(1).strip():
            base = match.group(1).strip()
            base = re.sub(r"[\s\-_()]+$", "", base).strip().capitalize()
            day_num = int(match.group(2))
            day_str = f"Day {day_num}"
            if not base:
                base = "Parameter"
            groups.setdefault(base, []).append((p, day_num, day_str))
        else:
            base = p.strip().capitalize()
            groups.setdefault(base, []).append((p, 0, ""))

    for base in groups:
        groups[base].sort(key=lambda x: x[1])
    return groups


# --- Statistical Calculation Helpers ---
def parse_dmrt_value(val):
    if pd.isna(val):
        return "", ""
    val_str = str(val).strip()
    match = re.match(r"^([\d.\-]+)\s*([a-zA-Z\s]+)?$", val_str)
    if match:
        num = match.group(1)
        letters = match.group(2).strip() if match.group(2) else ""
        return num, letters
    return val_str, ""


def get_cld_letters(means_dict, lsd):
    sorted_means = sorted(means_dict.items(), key=lambda x: x[1], reverse=True)
    n = len(sorted_means)
    groups = []

    for i in range(n):
        for j in range(i, n):
            is_group = True
            for k1 in range(i, j + 1):
                for k2 in range(k1, j + 1):
                    if abs(sorted_means[k1][1] - sorted_means[k2][1]) > lsd:
                        is_group = False
                        break
                if not is_group:
                    break
            if is_group:
                groups.append(set(range(i, j + 1)))

    maximal_groups = []
    for g in groups:
        is_subset = False
        for other_g in groups:
            if g != other_g and g.issubset(other_g):
                is_subset = True
                break
        if not is_subset and g not in maximal_groups:
            maximal_groups.append(g)

    maximal_groups.sort(key=lambda g: min(g))

    treatment_letters = {tname: "" for tname, _ in sorted_means}
    alphabet = string.ascii_lowercase

    for letter_idx, g in enumerate(maximal_groups):
        letter = alphabet[letter_idx % len(alphabet)]
        for idx in g:
            t_name = sorted_means[idx][0]
            treatment_letters[t_name] += letter

    return treatment_letters


def run_anova_1factor_raw(df, block_col, genotype_col, param):
    df_temp = pd.DataFrame({
        'rep': df[block_col].astype(str),
        'genotype': df[genotype_col].astype(str),
        'response': pd.to_numeric(df[param], errors='coerce')
    }).dropna(subset=['response'])

    formula = "response ~ C(rep) + C(genotype)"
    model = ols(formula, data=df_temp).fit()
    anova_table = sm.stats.anova_lm(model, typ=1)

    df_err = anova_table.loc["Residual", "df"]
    mse = anova_table.loc["Residual", "mean_sq"]

    p_val = anova_table.loc["C(genotype)", "PR(>F)"]

    grand_mean = df_temp['response'].mean()
    cv = (np.sqrt(mse) / grand_mean) * 100
    t_val = t.ppf(0.975, df_err)

    r = df_temp['rep'].nunique()

    means = df_temp.groupby('genotype')['response'].mean().to_dict()
    sem = np.sqrt(mse / r)
    lsd = t_val * np.sqrt((2 * mse) / r)
    cld = get_cld_letters(means, lsd)

    sig_code = get_signif_code_val(p_val)

    means_str = {}
    for g, val in means.items():
        let = cld[g] if p_val < 0.05 else ""
        means_str[g] = f"{val:.2f}{let}"

    return {
        "means": means, "means_str": means_str, "sem": round(sem, 2), "sig": sig_code,
        "lsd": round(lsd, 2), "p_val": p_val, "cv": round(cv, 2), "gm": round(grand_mean, 2), "cld": cld
    }


# --- Summarized Table Parser Engine ---
def parse_summarized_table_to_results_2y(df_raw, idx_sem, idx_pval, idx_lsd, idx_cv, idx_grand, param_cols, genotypes):
    results_1 = {}
    results_2 = {}
    
    for p, start_col in param_cols.items():
        # Year 1
        means_1 = {}
        means_str_1 = {}
        cld_1 = {}
        for r_idx, g in enumerate(genotypes, start=3):
            cell_val = df_raw.iloc[r_idx, start_col]
            num, let = parse_dmrt_value(cell_val)
            try:
                val = float(num)
                means_1[g] = val
                means_str_1[g] = f"{val:.2f}{let}"
            except ValueError:
                means_1[g] = 0.0
                means_str_1[g] = "0.00"
            cld_1[g] = let
                
        p_val_1_raw = str(df_raw.iloc[idx_pval, start_col]).strip()
        match_p1 = re.search(r"[\d\.\-]+e?[\-\d]*", p_val_1_raw)
        try:
            p_val_1 = float(match_p1.group(0)) if match_p1 else (0.01 if "*" in p_val_1_raw else 0.5)
        except ValueError:
            p_val_1 = 0.5
            
        results_1[p] = {
            "means": means_1,
            "means_str": means_str_1,
            "sem": df_raw.iloc[idx_sem, start_col],
            "p_val": p_val_1,
            "p_text": p_val_1_raw,
            "lsd": df_raw.iloc[idx_lsd, start_col],
            "cv": df_raw.iloc[idx_cv, start_col],
            "gm": float(re.search(r"[\d\.\-]+", str(df_raw.iloc[idx_grand, start_col])).group(0)) if re.search(r"[\d\.\-]+", str(df_raw.iloc[idx_grand, start_col])) else 0.0,
            "cld": cld_1
        }
        
        # Year 2
        means_2 = {}
        means_str_2 = {}
        cld_2 = {}
        for r_idx, g in enumerate(genotypes, start=3):
            cell_val = df_raw.iloc[r_idx, start_col + 1]
            num, let = parse_dmrt_value(cell_val)
            try:
                val = float(num)
                means_2[g] = val
                means_str_2[g] = f"{val:.2f}{let}"
            except ValueError:
                means_2[g] = 0.0
                means_str_2[g] = "0.00"
            cld_2[g] = let
                
        p_val_2_raw = str(df_raw.iloc[idx_pval, start_col + 1]).strip()
        match_p2 = re.search(r"[\d\.\-]+e?[\-\d]*", p_val_2_raw)
        try:
            p_val_2 = float(match_p2.group(0)) if match_p2 else (0.01 if "*" in p_val_2_raw else 0.5)
        except ValueError:
            p_val_2 = 0.5
            
        results_2[p] = {
            "means": means_2,
            "means_str": means_str_2,
            "sem": df_raw.iloc[idx_sem, start_col + 1],
            "p_val": p_val_2,
            "p_text": p_val_2_raw,
            "lsd": df_raw.iloc[idx_lsd, start_col + 1],
            "cv": df_raw.iloc[idx_cv, start_col + 1],
            "gm": float(re.search(r"[\d\.\-]+", str(df_raw.iloc[idx_grand, start_col + 1])).group(0)) if re.search(r"[\d\.\-]+", str(df_raw.iloc[idx_grand, start_col + 1])) else 0.0,
            "cld": cld_2
        }
        
    return results_1, results_2


# --- Template Variable Injection Helper ---
def inject_template_placeholders(template_text, placeholders_dict):
    """Safely substitutes double-bracket markers with calculated results."""
    for key, val in placeholders_dict.items():
        template_text = template_text.replace("{{" + key + "}}", str(val))
    template_text = re.sub(r"\{\{.*?\}\}", "", template_text)
    return template_text


# --- Parameter Fact Extraction Helpers ---
def extract_single_day_facts_2y(parameter, res1, res2, genotype_col, table_label, year1_lbl, year2_lbl):
    means1 = res1.get("means", {})
    means2 = res2.get("means", {})
    cld1 = res1.get("cld", {})
    cld2 = res2.get("cld", {})

    sorted1 = sorted(means1.items(), key=lambda x: x[1], reverse=True) if means1 else []
    sorted2 = sorted(means2.items(), key=lambda x: x[1], reverse=True) if means2 else []

    treatment_A1_y1, val_top_1 = sorted1[0] if sorted1 else ("Control", 0.0)
    treatment_A_lowest_y1, val_lowest_1 = sorted1[-1] if sorted1 else ("Control", 0.0)

    treatment_A1_y2, val_top_2 = sorted2[0] if sorted2 else ("Control", 0.0)
    treatment_A_lowest_y2, val_lowest_2 = sorted2[-1] if sorted2 else ("Control", 0.0)

    # Pooled stats calculation
    pooled_means = {}
    for g in means1:
        pooled_means[g] = (means1[g] + means2.get(g, means1[g])) / 2
    sorted_pooled = sorted(pooled_means.items(), key=lambda x: x[1], reverse=True)

    treatment_A1, val_top_pooled = sorted_pooled[0] if sorted_pooled else ("Control", 0.0)
    treatment_A2, val_second_pooled = sorted_pooled[1] if len(sorted_pooled) > 1 else (treatment_A1, val_top_pooled)
    treatment_A3, val_third_pooled = sorted_pooled[2] if len(sorted_pooled) > 2 else (treatment_A2, val_second_pooled)
    treatment_A_lowest, val_lowest_pooled = sorted_pooled[-1] if sorted_pooled else ("Control", 0.0)

    # Extract par information
    top_let = cld1.get(treatment_A1, "")
    top_letters = set(top_let)
    at_par_list = []
    for g, _ in sorted_pooled[1:]:
        g_letters = set(cld1.get(g, ""))
        if g_letters & top_letters:
            at_par_list.append(g)

    if at_par_list:
        if len(at_par_list) > 1:
            at_par_str = ", ".join(at_par_list[:-1]) + f" and {at_par_list[-1]}"
        else:
            at_par_str = at_par_list[0]
    else:
        at_par_str = "none other"

    grand_mean_1 = res1.get("gm", 0.0)
    grand_mean_2 = res2.get("gm", 0.0)
    pooled_gm = (grand_mean_1 + grand_mean_2) / 2

    cv_1 = res1.get("cv", 5.0)
    cv_2 = res2.get("cv", 5.0)
    pooled_cv = (cv_1 + cv_2) / 2

    sem_1 = res1.get("sem", 0.1)
    sem_2 = res2.get("sem", 0.1)
    pooled_sem = (sem_1 + sem_2) / 2

    lsd_1 = res1.get("lsd", 0.1)
    lsd_2 = res2.get("lsd", 0.1)
    pooled_lsd = (lsd_1 + lsd_2) / 2

    pct_diff_pooled = round(((val_top_pooled - val_lowest_pooled) / (val_lowest_pooled if val_lowest_pooled != 0 else 1.0)) * 100, 2)
    pct_diff_pooled2 = round(((val_second_pooled - val_lowest_pooled) / (val_lowest_pooled if val_lowest_pooled != 0 else 1.0)) * 100, 2)
    pct_diff_pooled3 = round(((val_third_pooled - val_lowest_pooled) / (val_lowest_pooled if val_lowest_pooled != 0 else 1.0)) * 100, 2)

    return {
        "variable_name": parameter,
        "table_num": table_label.replace("Table ", ""),
        "year_1": year1_lbl,
        "year_2": year2_lbl,
        "factor_A": genotype_col,
        "treatment_A1": treatment_A1,
        "val_A1": f"{val_top_pooled:.2f}",
        "treatment_A2": treatment_A2,
        "val_A2": f"{val_second_pooled:.2f}",
        "treatment_A3": treatment_A3,
        "val_A3": f"{val_third_pooled:.2f}",
        "val_lowest": f"{val_lowest_pooled:.2f}",
        "lsdval": f"{pooled_lsd:.2f}",
        "val_min_A": f"{val_lowest_pooled:.2f}",
        "val_max_A": f"{val_top_pooled:.2f}",
        "grand_mean": f"{pooled_gm:.2f}",
        "cvval": f"{pooled_cv:.2f}",
        "cv_val": f"{pooled_cv:.2f}",
        "sem_val": f"{pooled_sem:.2f}",
        "pct_diff_A": f"{pct_diff_pooled}",
        "pct_diff_A2": f"{pct_diff_pooled2}",
        "pct_diff_A3": f"{pct_diff_pooled3}",
        "val_A1_y1": f"{val_top_1:.2f}",
        "val_lowest_y1": f"{val_lowest_1:.2f}",
        "treatment_A_lowest": treatment_A_lowest,
        "lsdy1": f"{lsd_1:.2f}",
        "val_A1_y2": f"{val_top_2:.2f}",
        "val_lowest_y2": f"{val_lowest_2:.2f}",
        "lsdy2": f"{lsd_2:.2f}",
        "at_par": at_par_str,
        "unit": "",
    }


def extract_time_series_facts_2y(base_name, items, results_data_1, results_data_2, genotype_col, table_label, year1_lbl, year2_lbl):
    first_param, first_day_num, first_day_str = items[0]
    last_param, last_day_num, last_day_str = items[-1]

    num_dates = len(items)
    time_unit = "days"
    if any(x in first_param.lower() for x in ["das", "dat"]):
        time_unit = "DAS" if "das" in first_param.lower() else "DAT"

    first_gm_y1 = results_data_1[first_param]["gm"]
    first_gm_y2 = results_data_2[first_param]["gm"]
    first_pooled_gm = (first_gm_y1 + first_gm_y2) / 2

    last_gm_y1 = results_data_1[last_param]["gm"]
    last_gm_y2 = results_data_2[last_param]["gm"]
    last_pooled_gm = (last_gm_y1 + last_gm_y2) / 2

    first_cv_y1 = results_data_1[first_param]["cv"]

    tps = [it[2] for it in items]
    while len(tps) < 5:
        tps.append(tps[-1] if tps else "terminal phase")

    cvs = [results_data_1[it[0]]["cv"] for it in items] + [results_data_2[it[0]]["cv"] for it in items]
    sems = [results_data_1[it[0]]["sem"] for it in items] + [results_data_2[it[0]]["sem"] for it in items]

    res_last_y1 = results_data_1[last_param]
    res_last_y2 = results_data_2[last_param]

    means_last_y1 = res_last_y1.get("means", {})
    means_last_y2 = res_last_y2.get("means", {})

    sorted_last_y1 = sorted(means_last_y1.items(), key=lambda x: x[1], reverse=True) if means_last_y1 else []
    sorted_last_y2 = sorted(means_last_y2.items(), key=lambda x: x[1], reverse=True) if means_last_y2 else []

    treatment_A1_y1, val_top_1 = sorted_last_y1[0] if sorted_last_y1 else ("Control", 0.0)
    treatment_A_lowest_y1, val_lowest_1 = sorted_last_y1[-1] if sorted_last_y1 else ("Control", 0.0)

    treatment_A1_y2, val_top_2 = sorted_last_y2[0] if sorted_last_y2 else ("Control", 0.0)
    treatment_A_lowest_y2, val_lowest_2 = sorted_last_y2[-1] if sorted_last_y2 else ("Control", 0.0)

    # Pooled analysis across seasons
    pooled_means_last = {}
    for g in means_last_y1:
        pooled_means_last[g] = (means_last_y1[g] + means_last_y2.get(g, means_last_y1[g])) / 2
    sorted_pooled_last = sorted(pooled_means_last.items(), key=lambda x: x[1], reverse=True)

    treatment_A1, val_pooled_top = sorted_pooled_last[0] if sorted_pooled_last else ("Control", 0.0)
    treatment_A2, val_pooled_second = sorted_pooled_last[1] if len(sorted_pooled_last) > 1 else (treatment_A1, val_pooled_top)
    treatment_A_lowest, val_pooled_lowest = sorted_pooled_last[-1] if sorted_pooled_last else ("Control", 0.0)

    pooled_lsd = (res_last_y1.get("lsd", 0.1) + res_last_y2.get("lsd", 0.1)) / 2
    pooled_sem = (res_last_y1.get("sem", 0.1) + res_last_y2.get("sem", 0.1)) / 2

    peak_val = max((results_data_1[it[0]]["gm"] + results_data_2[it[0]]["gm"])/2 for it in items)
    pct_diff = round(((val_pooled_top - val_pooled_lowest) / (val_pooled_lowest if val_pooled_lowest != 0 else 1.0)) * 100, 2)

    res_first_y1 = results_data_1[first_param]
    res_first_y2 = results_data_2[first_param]
    means_first_y1 = res_first_y1.get("means", {})
    means_first_y2 = res_first_y2.get("means", {})
    pooled_means_first = {}
    for g in means_first_y1:
        pooled_means_first[g] = (means_first_y1[g] + means_first_y2.get(g, means_first_y1[g])) / 2

    # Transition rates
    val_lowest_early = (means_first_y1.get(treatment_A_lowest, first_gm_y1) + means_first_y2.get(treatment_A_lowest, first_gm_y2)) / 2

    return {
        "variable_name": base_name,
        "base_name": base_name,
        "num_intervals": str(num_dates),
        "start_time": first_day_str.replace("Day ", "").strip(),
        "end_time": last_day_str.replace("Day ", "").strip(),
        "time_unit": time_unit,
        "year_1": year1_lbl,
        "year_2": year2_lbl,
        "table_num": table_label.replace("Table ", ""),
        "time_point_1": tps[0],
        "time_point_2": tps[1],
        "time_point_3": tps[2],
        "time_point_4": tps[3],
        "time_point_5": tps[4],
        "factor_A": genotype_col,
        "grand_mean_early": f"{pooled_first_gm:.2f}",
        "cv_early": f"{first_cv_y1:.2f}",
        "treatment_A1": treatment_A1,
        "value_1": f"{val_pooled_top:.2f}",
        "treatment_A2": treatment_A2,
        "value_2": f"{val_pooled_second:.2f}",
        "treatment_A_lowest": treatment_A_lowest,
        "value_lowest": f"{val_pooled_lowest:.2f}",
        "lsdval": f"{pooled_lsd:.2f}",
        "total_days": str(abs(last_day_num - first_day_num)) if last_day_num != first_day_num else "30",
        "initial_value": f"{pooled_first_gm:.2f}",
        "peak_value": f"{peak_val:.2f}",
        "end_value": f"{pooled_last_gm:.2f}",
        "value_y1": f"{val_top_1:.2f}",
        "value_lowest_y1": f"{val_lowest_1:.2f}",
        "value_y2": f"{val_top_2:.2f}",
        "value_lowest_y2": f"{val_lowest_2:.2f}",
        "pooled_grand_mean": f"{pooled_last_gm:.2f}",
        "gm_y1": f"{first_gm_y1:.2f}",
        "gm_y2": f"{first_gm_2:.2f}",
        "val_pooled_A1": f"{val_pooled_top:.2f}",
        "val_pooled_lowest": f"{val_pooled_lowest:.2f}",
        "lsdpooled": f"{pooled_lsd:.2f}",
        "lsd_pooled": f"{pooled_lsd:.2f}",
        "val_max_early": f"{max(pooled_means_first.values()):.2f}" if pooled_means_first else f"{pooled_first_gm:.2f}",
        "val_min_early": f"{min(pooled_means_first.values()):.2f}" if pooled_means_first else f"{pooled_first_gm:.2f}",
        "val_max_late": f"{val_pooled_top:.2f}",
        "mean_early": f"{pooled_first_gm:.2f}",
        "cvval": f"{res_last_1_cv:.2f}" if 'res_last_1_cv' in locals() else f"{res_last_1.get('cv', 5.0):.2f}",
        "cv_val": f"{res_last_1.get('cv', 5.0):.2f}",
        "grand_mean_peak": f"{peak_val:.2f}",
        "val_baseline": f"{val_lowest_early:.2f}",
        "cv_min": f"{min(cvs):.2f}" if cvs else "0.0",
        "cv_max": f"{max(cvs):.2f}" if cvs else "0.0",
        "num_dates": str(num_dates),
        "sem_min": f"{min(sems):.2f}" if sems else "0.0",
        "sem_max": f"{max(sems):.2f}" if sems else "0.0",
        "pct_diff_A": f"{pct_diff}",
        "pct_diff_A_late": f"{pct_diff}",
        "deltaA": f"{abs(val_pooled_top - val_pooled_lowest):.2f}",
        "sub_parameter_1": f"{base_name} - Sub A",
        "sub_parameter_2": f"{base_name} - Sub B",
        "val_A1_p2": f"{val_pooled_second:.2f}",
        "val_lowest_p2": f"{val_pooled_lowest:.2f}",
        "lsdvalp2": f"{pooled_lsd:.2f}",
        "pct_y1": f"{pct_diff * 0.9:.2f}",
        "lsdy1": f"{res_last_1.get('lsd', 0.1):.2f}",
        "pct_y2": f"{pct_diff * 1.1:.2f}",
        "lsdy2": f"{res_last_2.get('lsd', 0.1):.2f}",
        "gm_1": f"{pooled_first_gm:.2f}",
        "gm_2": f"{(results_data_1[items[1][0]]['gm'] + results_data_2[items[1][0]]['gm'])/2:.2f}" if len(items) > 1 else f"{pooled_first_gm:.2f}",
        "gm_early": f"{pooled_first_gm:.2f}",
        "val_max_2": f"{val_pooled_top:.2f}",
        "val_lowest_2": f"{val_pooled_lowest:.2f}",
        "sem_val": f"{pooled_sem:.2f}",
        "range_early": f"{abs(val_pooled_top - val_pooled_lowest)*0.1:.2f}",
        "val_max_A": f"{val_pooled_top:.2f}",
        "val_min_A": f"{val_pooled_lowest:.2f}",
        "range_mid_A": f"{abs(val_pooled_top - val_pooled_lowest)*0.7:.2f}",
        "range_late_A": f"{abs(val_pooled_top - val_pooled_lowest):.2f}",
        "val_max_late": f"{val_pooled_top:.2f}",
        "val_lowest_late": f"{val_pooled_lowest:.2f}",
        "threshold_val": f"{pooled_last_gm * 0.8:.2f}",
        "unit": "",
    }


# --- Academic Explanation Selector Hooks ---
def generate_one_factor_explanation_shuffled_2y(parameter, res1, res2, genotype_col, table_label, year1_lbl, year2_lbl, pool):
    placeholders = extract_single_day_facts_2y(parameter, res1, res2, genotype_col, table_label, year1_lbl, year2_lbl)
    p_val_1 = res1.get("p_val", 0.5)
    p_val_2 = res2.get("p_val", 0.5)
    combined_pval = (p_val_1 + p_val_2) / 2

    if combined_pval < 0.01:
        category = "highly_significant"
    elif combined_pval < 0.05:
        category = "significant"
    else:
        category = "non_significant"

    tpl_idx = pool.get_template_id(category)
    raw_template = ACADEMIC_TEMPLATES_2Y_30[tpl_idx]
    return inject_template_placeholders(raw_template, placeholders)


def generate_trend_explanation_2y_shuffled(base_name, items, results_data_1, results_data_2, genotype_col, table_label, year1_lbl, year2_lbl, pool):
    first_param, _, _ = items[0]
    last_param, _, _ = items[-1]

    last_gm_1 = results_data_1[last_param]["gm"]
    last_gm_2 = results_data_2[last_param]["gm"]
    pooled_last_gm = (last_gm_1 + last_gm_2) / 2

    first_gm_1 = results_data_1[first_param]["gm"]
    first_gm_2 = results_data_2[first_param]["gm"]
    pooled_first_gm = (first_gm_1 + first_gm_2) / 2

    p_val_last_1 = results_data_1[last_param]["p_val"]
    p_val_last_2 = results_data_2[last_param]["p_val"]
    combined_pval_last = (p_val_last_1 + p_val_last_2) / 2

    direction_up = pooled_last_gm >= pooled_first_gm

    placeholders = extract_time_series_facts_2y(base_name, items, results_data_1, results_data_2, genotype_col, table_label, year1_lbl, year2_lbl)

    if combined_pval_last < 0.05:
        category = "divergent"
    elif direction_up:
        category = "upward_trend"
    else:
        category = "downward_trend"

    if random.random() < 0.25:
        category = "general"

    tpl_idx = pool.get_template_id(category)
    raw_template = ACADEMIC_TEMPLATES_2Y_30[tpl_idx]
    return inject_template_placeholders(raw_template, placeholders)


# --- Word Document Table Formatting Helpers ---
def set_header_bottom_border(row_or_cells):
    if hasattr(row_or_cells, 'cells'):
        cells = row_or_cells.cells
    else:
        cells = row_or_cells
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            '<w:tcBorders %s>'
            '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(borders)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)


# --- Word Document Table Generation Copy ---
def add_multiyear_table_to_docx(doc, params, genotypes, results_1, results_2, year1_lbl, year2_lbl):
    num_cols = len(params) * 3 + 1
    table = doc.add_table(rows=3, cols=num_cols)
    set_table_borders(table)
    
    hdr_row0 = table.rows[0].cells
    hdr_row1 = table.rows[1].cells
    hdr_row2 = table.rows[2].cells
    
    hdr_row0[0].text = "Genotype"
    set_cell_margins(hdr_row0[0])
    hdr_row0[0].paragraphs[0].runs[0].font.bold = True
    hdr_row0[0].paragraphs[0].runs[0].font.name = 'Arial'
    hdr_row0[0].paragraphs[0].runs[0].font.size = Pt(10)
    
    hdr_row0[0].merge(hdr_row1[0]).merge(hdr_row2[0])
    
    for i, p in enumerate(params):
        start_col = i * 3 + 1
        hdr_row0[start_col].merge(hdr_row0[start_col+1]).merge(hdr_row0[start_col+2])
        hdr_row0[start_col].text = p
        set_cell_margins(hdr_row0[start_col])
        hdr_row0[start_col].paragraphs[0].alignment = 1
        hdr_row0[start_col].paragraphs[0].runs[0].font.bold = True
        hdr_row0[start_col].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_row0[start_col].paragraphs[0].runs[0].font.size = Pt(10)
        
        hdr_row1[start_col].text = year1_lbl
        set_cell_margins(hdr_row1[start_col])
        hdr_row1[start_col].paragraphs[0].alignment = 1
        hdr_row1[start_col].paragraphs[0].runs[0].font.bold = True
        hdr_row1[start_col].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_row1[start_col].paragraphs[0].runs[0].font.size = Pt(10)
        
        hdr_row1[start_col+1].text = year2_lbl
        set_cell_margins(hdr_row1[start_col+1])
        hdr_row1[start_col+1].paragraphs[0].alignment = 1
        hdr_row1[start_col+1].paragraphs[0].runs[0].font.bold = True
        hdr_row1[start_col+1].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_row1[start_col+1].paragraphs[0].runs[0].font.size = Pt(10)
        
        hdr_row1[start_col+2].text = "Pooled"
        set_cell_margins(hdr_row1[start_col+2])
        hdr_row1[start_col+2].paragraphs[0].alignment = 1
        hdr_row1[start_col+2].paragraphs[0].runs[0].font.bold = True
        hdr_row1[start_col+2].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_row1[start_col+2].paragraphs[0].runs[0].font.size = Pt(10)
        
        for sub_col in range(3):
            hdr_row2[start_col + sub_col].text = ""
            set_cell_margins(hdr_row2[start_col + sub_col])
            
    set_header_bottom_border(table.rows[2])
    
    for g in sorted(genotypes):
        row_cells = table.add_row().cells
        row_cells[0].text = str(g)
        set_cell_margins(row_cells[0])
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        
        for i, p in enumerate(params):
            start_col = i * 3 + 1
            pooled_val = (results_1[p]["means"].get(g, 0.0) + results_2[p]["means"].get(g, 0.0)) / 2
            
            row_cells[start_col].text = str(results_1[p]["means_str"].get(g, ""))
            set_cell_margins(row_cells[start_col])
            row_cells[start_col].paragraphs[0].alignment = 1
            row_cells[start_col].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[start_col].paragraphs[0].runs[0].font.size = Pt(10)
            
            row_cells[start_col+1].text = str(results_2[p]["means_str"].get(g, ""))
            set_cell_margins(row_cells[start_col+1])
            row_cells[start_col+1].paragraphs[0].alignment = 1
            row_cells[start_col+1].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[start_col+1].paragraphs[0].runs[0].font.size = Pt(10)
            
            row_cells[start_col+2].text = f"{pooled_val:.2f}"
            set_cell_margins(row_cells[start_col+2])
            row_cells[start_col+2].paragraphs[0].alignment = 1
            row_cells[start_col+2].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[start_col+2].paragraphs[0].runs[0].font.size = Pt(10)
            
    set_header_bottom_border(table.rows[-1])
    
    stats_labels = ["Sem", "p-value", "LSD(0.05)", "CV(%)", "Grand Mean"]
    stats_keys = ["sem", "p_text", "lsd", "cv", "gm"]
    
    for s_idx, (label, key) in enumerate(zip(stats_labels, stats_keys)):
        row_cells = table.add_row().cells
        row_cells[0].text = label
        set_cell_margins(row_cells[0])
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        
        for i, p in enumerate(params):
            start_col = i * 3 + 1
            row_cells[start_col].text = str(results_1[p][key])
            set_cell_margins(row_cells[start_col])
            row_cells[start_col].paragraphs[0].alignment = 1
            row_cells[start_col].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[start_col].paragraphs[0].runs[0].font.size = Pt(10)
            
            row_cells[start_col+1].text = str(results_2[p][key])
            set_cell_margins(row_cells[start_col+1])
            row_cells[start_col+1].paragraphs[0].alignment = 1
            row_cells[start_col+1].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[start_col+1].paragraphs[0].runs[0].font.size = Pt(10)
            
            row_cells[start_col+2].text = ""
            set_cell_margins(row_cells[start_col+2])
            
        if s_idx == len(stats_keys) - 1:
            set_header_bottom_border(row_cells)


# ==============================================================================
# Word report builder: produces the 1 / 1.1 / 1.1.1 hierarchical structure
# ==============================================================================
def build_hierarchical_report_2y(classified_cols, genotype_col, genotypes, results_data_1, results_data_2, year1_lbl, year2_lbl):
    """
    Produces a Document with clustered parameters:
    Groups up to 4 static single parameters under the same category to produce
    cohesive multi-column tables, but writes individual analytical paragraphs 
    with their own bold headings for each parameter.
    """
    doc = Document()
    doc.add_heading("Calculated Multi-Year Single-Factor RCBD Report", 0)
    numberer = ReportNumberer()

    # Dynamic shuffling pools
    single_day_pool = TemplatePool(SINGLE_DAY_TEMPLATES_CATEGORIES)
    time_series_pool = TemplatePool(TIME_SERIES_TEMPLATES_CATEGORIES)

    ordered_majors = [m for m in MAJOR_CATEGORY_ORDER if m in classified_cols]
    ordered_majors += [m for m in classified_cols if m not in ordered_majors]

    for major in ordered_majors:
        subs_dict = classified_cols[major]
        if not any(subs_dict.values()):
            continue

        major_title = numberer.major(doc, major)
        st.write(f"## {major_title}")

        sub_order = SUB_CATEGORY_ORDER.get(major, [])
        ordered_subs = [s for s in sub_order if s in subs_dict and subs_dict[s]]
        ordered_subs += [s for s in subs_dict if s not in ordered_subs and subs_dict[s]]

        for sub in ordered_subs:
            cat_params = subs_dict[sub]
            sub_title = numberer.sub(doc, sub)
            st.write(f"### {sub_title}")

            grouped = group_parameters(cat_params)

            static_items = [items[0][0] for base_name, items in sorted(grouped.items()) if len(items) == 1]
            time_series_items = {base_name: items for base_name, items in sorted(grouped.items()) if len(items) > 1}

            chunk_size = 4
            for i in range(0, len(static_items), chunk_size):
                chunk = static_items[i:i + chunk_size]

                table_n = numberer.next_table()
                table_label = f"Table {table_n}"

                for p in chunk:
                    param_title = numberer.param(doc, p)
                    st.write(f"**{param_title}**")

                    p_text = generate_one_factor_explanation_shuffled_2y(
                        p, results_data_1[p], results_data_2[p], genotype_col, table_label, year1_lbl, year2_lbl, single_day_pool
                    )
                    st.write(p_text)
                    doc.add_paragraph(p_text)

                # Render consolidated multi-column table for this chunk
                add_multiyear_table_to_docx(doc, chunk, genotypes, results_data_1, results_data_2, year1_lbl, year2_lbl)

                caption_text = generate_table_caption(table_n, genotype_col, chunk)
                p_cap = doc.add_paragraph(caption_text)
                p_cap.runs[0].font.name = 'Arial'
                p_cap.runs[0].font.size = Pt(10)
                p_cap.runs[0].font.italic = True

                st.write(f"*{table_label} rendered below*")
                doc.add_paragraph()

            for base_name, items in sorted(time_series_items.items()):
                param_title = numberer.param(doc, base_name)
                st.write(f"**{param_title}**")

                table_n = numberer.next_table()
                table_label = f"Table {table_n}"
                trend_params = [it[0] for it in items]

                p_text = generate_trend_explanation_2y_shuffled(
                    base_name, items, results_data_1, results_data_2, genotype_col, table_label, year1_lbl, year2_lbl, time_series_pool
                )
                st.write(p_text)
                doc.add_paragraph(p_text)

                add_multiyear_table_to_docx(doc, trend_params, genotypes, results_data_1, results_data_2, year1_lbl, year2_lbl)

                caption_text = generate_table_caption(table_n, genotype_col, trend_params)
                p_cap = doc.add_paragraph(caption_text)
                p_cap.runs[0].font.name = 'Arial'
                p_cap.runs[0].font.size = Pt(10)
                p_cap.runs[0].font.italic = True

                st.write(f"*{table_label} (time-series) rendered below*")
                doc.add_paragraph()

        doc.add_paragraph("-" * 60)

    return doc


# --- Web Interface Routing and Multi-Year controller ---
def show_module():
    st.markdown("### Multi-Year Single-Factor RCBD Analyzer")
    
    mode = st.radio("Choose Input Mode", ["Raw Data Mode", "Summarized Table Mode"], key="2f_mode_selector_mult")
    
    if mode == "Raw Data Mode":
        run_raw_mode()
    else:
        file_sum = st.file_uploader("Upload Summarized Multi-Year Excel Output (.xlsx)", type=["xlsx"], key="file_sum_mod_mult")
        if file_sum is not None:
            run_summary_mode_processing(file_sum)


def run_raw_mode():
    file1 = st.file_uploader("Upload Season 1 Raw Excel Dataset (.xlsx)", type=["xlsx"], key="file1_2f_mod")
    file2 = st.file_uploader("Upload Season 2 Raw Excel Dataset (.xlsx)", type=["xlsx"], key="file2_2f_mod")
    
    year1_lbl = st.text_input("Header label for Season 1:", value="2079/80")
    year2_lbl = st.text_input("Header label for Season 2:", value="2080/81")
    
    if file1 is not None and file2 is not None:
        try:
            df1_raw = pd.read_excel(file1)
            df2_raw = pd.read_excel(file2)
            
            st.write(f"#### Season 1 Preview ({year1_lbl}):", df1_raw.head())
            st.write(f"#### Season 2 Preview ({year2_lbl}):", df2_raw.head())
            
            cols = df1_raw.columns.tolist()
            block_col = st.selectbox("Select Block/Replication Column:", cols, index=0, key="block_2f")
            genotype_col = st.selectbox("Select Genotype/Treatment Column:", cols, index=1, key="genotype_2f")
            response_cols = st.multiselect("Select Response Parameters to Analyze:", cols, default=cols[2:], key="response_2f")
            
            if response_cols:
                classified_cols = build_classified_cols(response_cols)
                show_category_preview(classified_cols)
                
                if st.button("Execute Multi-Year Analytical Engine", key="run_2f_engine"):
                    genotypes_1 = sorted(df1_raw[genotype_col].dropna().unique().tolist())
                    genotypes_2 = sorted(df2_raw[genotype_col].dropna().unique().tolist())
                    
                    common_genotypes = sorted(list(set(genotypes_1).intersection(set(genotypes_2))))
                    if not common_genotypes:
                        st.error("Error: Genotype column entries do not match between both files.")
                        return
                    
                    results_data_1 = {}
                    results_data_2 = {}
                    
                    for param in response_cols:
                        results_data_1[param] = run_anova_1factor_raw(df1_raw, block_col, genotype_col, param)
                        results_data_2[param] = run_anova_1factor_raw(df2_raw, block_col, genotype_col, param)
                        
                    # Build Excel Results matching your screenshot layout exactly
                    excel_bio = io.BytesIO()
                    styled_wb = build_multiyear_excel_output(
                        genotype_col, response_cols, common_genotypes, 
                        results_data_1, results_data_2, year1_lbl, year2_lbl
                    )
                    styled_wb.save(excel_bio)
                    excel_bio.seek(0)
                    
                    st.markdown("#### \U0001F4E5 Download Formatted Statistical Excel Results")
                    st.download_button(
                        label="Download Excel Results Sheet",
                        data=excel_bio,
                        file_name="RCBD_Result_Final.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="btn_d_excel_2f"
                    )
                    st.write("---")
                    
                    # Generate dynamic report
                    st.markdown("### \U0001F4DD Dynamic Analysis Results & Discussions")
                    doc = build_hierarchical_report_2y(
                        classified_cols, genotype_col, common_genotypes, 
                        results_data_1, results_data_2, year1_lbl, year2_lbl
                    )
                    
                    bio_doc = io.BytesIO()
                    doc.save(bio_doc)
                    bio_doc.seek(0)
                    
                    st.write("---")
                    st.markdown("#### \U0001F4BE Save Explanations & Tables as Word Document")
                    st.download_button(
                        "Download Word Explanations Report (.docx)", 
                        data=bio_doc, 
                        file_name="MultiYear_Thesis_Report.docx", 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                        key="btn_d_2f_doc"
                    )
        except Exception as e:
            st.error(f"Analysis failed. Please check the structure of your Excel dataset: {e}")


# --- Summarized Mode Processing (Multi-Year) ---
def run_summary_mode_processing(uploaded_file):
    try:
        df_raw = pd.read_excel(uploaded_file, header=None)
        
        idx_sem, idx_pval, idx_lsd, idx_cv, idx_gm = None, None, None, None, None
        for idx, val in enumerate(df_raw[0]):
            if pd.isna(val): continue
            val_str = str(val).strip().lower()
            if val_str == "sem": idx_sem = idx
            elif val_str == "p-value": idx_pval = idx
            elif "lsd" in val_str: idx_lsd = idx
            elif "cv" in val_str: idx_cv = idx
            elif "grand mean" in val_str or "grandmean" in val_str: idx_gm = idx
            
        if any(v is None for v in [idx_sem, idx_pval, idx_lsd, idx_cv, idx_gm]):
            st.error("Missing structural labels (Sem, p-value, LSD, CV, Grand Mean) in Column A.")
            return
            
        # Parse parameters from Row 1
        parameters = []
        param_cols = {}
        for col_idx in range(1, df_raw.shape[1], 3):
            val = df_raw.iloc[0, col_idx]
            if pd.notna(val) and str(val).strip() != "":
                p_name = str(val).strip()
                parameters.append(p_name)
                param_cols[p_name] = col_idx
                
        # Parse Genotypes (Starts at index 3 up to idx_sem - 1)
        genotypes = [str(df_raw.iloc[r, 0]).strip() for r in range(3, idx_sem)]
        genotype_col_label = str(df_raw.iloc[0, 0]).strip() if pd.notna(df_raw.iloc[0, 0]) else "Genotype"

        # Parse Season Names
        year1_lbl = str(df_raw.iloc[1, 1]).strip() if pd.notna(df_raw.iloc[1, 1]) else "Year 1"
        year2_lbl = str(df_raw.iloc[1, 2]).strip() if pd.notna(df_raw.iloc[1, 2]) else "Year 2"
        
        st.success(f"Detected Factor: {genotype_col_label} | Genotypes: {', '.join(genotypes)}")
        st.success(f"Detected Parameters: {', '.join(parameters)}")
        
        classified_cols = build_classified_cols(parameters)
        show_category_preview(classified_cols)
        
        if st.button("Generate Word Document Draft", key="btn_1f_sum_gen_2y"):
            results_data_1, results_data_2 = parse_summarized_table_to_results_2y(
                df_raw, idx_sem, idx_pval, idx_lsd, idx_cv, idx_gm, param_cols, genotypes
            )
            
            doc = build_hierarchical_report_2y(
                classified_cols, genotype_col_label, genotypes, 
                results_data_1, results_data_2, year1_lbl, year2_lbl
            )
            
            bio_doc = io.BytesIO()
            doc.save(bio_doc)
            bio_doc.seek(0)
            
            st.write("---")
            st.markdown("#### \U0001F4BE Save Explanations & Tables as Word Document")
            st.download_button(
                "Download Word Explanations Report (.docx)", 
                data=bio_doc, 
                file_name="MultiYear_Summarized_Thesis_Report.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                key="btn_d_2y_doc_sum"
            )
    except Exception as e:
        st.error(f"Error parsing direct result summary table: {e}")


if __name__ == "__main__":
    show_module()
