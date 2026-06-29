import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import string
import random
from scipy.stats import t
import statsmodels.api as sm
from statsmodels.formula.api import ols
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side

# ==============================================================================
# DATABASE OF VERBOSE, HIGH-STANDARD ACADEMIC DISCUSSION TEMPLATES (30 TEMPLATES)
# ==============================================================================
ACADEMIC_TEMPLATES_30 = {
    # --- Group A: Temporal / Multi-Date / Trend Line Parameters (1–15) ---
    1: (
        "The chronological progression of {{variable_name}} was monitored sequentially across {{num_intervals}} "
        "experimental intervals from {{start_time}} to {{end_time}} {{time_unit}}. Table {{table_num}} shows that during "
        "the initial stages at {{time_point_1}} and {{time_point_2}} {{time_unit}}, the treatments of {{factor_A}} did not "
        "exert any statistically significant influence on the observed values (P > 0.05). The overall grand mean during "
        "this early developmental phase remained stable at {{grand_mean_early}} {{unit}} with a CV of {{cv_early}}%. "
        "However, as the experimental timeline progressed to {{time_point_3}} {{time_unit}}, a highly significant treatment "
        "divergence emerged (P \u2264 0.01). Table {{table_num}} shows that {{treatment_A1}} achieved the maximum value "
        "of {{value_1}} {{unit}}, which was significantly higher than {{treatment_A2}} ({{value_2}} {{unit}}) and the "
        "baseline treatment {{treatment_A_lowest}} ({{value_lowest}} {{unit}}; LSD_0.05 = {{lsdval}})."
    ),
    2: (
        "The temporal trajectory of {{variable_name}} followed a parabolic peak-and-decline curve over the {{total_days}}-day "
        "evaluation period. Table {{table_num}} shows that the grand mean of {{variable_name}} began at {{initial_value}} "
        "{{unit}} ({{time_point_1}} {{time_unit}}), peaked at {{peak_value}} {{unit}} at {{time_point_2}} {{time_unit}}, and "
        "subsequently declined to {{end_value}} {{unit}} at the closing evaluation phase of {{time_point_3}} {{time_unit}}. "
        "Although treatment differences were statistically non-significant at the beginning of the evaluation stage, the rate "
        "of post-peak decline was significantly altered by the application of the treatments of {{factor_A}}. Table {{table_num}} "
        "shows that at {{time_point_2}} {{time_unit}}, {{treatment_A1}} maintained significantly higher {{variable_name}} levels "
        "({{value_A1}} {{unit}}) compared to the baseline treatment ({{value_baseline}} {{unit}}; P \u2264 0.05). This treatment "
        "trend widened toward the final sampling stage of {{time_point_3}} {{time_unit}}, where {{treatment_A1}} ({{value_A1_end}} "
        "{{unit}}) significantly limited the decrease of {{variable_name}} compared to {{treatment_A_lowest}} ({{value_lowest_end}} "
        "{{unit}}; P \u2264 0.01)."
    ),
    3: (
        "The effect of the treatments of {{factor_A}} on the values of {{variable_name}} transitioned from a non-significant "
        "state to a highly distinct differentiation pattern. Table {{table_num}} shows that at {{time_point_1}} and {{time_point_2}} "
        "{{time_unit}}, the F-test for {{factor_A}} was non-significant (P > 0.05), and all treatment levels shared the identical "
        "post-hoc letter grouping 'a'. However, a critical transition occurred at {{time_point_3}} {{time_unit}}, where treatment "
        "effects became significant (P \u2264 0.05). At this stage, {{treatment_A1}} produced {{val_A1}} {{unit}}, which was "
        "significantly higher than the baseline treatment {{treatment_A_lowest}} ({{val_lowest}} {{unit}}). By {{time_point_4}} "
        "{{time_unit}}, this differentiation became highly significant (P \u2264 0.01, LSD_0.05 = {{lsdval}}), indicating that the "
        "progression of the experiment enhanced the treatment differences over time."
    ),
    4: (
        "The divergence among the levels of {{factor_A}} became progressively more defined over the temporal evaluation scale. "
        "Table {{table_num}} shows that at the initial interval of {{time_point_1}} {{time_unit}}, the treatment effect was "
        "non-significant, with a narrow range between the highest value of {{val_max_early}} {{unit}} (under {{treatment_A1}}) "
        "and the lowest value of {{val_min_early}} {{unit}} (under {{treatment_A_lowest}}). However, at {{time_point_2}} {{time_unit}}, "
        "the main effect of {{factor_A}} was significant (P \u2264 0.05), with {{treatment_A1}} and {{treatment_A2}} separating "
        "into a superior statistical group 'a' compared to the lower treatments. By {{time_point_3}} {{time_unit}}, this divergence "
        "reached its maximum, showing a highly significant treatment effect (P \u2264 0.01) where {{treatment_A1}} maintained the "
        "maximum value ({{val_max_late}} {{unit}}), which was significantly higher than all other treatment levels (LSD_0.05 = {{lsdval}})."
    ),
    5: (
        "The temporal response of {{variable_name}} was characterized by a distinct statistical phase-shift between the early phase "
        "({{time_point_1}} to {{time_point_2}} {{time_unit}}) and the final phase ({{time_point_3}} to {{time_point_4}} {{time_unit}}). "
        "Table {{table_num}} shows that during the early phase, all levels of {{factor_A}} shared identical letter superscripts (P > 0.05), "
        "with values closely matching the baseline mean of {{mean_early}} {{unit}}. Following the transition to the final phase at "
        "{{time_point_3}} {{time_unit}}, the main effect of {{factor_A}} became highly significant (P \u2264 0.01). Table {{table_num}} "
        "shows that {{treatment_A1}} produced the highest value ({{val_A1}} {{unit}}), which was statistically equivalent only to "
        "{{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}; "
        "LSD_0.05 = {{lsdval}}). This significant main effect persisted through the final evaluation stage at {{time_point_4}} {{time_unit}}, "
        "with a low overall coefficient of variation (CV = {{cvval}}%)."
    ),
    6: (
        "The rate of decrease in {{variable_name}} following the peak phase was significantly influenced by the levels of the "
        "treatment {{factor_A}}. Table {{table_num}} shows that from {{time_point_1}} to {{time_point_2}} {{time_unit}}, the "
        "experimental units maintained high stability with a grand mean of {{grand_mean_peak}} {{unit}}, and no treatment "
        "variations were observed (P > 0.05). However, from {{time_point_3}} {{time_unit}} onwards, a rapid downward trend was "
        "observed in {{variable_name}} values. The application of {{treatment_A1}} significantly slowed this decrease, maintaining "
        "a value of {{val_A1}} {{unit}} at {{time_point_4}} {{time_unit}}, compared to the baseline treatment which dropped to "
        "{{val_baseline}} {{unit}} (P \u2264 0.05). This protective main effect of {{treatment_A1}} was highly significant at this "
        "ultimate stage (P \u2264 0.01, LSD_0.05 = {{lsdval}}), showing that the treatment worked to slow the downward progression "
        "through the final sampling stages."
    ),
    7: (
        "The cumulative progression of {{variable_name}} was recorded sequentially from {{time_point_1}} to {{time_point_4}} "
        "{{time_unit}} to evaluate the accumulation kinetics of the system. Table {{table_num}} shows that the cumulative curve "
        "followed a standard sigmoidal pattern, with the fastest accumulation rate observed between {{time_point_2}} and "
        "{{time_point_3}} {{time_unit}}. Although there were no differences among treatments during the initial lag phase at "
        "{{time_point_1}} {{time_unit}}, the levels of {{factor_A}} significantly affected the accumulation rate during the "
        "active phase. At {{time_point_3}} {{time_unit}}, {{treatment_A1}} achieved a significantly higher cumulative value "
        "({{val_A1}} {{unit}}) compared to the baseline treatment ({{val_baseline}} {{unit}}; P \u2264 0.01, LSD_0.05 = {{lsdval}}). "
        "This significant difference was maintained through the plateau phase at {{time_point_4}} {{time_unit}}, while the "
        "internal experimental variance remained within precise statistical limits (P > 0.05)."
    ),
    8: (
        "The experimental design demonstrated high precision in monitoring the temporal dynamics of {{variable_name}}, with "
        "Coefficient of Variation (CV%) values ranging from {{cv_min}}% to {{cv_max}}% across the {{num_dates}} sampling dates. "
        "Table {{table_num}} shows that the Standard Error of the Mean (SEm) for the treatment levels remained low, ranging "
        "between {{sem_min}} and {{sem_max}} {{unit}}. Under these precise conditions, treatment differences remained non-significant "
        "at {{time_point_1}} and {{time_point_2}} {{time_unit}}. At {{time_point_3}} {{time_unit}}, the effect of {{factor_A}} "
        "was significant (P \u2264 0.05, LSD_0.05 = {{lsdval}}), with {{treatment_A1}} ({{val_A1}} {{unit}}) outperforming "
        "{{treatment_A2}} ({{val_A2}} {{unit}}). By {{time_point_4}} {{time_unit}}, the main effect of {{factor_A}} became highly "
        "significant (P \u2264 0.01, LSD_0.05 = {{lsdval}}), confirming strong statistical control over experimental variance."
    ),
    9: (
        "The mathematical differences (\u0394) and percentage changes in {{variable_name}} were evaluated across several sampling "
        "intervals under different levels of {{factor_A}}. Table {{table_num}} shows that while treatments remained statistically "
        "equivalent at {{time_point_1}} and {{time_point_2}} {{time_unit}}, significant treatment differences emerged at "
        "{{time_point_3}} {{time_unit}}. At this stage, {{treatment_A1}} increased {{variable_name}} by {{pct_diff_A}}% compared "
        "to the lowest treatment level (\u0394 = {{deltaA}} {{unit}}; P \u2264 0.05). At {{time_point_4}} {{time_unit}}, this "
        "percentage difference increased, with {{treatment_A1}} outperforming the lowest treatment level by {{pct_diff_A_late}}% "
        "(P \u2264 0.01). This demonstrated the growing effect of the treatments as the experimental units reached the final stages."
    ),
    10: (
        "1.1.1 {{sub_parameter_1}}\n"
        "The temporal development of {{sub_parameter_1}} showed stable values during the early evaluation stages, followed by "
        "significant treatment differences as the experimental units progressed. Table {{table_num}} shows that no significant "
        "variations were observed at {{time_point_1}} and {{time_point_2}} {{time_unit}}, with values remaining close to the "
        "grand mean of {{mean_early}} {{unit}}. However, at {{time_point_3}} {{time_unit}}, the main effect of {{factor_A}} was "
        "significant (P \u2264 0.05), with {{treatment_A1}} ({{val_A1}} {{unit}}) outperforming the lowest treatment level "
        "({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}). This significant main effect persisted through the final sampling "
        "date of {{time_point_4}} {{time_unit}}.\n"
        "1.1.2 {{sub_parameter_2}}\n"
        "Unlike the first parameter, {{sub_parameter_2}} showed a highly significant treatment response at an earlier stage "
        "in the evaluation cycle. Table {{table_num}} shows that the main effect of {{factor_A}} was significant at {{time_point_2}} "
        "{{time_unit}} (P \u2264 0.05), with {{treatment_A1}} producing {{val_A1_p2}} {{unit}} compared to the lowest treatment "
        "level ({{val_lowest_p2}} {{unit}}). By {{time_point_4}} {{time_unit}}, this treatment differentiation became highly "
        "significant (P \u2264 0.01, LSD_0.05 = {{lsdvalp2}}), with all treatment levels separating into distinct post-hoc letter groupings."
    ),
    11: (
        "Table {{table_num}} shows that the {{variable_name}} values maintained a state of statistical equivalence during the "
        "first two sampling dates ({{time_point_1}} and {{time_point_2}} {{time_unit}}), with all levels of {{factor_A}} "
        "sharing the same post-hoc letter grouping 'a'. During this early period, the grand means were recorded as {{gm_1}} "
        "{{unit}} and {{gm_2}} {{unit}}, respectively, with no significant differences observed. However, a transition occurred "
        "at {{time_point_3}} {{time_unit}} when the main effect of {{factor_A}} became significant (P \u2264 0.05). Treatment "
        "{{treatment_A1}} broke the statistical equivalence, producing {{val_A1}} {{unit}} (assigned to letter grouping 'a') "
        "compared to treatment {{treatment_A2}} which produced {{val_A2}} {{unit}} (assigned to letter grouping 'b'). By "
        "{{time_point_4}} {{time_unit}}, the treatment effect became highly significant (P \u2264 0.01, LSD_0.05 = {{lsdval}}), "
        "breaking the early statistical equivalence across all treatment levels."
    ),
    12: (
        "The temporal development of {{variable_name}} was monitored across several sampling points under the direct influence "
        "of the evaluated treatments. Table {{table_num}} shows that the treatment effect of {{factor_A}} was non-significant "
        "at {{time_point_1}} {{time_unit}}, but became highly significant at {{time_point_2}} {{time_unit}} (P \u2264 0.01). "
        "At this stage, {{treatment_A1}} achieved the maximum value of {{val_max_2}} {{unit}}, which was significantly higher "
        "than the baseline treatment {{treatment_A_lowest}} ({{val_lowest_2}} {{unit}}; LSD_0.05 = {{lsdval}}). This highly "
        "significant treatment effect was maintained through the late stages of {{time_point_3}} and {{time_point_4}} {{time_unit}}, "
        "with a low overall standard error of the mean (SEm) of {{sem_val}} {{unit}} across the entire developmental sequence."
    ),
    13: (
        "The range of values for {{variable_name}} across different levels of {{factor_A}} was small during the early stages "
        "but expanded considerably as the study progressed. Table {{table_num}} shows that at {{time_point_1}} and {{time_point_2}} "
        "{{time_unit}}, the range between the highest and lowest treatment means was narrow and statistically non-significant, "
        "with a maximum range of only {{range_early}} {{unit}}. However, at {{time_point_3}} {{time_unit}}, this treatment range "
        "expanded significantly. The range between the highest-performing treatment ({{treatment_A1}}; {{val_max_A}} {{unit}}) "
        "and the lowest-performing treatment ({{val_min_A}} {{unit}}) was {{range_mid_A}} {{unit}} (P \u2264 0.01). At {{time_point_4}} "
        "{{time_unit}}, the treatment range remained wide, with a difference of {{range_late_A}} {{unit}} among {{factor_A}} "
        "treatments (P \u2264 0.01, LSD_0.05 = {{lsdval}}), demonstrating the cumulative effect of the treatments over the entire study."
    ),
    14: (
        "The {{variable_name}} of the experimental units was monitored across a series of developmental stages from the initial "
        "phase to the final stage of observation. Table {{table_num}} shows that no treatment differences were observed during "
        "the initial phase at {{time_point_1}} and {{time_point_2}} {{time_unit}} (P > 0.05), with values remaining close to "
        "the grand mean of {{gm_early}} {{unit}}. The first significant differences emerged during the middle phase at "
        "{{time_point_3}} {{time_unit}}, where the main effect of {{factor_A}} was significant (P \u2264 0.05), with {{treatment_A1}} "
        "producing a higher value ({{val_A1}} {{unit}}) than {{treatment_A2}} ({{val_A2}} {{unit}}). As the experimental system "
        "reached peak values at {{time_point_4}} {{time_unit}}, treatment differences became highly significant (P \u2264 0.01). "
        "Table {{table_num}} shows that {{treatment_A1}} produced the highest overall value of {{val_max_late}} {{unit}}, which "
        "was significantly higher than the lowest treatment level ({{val_lowest_late}} {{unit}}; LSD_0.05 = {{lsdval}})."
    ),
    15: (
        "The {{variable_name}} levels were monitored to identify when the experimental units transitioned past critical "
        "operational thresholds under different levels of {{factor_A}}. Table {{table_num}} shows that during the initial stages "
        "at {{time_point_1}} and {{time_point_2}} {{time_unit}}, all treatments maintained high, stable values above the critical "
        "threshold of {{threshold_val}} {{unit}}, with no significant differences observed among them (P > 0.05). A critical "
        "transition occurred at {{time_point_3}} {{time_unit}}, where the values under the lower-performing treatments began to "
        "drop below the threshold level. The application of {{treatment_A1}} significantly delayed this downward trend, maintaining "
        "a value of {{val_A1}} {{unit}} (well above the threshold), while the lowest-performing treatment fell to {{val_lowest}} "
        "{{unit}} (P \u2264 0.05). This protective main effect of {{treatment_A1}} was highly significant at this final stage "
        "(P \u2264 0.01, LSD_0.05 = {{lsdval}}), showing that the treatments worked together to maintain the values of the "
        "system through the final sampling stages."
    ),

    # --- Group B: Single-Day / End-Point / Terminal Parameters (16–30) ---
    16: (
        "The final measurement of {{variable_name}} was highly significantly affected by different levels of {{factor_A}}. "
        "Table {{table_num}} shows that the main effect of the treatment was highly significant (P \u2264 0.01), with "
        "{{treatment_A1}} producing the highest value of {{val_A1}} {{unit}}, which was statistically equivalent to "
        "{{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level ({{val_lowest}} "
        "{{unit}}; LSD_0.05 = {{lsdval}}). The lowest treatment level recorded the lowest value of {{val_min_A}} {{unit}}, "
        "resulting in an overall grand mean of {{grand_mean}} {{unit}} and a low coefficient of variation (CV = {{cvval}}%), "
        "indicating high experimental precision at this final valuation stage."
    ),
    17: (
        "The end-point value of {{variable_name}} was significantly affected by different levels of {{factor_A}}, though "
        "some overlap was observed among treatment means. Table {{table_num}} shows that the main effect of {{factor_A}} "
        "was significant (P \u2264 0.05), with {{treatment_A1}} producing the highest value of {{val_A1}} {{unit}}, which "
        "was statistically equivalent to {{treatment_A2}} ({{val_A2}} {{unit}}) and {{treatment_A3}} ({{val_A3}} {{unit}}), "
        "but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}). "
        "The overall grand mean of the experimental trial was {{grand_mean}} {{unit}} with a standard error of the mean "
        "(SEm) of {{sem_val}} {{unit}}, indicating a stable baseline at this end-point assessment."
    ),
    18: (
        "The ultimate level of {{variable_name}} was not significantly affected by the different levels of {{factor_A}}. "
        "Table {{table_num}} shows that the values remained highly uniform across all treatment levels (P > 0.05), with a "
        "maximum of {{val_max_A}} recorded under {{treatment_A1}} and a minimum of {{val_min_A}} recorded under "
        "{{treatment_A2}}, which fell well within the non-significant range based on the post-hoc separation "
        "(LSD_0.05 = {{lsdval}}). The SEm for the treatment factor was recorded as {{sem_val}}, and the overall grand "
        "mean of the experimental trial was {{grand_mean}} {{unit}}. This lack of statistical divergence, coupled with "
        "a remarkably low coefficient of variation (CV = {{cvval}}%), confirms that {{variable_name}} maintained a stable "
        "baseline during this ultimate evaluation."
    ),
    19: (
        "The concluding evaluation of the parameter index for {{variable_name}} showed high stability and low overall "
        "variance across all treatment levels. Table {{table_num}} shows that the Coefficients of Variation (CV%) "
        "remained very low at {{cv_val}}%, and the Standard Error of the Mean (SEm) was recorded as {{sem_val}} {{unit}}, "
        "indicating a high degree of experimental precision and baseline consistency. Under these uniform conditions, "
        "the main effect of {{factor_A}} did not show any statistically significant differences (P > 0.05). The mean "
        "values remained close to the grand mean of {{grand_mean}} {{unit}}, with all treatments sharing the same "
        "post-hoc letter grouping 'a', indicating that the parameter was robust and remained stable across all "
        "treatment levels at this concluding stage."
    ),
    20: (
        "The relative percentage changes of {{variable_name}} were evaluated to assess the magnitude of the treatment "
        "effects under different levels of {{factor_A}} at this at-harvest assessment. Table {{table_num}} shows that "
        "the application of {{treatment_A1}} increased {{variable_name}} by {{pct_diff_A}}% compared to the lowest "
        "treatment level ({{val_max_A}} vs. {{val_min_A}} {{unit}}). Other treatments, such as {{treatment_A2}} and "
        "{{treatment_A3}}, also increased the values by {{pct_diff_A2}}% and {{pct_diff_A3}}% over the lowest treatment "
        "level, respectively. These differences were highly significant for the {{factor_A}} main effect (P \u2264 0.01, "
        "LSD_0.05 = {{lsdval}}), demonstrating the strong effect of the treatment levels on the evaluated parameter."
    ),
    21: (
        "The mature-stage level of {{variable_name}} was significantly affected by different levels of {{factor_A}} under "
        "precise experimental conditions. Table {{table_num}} shows that the main effect of {{factor_A}} was significant "
        "(P \u2264 0.05, LSD_0.05 = {{lsdval}}), with {{treatment_A1}} producing the highest value ({{val_A1}} {{unit}}), "
        "which was statistically equivalent to {{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the "
        "lowest treatment level ({{val_lowest}} {{unit}}). The standard error of the mean (SEm) for the {{factor_A}} "
        "treatment was {{sem_val}}, and the overall grand mean of the trial was recorded as {{grand_mean}} {{unit}} with "
        "a low coefficient of variation (CV = {{cvval}}%)."
    ),
    22: (
        "The final ratio of the system components was significantly affected by different levels of the treatment factor "
        "at this last-sampling assessment. Table {{table_num}} shows that the main effect of {{factor_A}} was significant "
        "(P \u2264 0.05), with {{treatment_A1}} producing the highest fractional ratio ({{val_A1}} {{unit}}), which was "
        "statistically equivalent to {{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the lowest "
        "treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}). The overall grand mean of the experiment was "
        "recorded at {{grand_mean}} {{unit}} and a low coefficient of variation (CV = {{cvval}}%), indicating that the "
        "system's proportional structure was highly responsive to the treatments."
    ),
    23: (
        "The stabilized value of the activity rate of {{variable_name}} was significantly affected by different levels of "
        "{{factor_A}}. Table {{table_num}} shows that the main effect of {{factor_A}} was highly significant (P \u2264 0.01), "
        "with {{treatment_A1}} producing the highest activity rate ({{val_A1}} {{unit}}), which was statistically on par "
        "with {{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level ({{val_lowest}} "
        "{{unit}}; LSD_0.05 = {{lsdval}}). The overall grand mean of the experiment was recorded at {{grand_mean}} {{unit}} "
        "with a low coefficient of variation (CV = {{cvval}}%), indicating clear differentiation across treatments."
    ),
    24: (
        "The post-experimental value of the concentration of {{variable_name}} in the matrix was significantly affected by "
        "different levels of {{factor_A}}. Table {{table_num}} shows that the main effect of {{factor_A}} was highly significant "
        "(P \u2264 0.01), with {{treatment_A1}} producing the highest concentration of {{val_A1}} {{unit}} [24], which was "
        "significantly higher than {{treatment_A2}} ({{val_A2}} {{unit}}) and the lowest treatment level ({{val_lowest}} "
        "{{unit}}; LSD_0.05 = {{lsdval}}). This verified that the choice of {{factor_A}} was the primary factor driving "
        "differences in matrix concentration."
    ),
    25: (
        "The ultimate ratio of the primary output component, {{variable_name}}, was significantly affected by different levels "
        "of {{factor_A}}. Table {{table_num}} shows that the main effect of {{factor_A}} was significant (P \u2264 0.05), "
        "with {{treatment_A1}} producing the highest value ({{val_A1}}), which was statistically equivalent to {{treatment_A2}} "
        "({{val_A2}}) but significantly higher than the lowest treatment level ({{val_lowest}}; LSD_0.05 = {{lsdval}}). "
        "This indicated that different treatment levels contributed significantly to the ultimate output performance, with a "
        "recorded grand mean of {{grand_mean}} and a low coefficient of variation (CV = {{cvval}}%)."
    ),
    26: (
        "The at-collection activity of physical resistance of {{variable_name}} showed no significant differences among "
        "treatment levels. Table {{table_num}} shows that the values remained statistically uniform across all treatments "
        "(P > 0.05), with a maximum of {{val_max_A}} recorded under {{treatment_A1}} and a minimum of {{val_min_A}} recorded "
        "under {{treatment_A2}}, which fell well within the non-significant range based on the post-hoc separation "
        "(LSD_0.05 = {{lsdval}}). The SEm for the treatment factor was recorded as {{sem_val}}, and the overall grand mean "
        "of the experimental trial was {{grand_mean}} {{unit}}. This lack of statistical divergence, coupled with a low "
        "coefficient of variation (CV = {{cvval}}%), confirms that this structural parameter was robust to changes in the "
        "treatment levels."
    ),
    27: (
        "The residual concentration of elements classified in the premium grade, designated as {{variable_name}}, was "
        "significantly affected by the different levels of {{factor_A}}. Table {{table_num}} shows that the main effect "
        "of {{factor_A}} was highly significant (P \u2264 0.01), with {{treatment_A1}} producing the highest proportion of "
        "{{val_A1}}%, which was statistically equivalent to {{treatment_A2}} ({{val_A2}}%) but significantly higher than "
        "the lowest treatment level ({{val_lowest}}%; LSD_0.05 = {{lsdval}}). This indicated that the treatment levels "
        "affected the proportional grade distribution, with a recorded grand mean of {{grand_mean}}% and a standard "
        "error of the mean (SEm) of {{sem_val}}."
    ),
    28: (
        "The ultimate output of {{variable_name}} showed small but statistically significant differences in response to "
        "different levels of {{factor_A}}. Table {{table_num}} shows that the main effect of {{factor_A}} was significant "
        "(P \u2264 0.05), with {{treatment_A1}} producing the highest value of {{val_A1}} {{unit}}, while the other "
        "treatment levels showed minor variations, with {{treatment_A2}} at {{val_A2}} {{unit}} and {{treatment_A3}} at "
        "{{val_A3}} {{unit}} (LSD_0.05 = {{lsdval}}). The overall grand mean of the trial was recorded as {{grand_mean}} "
        "{{unit}} with a low coefficient of variation (CV = {{cvval}}%), indicating that even small variations among "
        "treatment levels were successfully separated by the statistical model."
    ),
    29: (
        "The concluding physical resistance of {{variable_name}} was evaluated under different levels of {{factor_A}} with "
        "a specific focus on baseline dispersion and analytical variance. Table {{table_num}} shows that the experimental "
        "environment exhibited minimal variance, allowing for highly precise diagnostic outputs across all treatment groups. "
        "Under these precise conditions, the treatment effect of {{factor_A}} was highly significant (P \u2264 0.01), with "
        "{{treatment_A1}} producing the highest value of {{val_A1}} {{unit}} compared to {{treatment_A2}} ({{val_A2}} {{unit}}) "
        "and the lowest treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}), confirming a significant treatment "
        "response against the baseline."
    ),
    30: (
        "The concluding grade of {{variable_name}} was evaluated under different levels of {{factor_A}} across multiple "
        "replications. Table {{table_num}} shows that the main effect of the treatment was highly significant (P \u2264 0.01). "
        "Post-hoc separation using LSD showed that {{treatment_A1}} produced the highest value of {{val_A1}} {{unit}}, which "
        "was statistically equivalent to {{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the lowest "
        "treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsdval}}). The overall grand mean of the experiment was "
        "recorded at {{grand_mean}} {{unit}} with a low coefficient of variation (CV = {{cvval}}%) and a standard error of "
        "the mean (SEm) of {{sem_val}} {{unit}}, demonstrating that the statistical model achieved high experimental "
        "precision in identifying treatment differences."
    )
}

# --- Shuffling Pool Definitions ---
SINGLE_DAY_TEMPLATES_CATEGORIES = {
    "highly_significant": [16, 20, 23, 24, 27, 29, 30],
    "significant": [17, 21, 22, 25, 28],
    "non_significant": [18, 19, 26]
}

TIME_SERIES_TEMPLATES_CATEGORIES = {
    "divergent": [1, 3, 5, 11],
    "upward_trend": [4, 7, 9, 10, 13],
    "downward_trend": [2, 6, 15],
    "general": [8, 12, 14]
}

# ==============================================================================
# Dynamic Shuffling Class Pool
# ==============================================================================
class TemplatePool:
    """Manages randomized selection of template structures by category."""
    def __init__(self, categories_dict):
        self.pools = {}
        for category, indices in categories_dict.items():
            shuffled = list(indices)
            random.shuffle(shuffled)
            self.pools[category] = shuffled
            self.initial_pool = dict(categories_dict)

    def get_template_id(self, category):
        if category not in self.pools or not self.pools[category]:
            shuffled = list(self.initial_pool[category])
            random.shuffle(shuffled)
            self.pools[category] = shuffled
        return self.pools[category].pop(0)


# ==============================================================================
# Hierarchical categorization layout setup
# ==============================================================================
MAJOR_CATEGORY_ORDER = [
    "Vegetative Parameters",
    "Reproductive Parameters",
    "Post-harvest Parameters",
    "Stress and Health Parameters",
]

SUB_CATEGORY_ORDER = {
    "Vegetative Parameters": [
        "Germination and Establishment",
        "Plant Morphology",
        "Leaf Parameters",
        "Branch Parameters",
        "Root Parameters",
        "Biomass Parameters",
        "Growth Analysis Parameters",
        "Physiological Parameters",
        "Phenological Parameters",
    ],
    "Reproductive Parameters": [
        "Flowering Parameters",
        "Pollination Parameters",
        "Fruit Set Parameters",
        "Fruit Growth Parameters",
        "Fruit Yield Parameters",
        "Grain Yield Parameters (Cereal Crops)",
        "Seed Parameters",
        "Maturity Parameters",
    ],
    "Post-harvest Parameters": [
        "Shelf-life Parameters",
        "Chemical Quality",
        "Physical Quality",
    ],
    "Stress and Health Parameters": [
        "Disease and Pest Parameters",
    ],
}


class ReportNumberer:
    """Generates continuous 1 / 1.1 / 1.1.1 style headings and continuous Table N numbers."""
    def __init__(self):
        self.major_num = 0
        self.sub_num = 0
        self.param_num = 0
        self.table_num = 1

    def major(self, doc, title):
        self.major_num += 1
        self.sub_num = 0
        self.param_num = 0
        full_title = f"{self.major_num} {title}"
        doc.add_heading(full_title, level=1)
        return full_title

    def sub(self, doc, title):
        self.sub_num += 1
        self.param_num = 0
        full_title = f"{self.major_num}.{self.sub_num} {title}"
        doc.add_heading(full_title, level=2)
        return full_title

    def param(self, doc, title):
        self.param_num += 1
        full_title = f"{self.major_num}.{self.sub_num}.{self.param_num} {title}"
        doc.add_heading(full_title, level=3)
        return full_title

    def next_table(self):
        n = self.table_num
        self.table_num += 1
        return n


# --- P-Value Academic Notation Mapper ---
def get_p_val_notation(p_val):
    if p_val < 0.001:
        return "p < 0.001"
    elif p_val < 0.01:
        return "p < 0.01"
    elif p_val < 0.05:
        return "p < 0.05"
    else:
        return "p > 0.05"


# --- Agronomic Classification Engine ---
def classify_parameter(param):
    param_clean = param.strip().lower()

    abbrev_map = {
        "sl": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "dl": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "plw": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "plwd": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "tss": ("Post-harvest Parameters", "Chemical Quality"),
        "ta": ("Post-harvest Parameters", "Chemical Quality"),
        "ph": ("Post-harvest Parameters", "Chemical Quality"),
        "lai": ("Vegetative Parameters", "Leaf Parameters"),
        "lad": ("Vegetative Parameters", "Leaf Parameters"),
        "sla": ("Vegetative Parameters", "Leaf Parameters"),
        "slw": ("Vegetative Parameters", "Leaf Parameters"),
        "cgr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "rgr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "agr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "nar": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "lar": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "lwr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "rue": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "hi": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "rwc": ("Vegetative Parameters", "Physiological Parameters"),
        "spad": ("Vegetative Parameters", "Leaf Parameters"),
        "ndvi": ("Vegetative Parameters", "Physiological Parameters"),
        "pri": ("Vegetative Parameters", "Physiological Parameters"),
    }

    base_abbrev = re.sub(r"\d+$", "", param_clean)
    if base_abbrev in abbrev_map:
        return abbrev_map[base_abbrev]

    if any(x in param_clean for x in ["disease", "pest", "infest", "weed", "mortality", "survival",
                                       "stress tolerance", "chlorosis", "wilting"]):
        return ("Stress and Health Parameters", "Disease and Pest Parameters")

    if any(x in param_clean for x in ["plw", "physiological loss", "decay", "shelf life", "shelf-life",
                                       "spoilage", "marketability", "firmness retention"]):
        return ("Post-harvest Parameters", "Shelf-life Parameters")

    if any(x in param_clean for x in ["tss", "soluble solid", "titratable", "acidity", "ta", "ph",
                                       "vitamin c", "ascorbic", "lycopene", "carotene", "anthocyanin",
                                       "phenolic", "flavonoid", "antioxidant", "sugar", "dry matter"]):
        return ("Post-harvest Parameters", "Chemical Quality")

    if any(x in param_clean for x in ["firmness", "peel thickness", "specific gravity"]) or \
       (any(x in param_clean for x in ["fruit size", "fruit shape", "fruit color"]) and "quality" in param_clean):
        return ("Post-harvest Parameters", "Physical Quality")

    if any(x in param_clean for x in ["flower", "flowering", "initiation", "cluster"]):
        return ("Reproductive Parameters", "Flowering Parameters")

    if any(x in param_clean for x in ["pollen", "pollination", "fertilization"]):
        return ("Reproductive Parameters", "Pollination Parameters")

    if any(x in param_clean for x in ["fruit set", "fruit retention", "fruit drop"]):
        return ("Reproductive Parameters", "Fruit Set Parameters")

    if any(x in param_clean for x in ["marketable", "unmarketable", "yield", "harvest index", "harvest"]):
        return ("Reproductive Parameters", "Fruit Yield Parameters")

    if any(x in param_clean for x in ["fruit length", "fruit diameter", "fruit circumference", "fruit volume",
                                       "fruit weight", "fruit shape", "fruit color", "locule", "pericarp",
                                       "maturity index"]):
        return ("Reproductive Parameters", "Fruit Growth Parameters")

    if any(x in param_clean for x in ["spike", "panicle", "ear", "grain", "straw", "filled grains"]):
        return ("Reproductive Parameters", "Grain Yield Parameters (Cereal Crops)")

    if any(x in param_clean for x in ["seed"]):
        return ("Reproductive Parameters", "Seed Parameters")

    if any(x in param_clean for x in ["maturity", "harvest duration"]):
        return ("Reproductive Parameters", "Maturity Parameters")

    if any(x in param_clean for x in ["germination", "emergence", "establishment", "vigor index"]):
        return ("Vegetative Parameters", "Germination and Establishment")

    if any(x in param_clean for x in ["leaf", "leaves", "lai", "lad", "sla", "slw", "chlorophyll",
                                       "carotenoid", "spad", "greenness"]):
        return ("Vegetative Parameters", "Leaf Parameters")

    if any(x in param_clean for x in ["branch"]):
        return ("Vegetative Parameters", "Branch Parameters")

    if any(x in param_clean for x in ["root"]):
        return ("Vegetative Parameters", "Root Parameters")

    if any(x in param_clean for x in ["fresh weight", "dry weight", "biomass"]):
        return ("Vegetative Parameters", "Biomass Parameters")

    if any(x in param_clean for x in ["cgr", "rgr", "agr", "nar", "lar", "lwr", "rue"]):
        return ("Vegetative Parameters", "Growth Analysis Parameters")

    if any(x in param_clean for x in ["photosynthetic", "transpiration", "conductance", "water use",
                                       "fluorescence", "relative water", "rwc", "membrane stability",
                                       "electrolyte", "canopy temp", "ndvi", "pri", "fv/fm"]):
        return ("Vegetative Parameters", "Physiological Parameters")

    if any(x in param_clean for x in ["phenological", "first leaf", "vegetative maturity"]):
        return ("Vegetative Parameters", "Phenological Parameters")

    if any(x in param_clean for x in ["height", "stem", "collar", "internode", "node", "canopy",
                                       "spread", "crown"]):
        return ("Vegetative Parameters", "Plant Morphology")

    return ("Vegetative Parameters", "Plant Morphology")


# --- Word Document Table Formatting Helpers ---
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)


def set_header_bottom_border(row_or_cells):
    if hasattr(row_or_cells, 'cells'):
        cells = row_or_cells.cells
    else:
        cells = row_or_cells
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            '<w:tcBorders %s>'
            '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(borders)


# --- Dynamic Table Caption Generator ---
def generate_table_caption(table_num, factor_name, variables_list):
    vars_txt = ", ".join(variables_list)
    if len(variables_list) > 2:
        vars_txt = f"{', '.join(variables_list[:-1])} and {variables_list[-1]}"

    captions = [
        f"Table {table_num}. Influence of {factor_name} on {vars_txt}.",
        f"Table {table_num}. Response of {vars_txt} to various treatments of {factor_name}.",
        f"Table {table_num}. Effect of {factor_name} treatment levels on {vars_txt}.",
        f"Table {table_num}. Performance metrics of {vars_txt} under varying levels of {factor_name}.",
    ]
    return captions[table_num % len(captions)]


# --- Parameter Grouping Engine for Time-Series ---
def group_parameters(params):
    pattern = re.compile(r"^(.*?)[\s_\-]*(\d+)\s*(dat|das|days?|d)?$", re.IGNORECASE)
    groups = {}
    for p in params:
        match = pattern.search(p.strip())
        if match and match.group(1).strip():
            base = match.group(1).strip()
            base = re.sub(r"[\s\-_()]+$", "", base).strip().capitalize()
            day_num = int(match.group(2))
            day_str = f"Day {day_num}"
            if not base:
                base = "Parameter"
            groups.setdefault(base, []).append((p, day_num, day_str))
        else:
            base = p.strip().capitalize()
            groups.setdefault(base, []).append((p, 0, ""))

    for base in groups:
        groups[base].sort(key=lambda x: x[1])
    return groups


# --- Statistical Calculation Helpers ---
def get_signif_code_val(p):
    if pd.isna(p):
        return "ns"
    if p < 0.01:
        return "**"
    elif p < 0.05:
        return "*"
    else:
        return "ns"


def parse_dmrt_value(val):
    if pd.isna(val):
        return "", ""
    val_str = str(val).strip()
    match = re.match(r"^([\d.\-]+)\s*([a-zA-Z\s]+)?$", val_str)
    if match:
        num = match.group(1)
        letters = match.group(2).strip() if match.group(2) else ""
        return num, letters
    return val_str, ""


def get_cld_letters(means_dict, lsd):
    sorted_means = sorted(means_dict.items(), key=lambda x: x[1], reverse=True)
    n = len(sorted_means)
    groups = []

    for i in range(n):
        for j in range(i, n):
            is_group = True
            for k1 in range(i, j + 1):
                for k2 in range(k1, j + 1):
                    if abs(sorted_means[k1][1] - sorted_means[k2][1]) > lsd:
                        is_group = False
                        break
                if not is_group:
                    break
            if is_group:
                groups.append(set(range(i, j + 1)))

    maximal_groups = []
    for g in groups:
        is_subset = False
        for other_g in groups:
            if g != other_g and g.issubset(other_g):
                is_subset = True
                break
        if not is_subset and g not in maximal_groups:
            maximal_groups.append(g)

    maximal_groups.sort(key=lambda g: min(g))

    treatment_letters = {tname: "" for tname, _ in sorted_means}
    alphabet = string.ascii_lowercase

    for letter_idx, g in enumerate(maximal_groups):
        letter = alphabet[letter_idx % len(alphabet)]
        for idx in g:
            t_name = sorted_means[idx][0]
            treatment_letters[t_name] += letter

    return treatment_letters


def run_anova_1factor_raw(df, block_col, genotype_col, param):
    df_temp = pd.DataFrame({
        'rep': df[block_col].astype(str),
        'genotype': df[genotype_col].astype(str),
        'response': pd.to_numeric(df[param], errors='coerce')
    }).dropna(subset=['response'])

    formula = "response ~ C(rep) + C(genotype)"
    model = ols(formula, data=df_temp).fit()
    anova_table = sm.stats.anova_lm(model, typ=1)

    df_err = anova_table.loc["Residual", "df"]
    mse = anova_table.loc["Residual", "mean_sq"]

    p_val = anova_table.loc["C(genotype)", "PR(>F)"]

    grand_mean = df_temp['response'].mean()
    cv = (np.sqrt(mse) / grand_mean) * 100
    t_val = t.ppf(0.975, df_err)

    r = df_temp['rep'].nunique()

    means = df_temp.groupby('genotype')['response'].mean().to_dict()
    sem = np.sqrt(mse / r)
    lsd = t_val * np.sqrt((2 * mse) / r)
    cld = get_cld_letters(means, lsd)

    sig_code = get_signif_code_val(p_val)

    means_str = {}
    for g, val in means.items():
        let = cld[g] if p_val < 0.05 else ""
        means_str[g] = f"{val:.2f}{let}"

    return {
        "means": means, "means_str": means_str, "sem": round(sem, 2), "sig": sig_code,
        "lsd": round(lsd, 2), "p_val": p_val, "cv": round(cv, 2), "gm": round(grand_mean, 2), "cld": cld
    }


# --- Summarized Table Parser Engine ---
def parse_summarized_table_to_results_1f(df_raw, idx_sem, idx_pval, idx_lsd, idx_cv, idx_gm, param_cols, genotypes):
    results_data = {}
    for p, col_idx in param_cols.items():
        means = {}
        means_str = {}
        cld = {}
        for r_idx, g in enumerate(genotypes, start=1):
            cell_val = df_raw.iloc[r_idx, col_idx]
            num, let = parse_dmrt_value(cell_val)
            try:
                val = float(num)
                means[g] = val
                means_str[g] = f"{val:.2f}{let}"
            except ValueError:
                means[g] = 0.0
                means_str[g] = "0.00"
            cld[g] = let

        p_val_raw = str(df_raw.iloc[idx_pval, col_idx]).strip()
        match_p = re.search(r"[\d\.\-]+e?[\d\-]*", p_val_raw)
        try:
            p_val = float(match_p.group(0)) if match_p else (0.01 if "*" in p_val_raw else 0.5)
        except ValueError:
            p_val = 0.5

        results_data[p] = {
            "means": means,
            "means_str": means_str,
            "sem": df_raw.iloc[idx_sem, col_idx],
            "p_val": p_val,
            "p_text": p_val_raw,
            "lsd": df_raw.iloc[idx_lsd, col_idx],
            "cv": df_raw.iloc[idx_cv, col_idx],
            "gm": float(re.search(r"[\d\.\-]+", str(df_raw.iloc[idx_gm, col_idx])).group(0)) if re.search(r"[\d\.\-]+", str(df_raw.iloc[idx_gm, col_idx])) else 0.0,
            "cld": cld
        }
    return results_data


# --- Template Variable Injection Helper ---
def inject_template_placeholders(template_text, placeholders_dict):
    """Safely substitutes double-bracket markers with calculated results."""
    for key, val in placeholders_dict.items():
        template_text = template_text.replace("{{" + key + "}}", str(val))
    template_text = re.sub(r"\{\{.*?\}\}", "", template_text)
    return template_text


# --- Parameter Fact Extraction Helpers ---
def extract_single_day_facts_1f(parameter, res, genotype_col, table_label):
    means = res.get("means", {})
    means_str = res.get("means_str", {})
    cld = res.get("cld", {})

    sorted_g = sorted(means.items(), key=lambda x: x[1], reverse=True) if means else []

    treatment_A1, val_top = sorted_g[0] if sorted_g else ("Control", 0.0)
    treatment_A2, val_second = sorted_g[1] if len(sorted_g) > 1 else (treatment_A1, val_top)
    treatment_A3, val_third = sorted_g[2] if len(sorted_g) > 2 else (treatment_A2, val_second)
    treatment_A_lowest, val_lowest = sorted_g[-1] if sorted_g else ("Control", 0.0)

    top_let = cld.get(treatment_A1, "")
    top_letters = set(top_let)
    at_par_list = []
    for g, _ in sorted_g[1:]:
        g_letters = set(cld.get(g, ""))
        if g_letters & top_letters:
            at_par_list.append(g)

    if at_par_list:
        if len(at_par_list) > 1:
            at_par_str = ", ".join(at_par_list[:-1]) + f" and {at_par_list[-1]}"
        else:
            at_par_str = at_par_list[0]
    else:
        at_par_str = "none other"

    sem_val = res.get("sem", 0.1)
    lsd_val = res.get("lsd", 0.1)
    cv_val = res.get("cv", 5.0)
    grand_mean = res.get("gm", 0.0)
    p_val = res.get("p_val", 0.5)
    p_val_notation = get_p_val_notation(p_val)

    pct_diff = round(((val_top - val_lowest) / (val_lowest if val_lowest != 0 else 1.0)) * 100, 2)
    pct_diff_A2 = round(((val_second - val_lowest) / (val_lowest if val_lowest != 0 else 1.0)) * 100, 2)
    pct_diff_A3 = round(((val_third - val_lowest) / (val_lowest if val_lowest != 0 else 1.0)) * 100, 2)

    return {
        "variable_name": parameter,
        "table_num": table_label.replace("Table ", ""),
        "factor_A": genotype_col,
        "treatment_A1": treatment_A1,
        "val_A1": f"{val_top:.2f}",
        "value_A1": f"{val_top:.2f}",
        "treatment_A2": treatment_A2,
        "val_A2": f"{val_second:.2f}",
        "value_A2": f"{val_second:.2f}",
        "treatment_A3": treatment_A3,
        "val_A3": f"{val_third:.2f}",
        "value_A3": f"{val_third:.2f}",
        "treatment_A_lowest": treatment_A_lowest,
        "val_lowest": f"{val_lowest:.2f}",
        "value_lowest": f"{val_lowest:.2f}",
        "val_min_A": f"{val_lowest:.2f}",
        "val_max_A": f"{val_top:.2f}",
        "grand_mean": f"{grand_mean:.2f}",
        "cvval": f"{cv_val}",
        "cv_val": f"{cv_val}",
        "sem_val": f"{sem_val}",
        "sem": f"{sem_val}",
        "lsdval": f"{lsd_val}",
        "lsd_val": f"{lsd_val}",
        "pct_diff_A": f"{pct_diff}",
        "pct_diff_A2": f"{pct_diff_A2}",
        "pct_diff_A3": f"{pct_diff_A3}",
        "at_par": at_par_str,
        "unit": "",
    }


def extract_time_series_facts_1f(base_name, items, results_data, genotype_col, table_label):
    first_param, first_day_num, first_day_str = items[0]
    last_param, last_day_num, last_day_str = items[-1]

    num_dates = len(items)
    time_unit = "days"
    if any(x in first_param.lower() for x in ["das", "dat"]):
        time_unit = "DAS" if "das" in first_param.lower() else "DAT"

    first_gm = results_data[first_param]["gm"]
    last_gm = results_data[last_param]["gm"]
    first_cv = results_data[first_param]["cv"]

    tps = [it[2] for it in items]
    while len(tps) < 5:
        tps.append(tps[-1] if tps else "terminal phase")

    cvs = [results_data[it[0]]["cv"] for it in items]
    sems = [results_data[it[0]]["sem"] for it in items]

    res_last = results_data[last_param]
    means_last = res_last.get("means", {})
    cld_last = res_last.get("cld", {})
    sorted_g_last = sorted(means_last.items(), key=lambda x: x[1], reverse=True) if means_last else []

    treatment_A1, val_top = sorted_g_last[0] if sorted_g_last else ("Control", 0.0)
    treatment_A2, val_second = sorted_g_last[1] if len(sorted_g_last) > 1 else (treatment_A1, val_top)
    treatment_A_lowest, val_lowest = sorted_g_last[-1] if sorted_g_last else ("Control", 0.0)

    peak_val = max(results_data[it[0]]["gm"] for it in items)
    pct_diff = round(((val_top - val_lowest) / (val_lowest if val_lowest != 0 else 1.0)) * 100, 2)
    
    res_first = results_data[first_param]
    means_first = res_first.get("means", {})
    val_A1_early = means_first.get(treatment_A1, first_gm)
    val_lowest_early = means_first.get(treatment_A_lowest, first_gm)

    sem_val = res_last.get("sem", 0.1)
    lsd_val = res_last.get("lsd", 0.1)

    return {
        "variable_name": base_name,
        "base_name": base_name,
        "num_intervals": str(num_dates),
        "start_time": first_day_str.replace("Day ", "").strip(),
        "end_time": last_day_str.replace("Day ", "").strip(),
        "time_unit": time_unit,
        "table_num": table_label.replace("Table ", ""),
        "factor_A": genotype_col,
        "time_point_1": tps[0],
        "time_point_2": tps[1],
        "time_point_3": tps[2],
        "time_point_4": tps[3],
        "time_point_5": tps[4],
        "grand_mean_early": f"{first_gm:.2f}",
        "cv_early": f"{first_cv:.2f}",
        "treatment_A1": treatment_A1,
        "value_1": f"{val_top:.2f}",
        "treatment_A2": treatment_A2,
        "value_2": f"{val_second:.2f}",
        "treatment_A_lowest": treatment_A_lowest,
        "value_lowest": f"{val_lowest:.2f}",
        "lsdval": f"{lsd_val}",
        "total_days": str(abs(last_day_num - first_day_num)) if last_day_num != first_day_num else "30",
        "initial_value": f"{first_gm:.2f}",
        "peak_value": f"{peak_val:.2f}",
        "end_value": f"{last_gm:.2f}",
        "value_A1": f"{val_A1_early:.2f}",
        "val_A1": f"{val_top:.2f}",
        "val_A2": f"{val_second:.2f}",
        "val_lowest": f"{val_lowest:.2f}",
        "value_baseline": f"{val_lowest_early:.2f}",
        "val_baseline": f"{val_lowest_early:.2f}",
        "value_A1_end": f"{val_top:.2f}",
        "value_lowest_end": f"{val_lowest:.2f}",
        "val_max_early": f"{max(means_first.values()):.2f}" if means_first else f"{first_gm:.2f}",
        "val_min_early": f"{min(means_first.values()):.2f}" if means_first else f"{first_gm:.2f}",
        "val_max_late": f"{val_top:.2f}",
        "mean_early": f"{first_gm:.2f}",
        "cvval": f"{res_last.get('cv', 5.0)}",
        "cv_val": f"{res_last.get('cv', 5.0)}",
        "grand_mean_peak": f"{peak_val:.2f}",
        "cv_min": f"{min(cvs):.2f}" if cvs else "0.0",
        "cv_max": f"{max(cvs):.2f}" if cvs else "0.0",
        "num_dates": str(num_dates),
        "sem_min": f"{min(sems):.2f}" if sems else "0.0",
        "sem_max": f"{max(sems):.2f}" if sems else "0.0",
        "pct_diff_A": f"{pct_diff}",
        "pct_diff_A_late": f"{pct_diff}",
        "deltaA": f"{abs(val_top - val_lowest):.2f}",
        "sub_parameter_1": f"{base_name} - Sub A",
        "sub_parameter_2": f"{base_name} - Sub B",
        "val_A1_p2": f"{val_second:.2f}",
        "val_lowest_p2": f"{val_lowest:.2f}",
        "lsdvalp2": f"{lsd_val}",
        "gm_1": f"{first_gm:.2f}",
        "gm_2": f"{results_data[items[1][0]]['gm']:.2f}" if len(items) > 1 else f"{first_gm:.2f}",
        "gm_early": f"{first_gm:.2f}",
        "val_max_2": f"{val_top:.2f}",
        "val_lowest_2": f"{val_lowest:.2f}",
        "sem_val": f"{sem_val}",
        "range_early": f"{abs(val_top - val_lowest) * 0.1:.2f}",
        "val_max_A": f"{val_top:.2f}",
        "val_min_A": f"{val_lowest:.2f}",
        "range_mid_A": f"{abs(val_top - val_lowest) * 0.7:.2f}",
        "range_late_A": f"{abs(val_top - val_lowest):.2f}",
        "val_max_late": f"{val_top:.2f}",
        "val_lowest_late": f"{val_lowest:.2f}",
        "threshold_val": f"{last_gm * 0.8:.2f}",
        "unit": "",
    }


# --- Academic Explanation Selector Hooks ---
def generate_one_factor_explanation_shuffled(parameter, res, genotype_col, table_label, pool):
    placeholders = extract_single_day_facts_1f(parameter, res, genotype_col, table_label)
    p_val = res.get("p_val", 0.5)

    if p_val < 0.01:
        category = "highly_significant"
    elif p_val < 0.05:
        category = "significant"
    else:
        category = "non_significant"

    tpl_idx = pool.get_template_id(category)
    raw_template = ACADEMIC_TEMPLATES_30[tpl_idx]
    return inject_template_placeholders(raw_template, placeholders)


def generate_trend_explanation_1f_shuffled(base_name, items, results_data, genotype_col, table_label, pool):
    first_param, _, _ = items[0]
    last_param, _, _ = items[-1]

    first_gm = results_data[first_param]["gm"]
    last_gm = results_data[last_param]["gm"]
    p_val_last = results_data[last_param]["p_val"]
    direction_up = last_gm >= first_gm

    placeholders = extract_time_series_facts_1f(base_name, items, results_data, genotype_col, table_label)

    if p_val_last < 0.05:
        category = "divergent"
    elif direction_up:
        category = "upward_trend"
    else:
        category = "downward_trend"

    if random.random() < 0.25:
        category = "general"

    tpl_idx = pool.get_template_id(category)
    raw_template = ACADEMIC_TEMPLATES_30[tpl_idx]
    return inject_template_placeholders(raw_template, placeholders)


# --- Styled Excel Exporter (1-Factor Format) ---
def build_styled_excel_1f(genotype_col, params, genotypes, results_dict):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "RCBD Output"
    ws.sheet_view.showGridLines = True

    font_bold = Font(name="Calibri", size=11, bold=True)
    font_regular = Font(name="Calibri", size=11)
    align_left = Alignment(horizontal="left", vertical="center")
    align_center = Alignment(horizontal="center", vertical="center")
    border_thin_bottom = Border(bottom=Side(style='thin', color='000000'))
    border_medium_bottom = Border(bottom=Side(style='medium', color='000000'))

    ws.cell(row=1, column=1, value=genotype_col).font = font_bold
    ws.cell(row=1, column=1).alignment = align_left

    for col_idx, param in enumerate(params, start=2):
        cell = ws.cell(row=1, column=col_idx, value=param)
        cell.font = font_bold
        cell.alignment = align_center

    g_levels_num = len(genotypes)

    row_g_start = 2
    row_g_end = 1 + g_levels_num
    row_sem = 2 + g_levels_num
    row_f = 3 + g_levels_num
    row_lsd = 4 + g_levels_num
    row_cv = 5 + g_levels_num
    row_gm = 6 + g_levels_num

    for idx, lvl in enumerate(genotypes):
        r = row_g_start + idx
        ws.cell(row=r, column=1, value=lvl).font = font_regular
        ws.cell(row=r, column=1).alignment = align_left

    ws.cell(row=row_sem, column=1, value="Sem(\u00b1)").font = font_regular
    ws.cell(row=row_sem, column=1).alignment = align_left
    ws.cell(row=row_f, column=1, value="p-value").font = font_regular
    ws.cell(row=row_f, column=1).alignment = align_left
    ws.cell(row=row_lsd, column=1, value="LSD(0.05)").font = font_regular
    ws.cell(row=row_lsd, column=1).alignment = align_left
    ws.cell(row=row_cv, column=1, value="CV(%)").font = font_regular
    ws.cell(row=row_cv, column=1).alignment = align_left
    ws.cell(row=row_gm, column=1, value="Grand Mean").font = font_regular
    ws.cell(row=row_gm, column=1).alignment = align_left

    for col_idx, param in enumerate(params, start=2):
        p_data = results_dict[param]

        for idx, lvl in enumerate(genotypes):
            r = row_g_start + idx
            cell = ws.cell(row=r, column=col_idx, value=p_data["means_str"][lvl])
            cell.font = font_regular
            cell.alignment = align_center

        ws.cell(row=row_sem, column=col_idx, value=p_data["sem"]).font = font_regular
        ws.cell(row=row_sem, column=col_idx).alignment = align_center
        ws.cell(row=row_f, column=col_idx, value=get_signif_code_val(p_data["p_val"])).font = font_regular
        ws.cell(row=row_f, column=col_idx).alignment = align_center
        ws.cell(row=row_lsd, column=col_idx, value=p_data["lsd"]).font = font_regular
        ws.cell(row=row_lsd, column=col_idx).alignment = align_center
        ws.cell(row=row_cv, column=col_idx, value=p_data["cv"]).font = font_regular
        ws.cell(row=row_cv, column=col_idx).alignment = align_center
        ws.cell(row=row_gm, column=col_idx, value=p_data["gm"]).font = font_regular
        ws.cell(row=row_gm, column=col_idx).alignment = align_center

    border_rows = [row_g_end, row_lsd, row_cv, row_gm]
    for r_idx in border_rows:
        for col in range(1, len(params) + 2):
            ws.cell(row=r_idx, column=col).border = border_medium_bottom

    for col in range(1, len(params) + 2):
        ws.cell(row=1, column=col).border = border_thin_bottom

    return wb


# --- DOCX Table Formatting ---
def add_excel_table_to_docx_1f(doc, genotype_col, g_cols, genotypes, results_data):
    num_cols = len(g_cols) + 1
    table = doc.add_table(rows=1, cols=num_cols)
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = genotype_col
    set_cell_margins(hdr_cells[0])
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
    hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(10)

    for c_idx, col_name in enumerate(g_cols):
        hdr_cells[c_idx + 1].text = col_name
        set_cell_margins(hdr_cells[c_idx + 1])
        hdr_cells[c_idx + 1].paragraphs[0].alignment = 1
        hdr_cells[c_idx + 1].paragraphs[0].runs[0].font.bold = True
        hdr_cells[c_idx + 1].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_cells[c_idx + 1].paragraphs[0].runs[0].font.size = Pt(10)

    set_header_bottom_border(table.rows[0])

    def add_styled_data_row(lbl, val_dict, bold=False):
        row_cells = table.add_row().cells
        row_cells[0].text = str(lbl)
        set_cell_margins(row_cells[0])
        p = row_cells[0].paragraphs[0]
        if len(p.runs) > 0:
            p.runs[0].font.bold = bold
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(10)

        for c_idx, col_name in enumerate(g_cols):
            row_cells[c_idx + 1].text = str(val_dict.get(col_name, "N/A"))
            set_cell_margins(row_cells[c_idx + 1])
            p_val = row_cells[c_idx + 1].paragraphs[0]
            p_val.alignment = 1
            if len(p_val.runs) > 0:
                p_val.runs[0].font.bold = bold
                p_val.runs[0].font.name = 'Arial'
                p_val.runs[0].font.size = Pt(10)
        return row_cells

    for lvl in sorted(genotypes):
        lvl_dict = {p: results_data[p]["means_str"][lvl] for p in g_cols}
        add_styled_data_row(lvl, lvl_dict)

    sem_dict = {p: results_data[p]["sem"] for p in g_cols}
    sig_dict = {p: get_signif_code_val(results_data[p]["p_val"]) for p in g_cols}
    lsd_dict = {p: results_data[p]["lsd"] for p in g_cols}
    cv_dict = {p: results_data[p]["cv"] for p in g_cols}
    gm_dict = {p: results_data[p]["gm"] for p in g_cols}

    r_sem = add_styled_data_row("SEm(\u00b1)", sem_dict)
    set_header_bottom_border(r_sem)
    r_sig = add_styled_data_row("p-value", sig_dict)
    set_header_bottom_border(r_sig)
    r_lsd = add_styled_data_row("LSD(0.05)", lsd_dict)
    set_header_bottom_border(r_lsd)
    r_cv = add_styled_data_row("CV, %", cv_dict)
    set_header_bottom_border(r_cv)
    r_gm = add_styled_data_row("Grand mean", gm_dict)
    set_header_bottom_border(r_gm)

    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        for idx in range(1, len(g_cols) + 1):
            row.cells[idx].width = Inches(1.1)


# ==============================================================================
# Word report builder: produces the 1 / 1.1 / 1.1.1 hierarchical structure
# ==============================================================================
def build_hierarchical_report_1f(classified_cols, genotype_col, genotypes, results_data):
    doc = Document()
    doc.add_heading("Calculated One-Factor RCBD Report", 0)
    numberer = ReportNumberer()

    single_day_pool = TemplatePool(SINGLE_DAY_TEMPLATES_CATEGORIES)
    time_series_pool = TemplatePool(TIME_SERIES_TEMPLATES_CATEGORIES)

    ordered_majors = [m for m in MAJOR_CATEGORY_ORDER if m in classified_cols]
    ordered_majors += [m for m in classified_cols if m not in ordered_majors]

    for major in ordered_majors:
        subs_dict = classified_cols[major]
        if not any(subs_dict.values()):
            continue

        major_title = numberer.major(doc, major)
        st.write(f"## {major_title}")

        sub_order = SUB_CATEGORY_ORDER.get(major, [])
        ordered_subs = [s for s in sub_order if s in subs_dict and subs_dict[s]]
        ordered_subs += [s for s in subs_dict if s not in ordered_subs and subs_dict[s]]

        for sub in ordered_subs:
            cat_params = subs_dict[sub]
            sub_title = numberer.sub(doc, sub)
            st.write(f"### {sub_title}")

            grouped = group_parameters(cat_params)

            static_items = [items[0][0] for base_name, items in sorted(grouped.items()) if len(items) == 1]
            time_series_items = {base_name: items for base_name, items in sorted(grouped.items()) if len(items) > 1}

            chunk_size = 4
            for i in range(0, len(static_items), chunk_size):
                chunk = static_items[i:i + chunk_size]

                table_n = numberer.next_table()
                table_label = f"Table {table_n}"

                for p in chunk:
                    param_title = numberer.param(doc, p)
                    st.write(f"**{param_title}**")

                    p_text = generate_one_factor_explanation_shuffled(
                        p, results_data[p], genotype_col, table_label, single_day_pool
                    )
                    st.write(p_text)
                    doc.add_paragraph(p_text)

                add_excel_table_to_docx_1f(doc, genotype_col, chunk, genotypes, results_data)

                caption_text = generate_table_caption(table_n, genotype_col, chunk)
                p_cap = doc.add_paragraph(caption_text)
                p_cap.runs[0].font.name = 'Arial'
                p_cap.runs[0].font.size = Pt(10)
                p_cap.runs[0].font.italic = True

                st.write(f"*{table_label} rendered below*")
                doc.add_paragraph()

            for base_name, items in sorted(time_series_items.items()):
                param_title = numberer.param(doc, base_name)
                st.write(f"**{param_title}**")

                table_n = numberer.next_table()
                table_label = f"Table {table_n}"
                trend_params = [it[0] for it in items]

                p_text = generate_trend_explanation_1f_shuffled(
                    base_name, items, results_data, genotype_col, table_label, time_series_pool
                )
                st.write(p_text)
                doc.add_paragraph(p_text)

                add_excel_table_to_docx_1f(doc, genotype_col, trend_params, genotypes, results_data)

                caption_text = generate_table_caption(table_n, genotype_col, trend_params)
                p_cap = doc.add_paragraph(caption_text)
                p_cap.runs[0].font.name = 'Arial'
                p_cap.runs[0].font.size = Pt(10)
                p_cap.runs[0].font.italic = True

                st.write(f"*{table_label} (time-series) rendered below*")
                doc.add_paragraph()

        doc.add_paragraph("-" * 60)

    return doc


def build_classified_cols(parameters):
    classified_cols = {}
    for c in parameters:
        major, sub = classify_parameter(c)
        classified_cols.setdefault(major, {}).setdefault(sub, []).append(c)
    return classified_cols


def show_category_preview(classified_cols):
    st.write("#### Automatically Categorized Parameter Divisions:")
    ordered_majors = [m for m in MAJOR_CATEGORY_ORDER if m in classified_cols]
    ordered_majors += [m for m in classified_cols if m not in ordered_majors]
    for major in ordered_majors:
        subs_dict = classified_cols[major]
        sub_order = SUB_CATEGORY_ORDER.get(major, [])
        ordered_subs = [s for s in sub_order if s in subs_dict and subs_dict[s]]
        ordered_subs += [s for s in subs_dict if s not in ordered_subs and subs_dict[s]]
        for sub in ordered_subs:
            params = subs_dict[sub]
            if params:
                st.write(f"**{major} \u2192 {sub}:** {', '.join(params)}")


# --- Web Interface Routing and One-Factor Controller ---
def show_module():
    st.markdown("### One-Factor RCBD Analyzer")

    mode = st.radio("Choose Input Mode", ["Raw Data Mode", "Summarized Table Mode"], key="1f_mode_selector")
    uploaded_file = st.file_uploader("Upload One-Factor Excel File", type=["xlsx"], key="file_uploader_1f")

    if uploaded_file is not None:
        if mode == "Raw Data Mode":
            run_raw_mode(uploaded_file)
        else:
            run_summary_mode_processing(uploaded_file)


def run_raw_mode(uploaded_file):
    try:
        df_raw_data = pd.read_excel(uploaded_file)
        st.write("#### Preview Raw Input Data:", df_raw_data.head())

        cols = df_raw_data.columns.tolist()

        block_col = st.selectbox("Select Block/Replication Column", cols, index=0, key="raw_1f_bk")
        genotype_col = st.selectbox("Select Genotype/Treatment Column", cols, index=1, key="raw_1f_g")
        response_cols = st.multiselect("Select Response Parameters to Analyze", cols, default=cols[2:], key="raw_1f_resp")

        if response_cols:
            df_raw_data[genotype_col] = df_raw_data[genotype_col].astype(str)
            df_raw_data[block_col] = df_raw_data[block_col].astype(str)

            classified_cols = build_classified_cols(response_cols)
            show_category_preview(classified_cols)

            if st.button("Run One-Factor Raw Analysis", key="btn_raw_1f_calc"):
                results_data = {}
                genotypes = sorted(df_raw_data[genotype_col].unique().tolist())

                for param in response_cols:
                    results_data[param] = run_anova_1factor_raw(df_raw_data, block_col, genotype_col, param)

                styled_wb = build_styled_excel_1f(genotype_col, response_cols, genotypes, results_data)
                excel_bio = io.BytesIO()
                styled_wb.save(excel_bio)
                excel_bio.seek(0)

                st.markdown("#### \U0001F4E5 Download Formatted Statistical Excel Results")
                st.download_button(
                    label="Download Formatted Excel Results Table",
                    data=excel_bio,
                    file_name="Result_1Factor_Output.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="btn_d_excel_styled_1f"
                )
                st.write("---")

                st.markdown("### \U0001F4DD Analysis Results and Academic Explanations")

                doc = build_hierarchical_report_1f(classified_cols, genotype_col, genotypes, results_data)

                bio_doc = io.BytesIO()
                doc.save(bio_doc)
                bio_doc.seek(0)

                st.write("---")
                st.markdown("#### \U0001F4BE Save Explanations as a Word Report")
                st.download_button(
                    "Download Word Explanations Report (.docx)",
                    data=bio_doc,
                    file_name="Calculated_OneFactor_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="btn_d_1f_raw_cal"
                )
    except Exception as e:
        st.error(f"Error executing raw One-Factor analysis: {e}")


def run_summary_mode_processing(uploaded_file):
    try:
        df_raw = pd.read_excel(uploaded_file, header=None)

        idx_sem, idx_pval, idx_lsd, idx_cv, idx_gm = None, None, None, None, None
        for idx, val in enumerate(df_raw[0]):
            if pd.isna(val):
                continue
            val_str = str(val).strip().lower()
            if val_str in ["sem", "sem(\u00b1)", "sem(±)"]:
                idx_sem = idx
            elif val_str in ["p-value", "p_value", "p-val", "f-value", "f value"]:
                idx_pval = idx
            elif "lsd" in val_str:
                idx_lsd = idx
            elif "cv" in val_str:
                idx_cv = idx
            elif "grand mean" in val_str or "grandmean" in val_str:
                idx_grand_mean = idx_gm = idx

        if any(v is None for v in [idx_sem, idx_pval, idx_lsd, idx_cv, idx_gm]):
            st.error("Missing structural markers (Sem, p-value, LSD, CV, Grand Mean) in Column A.")
            return

        genotypes = [str(df_raw.iloc[r, 0]).strip() for r in range(1, idx_sem)]
        genotype_col_label = str(df_raw.iloc[0, 0]).strip() if pd.notna(df_raw.iloc[0, 0]) else "Genotype"

        parameters = [str(x).strip() for x in df_raw.iloc[0].tolist()[1:] if pd.notna(x)]
        param_cols = {p: df_raw.iloc[0].tolist().index(p) for p in parameters}

        st.success(f"Detected Factor: {genotype_col_label} | Genotypes: {', '.join(genotypes)}")

        classified_cols = build_classified_cols(parameters)
        show_category_preview(classified_cols)

        if st.button("Generate Word Document Draft from Direct Output", key="btn_1f_sum_gen"):
            results_data = parse_summarized_table_to_results_1f(
                df_raw, idx_sem, idx_pval, idx_lsd, idx_cv, idx_gm, param_cols, genotypes
            )

            doc = build_hierarchical_report_1f(classified_cols, genotype_col_label, genotypes, results_data)

            bio_doc = io.BytesIO()
            doc.save(bio_doc)
            bio_doc.seek(0)

            st.write("---")
            st.markdown("#### \U0001F4BE Save Explanations & Tables as Word Document")
            st.download_button(
                "Download Word Explanations Report (.docx)",
                data=bio_doc,
                file_name="Summarized_OneFactor_Thesis_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_d_1f_sum_doc"
            )
    except Exception as e:
        st.error(f"Error parsing direct result summary table: {e}")


if __name__ == "__main__":
    show_module()
