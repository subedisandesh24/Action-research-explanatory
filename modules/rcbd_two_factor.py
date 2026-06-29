import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import string
import random
from scipy.stats import t
import statsmodels.api as sm
from statsmodels.formula.api import ols
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side

# ==============================================================================
# DATABASE OF VERBOSE, HIGH-STANDARD ACADEMIC DISCUSSION TEMPLATES (30 TEMPLATES)
# ==============================================================================
ACADEMIC_TEMPLATES_30 = {
    # --- Group A: Time-Series / Chronological / Trend Parameters (1–15) ---
    1: (
        "The temporal development of {{variable_name}} was monitored sequentially across {{num_intervals}} "
        "experimental phases from {{start_time}} to {{end_time}} {{time_unit}}. Table {{table_num}} shows that during "
        "the initial evaluation stages, specifically at {{time_point_1}} and {{time_point_2}} {{time_unit}}, the main "
        "effects of {{factor_A}} and {{factor_B}}, as well as their interactive combinations (A \u00d7 B), did not "
        "exert any statistically significant influence on the observed values (P > 0.05). The overall grand mean during "
        "this early phase remained stable at {{grand_mean_early}} {{unit}}. However, as the experimental units "
        "progressed, a distinct divergence among treatments emerged. Table {{table_num}} shows that at {{time_point_3}} "
        "{{time_unit}}, the main effect of {{factor_B}} was significant (P \u2264 0.05), with {{treatment_B1}} "
        "({{value_B1}} {{unit}}) significantly outperforming {{treatment_B2}} ({{value_B2}} {{unit}}). By {{time_point_4}} "
        "{{time_unit}}, both main effects reached a highly significant state (P \u2264 0.01), which was accompanied by "
        "a significant interaction (P \u2264 0.05), demonstrating that the response of {{variable_name}} to {{factor_B}} "
        "became highly dependent on the choice of {{factor_A}} during the final monitoring intervals."
    ),
    2: (
        "The overall timeline of {{variable_name}} followed a classic parabolic peak-and-decline trajectory over the "
        "{{total_days}}-day study. Table {{table_num}} shows that the grand mean of {{variable_name}} began at "
        "{{initial_value}} {{unit}} ({{time_point_1}} {{time_unit}}), peaked at {{peak_value}} {{unit}} at {{time_point_2}} "
        "{{time_unit}}, and subsequently declined to {{end_value}} {{unit}} at the terminal evaluation phase of "
        "{{time_point_3}} {{time_unit}}. Although treatments were statistically equivalent at the beginning of the "
        "evaluation stage, the rate of post-peak decline was significantly altered by the application of both "
        "{{factor_A}} and {{factor_B}}. Table {{table_num}} shows that at {{time_point_2}} {{time_unit}}, {{treatment_A1}} "
        "maintained significantly higher {{variable_name}} levels ({{value_A1}} {{unit}}) compared to the baseline "
        "treatment ({{value_baseline}} {{unit}}; P \u2264 0.05). This treatment trend widened toward the final "
        "sampling stage of {{time_point_3}} {{time_unit}}, where {{treatment_B1}} ({{value_B1}} {{unit}}) significantly "
        "limited the decrease of {{variable_name}} compared to {{treatment_B2}} ({{value_B2}} {{unit}}; P \u2264 0.01), "
        "highlighting the role of the treatments in slowing the downward progression of the evaluated parameter."
    ),
    3: (
        "Table {{table_num}} shows that the interaction effect between {{factor_A}} and {{factor_B}} was "
        "statistically non-significant at the early dates of {{time_point_1}} and {{time_point_2}} {{time_unit}}, but "
        "transitioned to a significant interactive response at {{time_point_3}} {{time_unit}} (P \u2264 0.05) and "
        "{{time_point_4}} {{time_unit}} (P \u2264 0.05). Under this significant interaction, the combined "
        "application of {{treatment_A1}} and {{treatment_B1}} produced the highest overall value of {{max_val}} {{unit}}, "
        "which was significantly higher than all other treatment combinations. For the final measurement date at "
        "{{time_point_5}} {{time_unit}}, the interaction returned to non-significant status (P > 0.05). However, both "
        "main effects remained highly significant on this final date. Table {{table_num}} shows that {{treatment_A1}} "
        "({{val_A1}} {{unit}}) and {{treatment_B1}} ({{val_B1}} {{unit}}) independently preserved higher levels of "
        "{{variable_name}} compared to their respective minimum treatments ({{val_minA}} and {{val_minB}} {{unit}}), "
        "indicating a sustained, independent treatment effect over the experimental timeline."
    ),
    4: (
        "The main effect of {{factor_A}} was the primary source of variation for {{variable_name}} across "
        "almost the entire temporal scale, whereas the effect of {{factor_B}} emerged only during the final "
        "sampling points. Table {{table_num}} shows that {{factor_A}} had a highly significant effect "
        "(P \u2264 0.01) from {{time_point_1}} to {{time_point_4}} {{time_unit}}, with {{treatment_A1}} "
        "consistently maintaining the highest level of {{variable_name}} (averaging {{avg_val}} {{unit}}) compared "
        "to the other treatments. In contrast, the main effect of {{factor_B}} remained non-significant "
        "(P > 0.05) until the final stage of {{time_point_4}} {{time_unit}}, where {{treatment_B1}} finally yielded "
        "a significantly higher value of {{val_B1}} {{unit}} compared to {{treatment_B2}} ({{val_B2}} {{unit}}; "
        "P \u2264 0.05). The interactive effect between {{factor_A}} and {{factor_B}} was non-significant "
        "throughout the entire temporal timeline, indicating that these two factors operated independently "
        "across the observation period."
    ),
    5: (
        "The temporal response of {{variable_name}} was characterized by a distinct statistical "
        "phase-shift between the initial phase ({{time_point_1}} to {{time_point_2}} {{time_unit}}) and the "
        "final phase ({{time_point_3}} to {{time_point_4}} {{time_unit}}). Table {{table_num}} shows that "
        "during the initial phase, all treatment combinations shared identical letter superscripts "
        "(P > 0.05), with values closely matching the baseline mean of {{mean_early}} {{unit}}. Following "
        "the transition to the final phase at {{time_point_3}} {{time_unit}}, the main effect of {{factor_A}} "
        "became highly significant (P \u2264 0.01). Table {{table_num}} shows that {{treatment_A1}} produced "
        "the highest value ({{val_A1}} {{unit}}), which was statistically equivalent only to {{treatment_A2}} "
        "({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}; "
        "LSD_0.05 = {{lsdval}}). This significant main effect persisted through the final evaluation stage at "
        "{{time_point_4}} {{time_unit}}, while the interactive combinations remained non-significant."
    ),
    6: (
        "The rate of decrease in {{variable_name}} following the peak phase was significantly "
        "influenced by the application of the treatment factors. Table {{table_num}} shows that from "
        "{{time_point_1}} to {{time_point_2}} {{time_unit}}, the experimental units maintained high stability "
        "with a grand mean of {{grand_mean_peak}} {{unit}}, and no treatment variations were observed "
        "(P > 0.05). However, from {{time_point_3}} {{time_unit}} onwards, a rapid downward trend was "
        "observed in {{variable_name}} values. The application of {{treatment_A1}} significantly slowed this "
        "decrease, maintaining a value of {{val_A1}} {{unit}} at {{time_point_4}} {{time_unit}}, compared to "
        "the baseline treatment which dropped to {{val_baseline}} {{unit}} (P \u2264 0.05). This protective main "
        "effect was complemented by the highly significant main effect of {{treatment_B1}} ({{val_B1}} {{unit}}; "
        "P \u2264 0.01), showing that the treatments worked to slow the downward progression through the final "
        "sampling stages."
    ),
    7: (
        "The cumulative progression of {{variable_name}} was recorded sequentially from {{time_point_1}} "
        "to {{time_point_4}} {{time_unit}} to evaluate the accumulation kinetics of the system. Table {{table_num}} "
        "shows that the cumulative curve followed a standard sigmoidal pattern, with the fastest accumulation rate "
        "observed between {{time_point_2}} and {{time_point_3}} {{time_unit}}. Although there were no "
        "differences among treatments during the initial lag phase at {{time_point_1}} {{time_unit}}, both main "
        "factors significantly affected the accumulation rate during the active phase. At {{time_point_3}} "
        "{{time_unit}}, {{treatment_A1}} achieved a significantly higher cumulative value ({{val_A1}} {{unit}}) "
        "compared to the baseline treatment ({{val_baseline}} {{unit}}; P \u2264 0.01). Similarly, the main "
        "effect of {{treatment_B1}} was highly significant at this stage, with a value of {{val_B1}} {{unit}} "
        "compared to {{val_B2}} {{unit}} (P \u2264 0.01). This significant difference was maintained through the "
        "plateau phase at {{time_point_4}} {{time_unit}}, while the interaction effect remained non-significant."
    ),
    8: (
        "The experimental design demonstrated high precision in monitoring the temporal dynamics "
        "of {{variable_name}}, with Coefficient of Variation (CV%) values ranging from {{cv_min}}% to {{cv_max}}% "
        "across the {{num_dates}} sampling dates. Table {{table_num}} shows that the Standard Error of the Mean "
        "(SEm) for both factor levels remained low, ranging between {{sem_min}} and {{sem_max}} {{unit}}. "
        "Under these precise conditions, treatment differences remained non-significant at {{time_point_1}} and "
        "{{time_point_2}} {{time_unit}}. At {{time_point_3}} {{time_unit}}, the main effect of {{factor_B}} was "
        "significant (P \u2264 0.05, LSD_0.05 = {{lsdB}}), with {{treatment_B1}} ({{val_B1}} {{unit}}) outperforming "
        "{{treatment_B2}} ({{val_B2}} {{unit}}). By {{time_point_4}} {{time_unit}}, both the main effect of "
        "{{factor_A}} (P \u2264 0.05, LSD_0.05 = {{lsdA}}) and {{factor_B}} (P \u2264 0.01, LSD_0.05 = "
        "{{lsdB}}) were significant, while their interactive effects remained statistically non-significant."
    ),
    9: (
        "The mathematical differences (\u0394) and percentage changes in {{variable_name}} were "
        "evaluated across several sampling intervals. Table {{table_num}} shows that while treatments "
        "remained statistically equivalent at {{time_point_1}} and {{time_point_2}} {{time_unit}}, significant "
        "treatment differences emerged at {{time_point_3}} {{time_unit}}. At this stage, {{treatment_A1}} "
        "increased {{variable_name}} by {{pct_diff_A}}% compared to the lowest treatment level (\u0394 = {{deltaA}} "
        "{{unit}}; P \u2264 0.05). At {{time_point_4}} {{time_unit}}, this percentage difference increased, with "
        "{{treatment_A1}} outperforming the lowest treatment level by {{pct_diff_A_late}}% (P \u2264 0.01). Concurrently, "
        "the application of {{treatment_B1}} resulted in a {{pct_diff_B}}% increase compared to {{treatment_B2}} "
        "(\u0394 = {{deltaB}} {{unit}}; P \u2264 0.01), demonstrating the growing effect of the treatments as the "
        "experimental units reached the final stages."
    ),
    10: (
        "The temporal development of {{variable_name}} showed stable values during the early evaluation stages, "
        "followed by significant treatment differences as the experimental units progressed. Table {{table_num}} shows "
        "that no significant variations were observed at {{time_point_1}} and {{time_point_2}} {{time_unit}}, with "
        "values remaining close to the grand mean of {{mean_early}} {{unit}}. However, at {{time_point_3}} "
        "{{time_unit}}, both the main effect of {{factor_A}} (P \u2264 0.05) and {{factor_B}} (P \u2264 0.01) "
        "were significant, with {{treatment_A1}} ({{val_A1}} {{unit}}) and {{treatment_B1}} ({{val_B1}} {{unit}}) "
        "outperforming their respective minimum levels. This significant main effect persisted through the final "
        "sampling date of {{time_point_4}} {{time_unit}}."
    ),
    11: (
        "Table {{table_num}} shows that the {{variable_name}} values maintained a state of statistical "
        "equivalence during the first two sampling dates ({{time_point_1}} and {{time_point_2}} {{time_unit}}), with "
        "all treatment combinations sharing the same post-hoc letter grouping 'a'. During this early period, the grand "
        "means were recorded as {{gm_1}} {{unit}} and {{gm_2}} {{unit}}, respectively, with no significant differences "
        "observed. However, a transition occurred at {{time_point_3}} {{time_unit}} when the main effect of {{factor_B}} "
        "application became significant (P \u2264 0.05). Treatment {{treatment_B1}} broke the statistical equivalence, "
        "producing {{val_B1}} {{unit}} (assigned to letter grouping 'a') compared to treatment {{treatment_B2}} which "
        "produced {{val_B2}} {{unit}} (assigned to letter grouping 'b'). By {{time_point_4}} {{time_unit}}, both "
        "{{factor_A}} (P \u2264 0.05) and {{factor_B}} (P \u2264 0.01) showed significant differences, breaking the "
        "early statistical equivalence across both factor levels."
    ),
    12: (
        "The interactive effects of different treatments on {{variable_name}} became significant "
        "during the middle of the monitoring cycle. Table {{table_num}} shows that the interaction (A \u00d7 B) "
        "was significant at {{time_point_1}} {{time_unit}} (P \u2264 0.05) and highly significant at {{time_point_2}} "
        "{{time_unit}} (P \u2264 0.01). Under this interactive response, the positive effect of {{treatment_B1}} "
        "was significantly enhanced when combined with {{treatment_A1}}, yielding the highest overall values of "
        "{{val_inter_1}} {{unit}} and {{val_inter_2}} {{unit}} on those dates. At the final sampling points of "
        "{{time_point_3}} and {{time_point_4}} {{time_unit}}, this interaction was non-significant (P > 0.05). Table "
        "{{table_num}} shows that despite the lack of interaction, both main effects remained highly significant on "
        "these final dates. Under these non-interactive conditions, {{treatment_A1}} ({{val_A1}} {{unit}}) and "
        "{{treatment_B1}} ({{val_B1}} {{unit}}) independently maintained the highest levels of {{variable_name}} "
        "compared to their respective minimum treatments."
    ),
    13: (
        "The range of values for {{variable_name}} across different treatments was small during "
        "the early stages but expanded considerably as the study progressed. Table {{table_num}} shows that at "
        "{{time_point_1}} and {{time_point_2}} {{time_unit}}, the range between the highest and lowest treatment "
        "means was narrow and statistically non-significant, with a maximum range of only {{range_early}} {{unit}}. "
        "However, at {{time_point_3}} {{time_unit}}, this treatment range expanded significantly. For Factor A, "
        "the range between the highest-performing treatment ({{treatment_A1}}; {{val_max_A}} {{unit}}) and the "
        "lowest-performing treatment ({{val_min_A}} {{unit}}) was {{range_mid_A}} {{unit}} (P \u2264 0.01). For Factor B, "
        "the range was {{range_mid_B}} {{unit}} (P \u2264 0.01). At {{time_point_4}} {{time_unit}}, the treatment "
        "range remained wide, with a difference of {{range_late_A}} {{unit}} among {{factor_A}} treatments "
        "(P \u2264 0.05) and {{range_late_B}} {{unit}} between {{factor_B}} treatments (P \u2264 0.01), demonstrating "
        "the cumulative effect of the treatments over the entire study."
    ),
    14: (
        "The {{variable_name}} of the experimental units was monitored across a series of "
        "developmental stages from the initial phase to the final terminal stage. Table {{table_num}} shows that "
        "no treatment differences were observed during the initial phase at {{time_point_1}} and {{time_point_2}} "
        "{{time_unit}} (P > 0.05). The first significant differences emerged during the middle phase at "
        "{{time_point_3}} {{time_unit}}, where the main effect of {{factor_B}} was significant (P \u2264 0.05), "
        "with {{treatment_B1}} producing a higher value ({{val_B1}} {{unit}}) than {{treatment_B2}} ({{val_B2}} "
        "{{unit}}). As the experimental system reached peak values at {{time_point_4}} {{time_unit}}, treatment "
        "differences became highly significant. Table {{table_num}} shows that both {{factor_A}} (P \u2264 0.01) "
        "and {{factor_B}} (P \u2264 0.01) had highly significant effects, and their interaction was "
        "significant (P \u2264 0.05). Under this interactive response, the combination of {{treatment_A1}} and "
        "{{treatment_B1}} produced the highest overall values. At the final stage of {{time_point_5}} {{time_unit}}, "
        "both main effects remained significant, although their interaction returned to non-significant status."
    ),
    15: (
        "The {{variable_name}} levels were monitored to identify when the experimental units "
        "transitioned past critical operational thresholds. Table {{table_num}} shows that during the "
        "initial stages at {{time_point_1}} and {{time_point_2}} {{time_unit}}, all treatments maintained high, "
        "stable values above the critical threshold of {{threshold_val}} {{unit}}, with no significant "
        "differences observed among them (P > 0.05). A critical transition occurred at {{time_point_3}} "
        "{{time_unit}}, where the values under the lower-performing treatments began to drop below the "
        "threshold level. The application of {{treatment_A1}} significantly delayed this downward trend, "
        "maintaining a value of {{val_A1}} {{unit}} (well above the threshold), while the lowest-performing "
        "treatment fell to {{val_lowest}} {{unit}} (P \u2264 0.05). This protective main effect was complemented "
        "by the highly significant main effect of {{treatment_B1}} ({{val_B1}} {{unit}}; P \u2264 0.01), "
        "showing that the treatments worked together to maintain the values of the system through the "
        "final sampling stages."
    ),

    # --- Group B: Single-Day / End-Point / Terminal Parameters (16–30) ---
    16: (
        "The terminal measurement of {{variable_name}} was significantly affected by both "
        "experimental factors and their interaction. Table {{table_num}} shows that the main effect of "
        "{{factor_A}} was highly significant (P \u2264 0.01), with {{treatment_A1}} producing the highest value of "
        "{{val_A1}} {{unit}}, which was statistically equivalent to {{treatment_A2}} ({{val_A2}} {{unit}}) "
        "but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsd_A}}). "
        "The main effect of {{factor_B}} was also highly significant (P \u2264 0.001), with {{treatment_B1}} "
        "producing a significantly higher value compared to {{treatment_B2}} ({{val_B1}} vs. {{val_B2}} {{unit}}; "
        "LSD_0.05 = {{lsd_B}}). A statistically significant interactive response (A \u00d7 B) was also observed "
        "for this parameter (P \u2264 0.05). Table {{table_num}} shows that under this significant interaction, "
        "the positive response to {{treatment_B1}} was significantly enhanced when combined with {{treatment_A1}}. "
        "Under this interactive response, the combination of {{treatment_A1}} and {{treatment_B1}} produced the "
        "highest overall value of {{max_val}} {{unit}}, while the lowest value occurred under the lowest "
        "treatment level of Factor A paired with {{treatment_B2}} ({{min_val}} {{unit}}), with a low overall "
        "coefficient of variation (CV = {{cv_val}}%)."
    ),
    17: (
        "The terminal value of {{variable_name}} was significantly affected by both experimental "
        "factors, though their interactive effect was non-significant. Table {{table_num}} shows that the "
        "main effect of {{factor_A}} was significant (P \u2264 0.05), with {{treatment_A1}} producing the highest "
        "value ({{val_A1}} {{unit}}), which was statistically equivalent to {{treatment_A2}} ({{val_A2}} {{unit}}) "
        "but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsd_A}}). "
        "The lowest treatment level recorded the lowest value of {{val_min_A}} {{unit}}. Similarly, the main "
        "effect of {{factor_B}} was highly significant (P \u2264 0.01). Table {{table_num}} shows that "
        "{{treatment_B1}} produced a significantly higher value of {{val_B1}} {{unit}} compared to {{treatment_B2}} "
        "({{val_B2}} {{unit}}; LSD_0.05 = {{lsd_B}}). The interaction between the two factors (A \u00d7 B) "
        "was non-significant (P > 0.05), indicating that the two treatments acted independently on this parameter, "
        "with the overall grand mean of the trial recorded as {{grand_mean}} {{unit}} and a low coefficient "
        "of variation (CV = {{cv_val}}%)."
    ),
    18: (
        "The terminal value of {{variable_name}} was significantly affected by different levels of "
        "{{factor_A}}, while {{factor_B}} and their interaction had no significant effect. Table {{table_num}} "
        "shows that the main effect of {{factor_A}} was significant (P \u2264 0.05), with {{treatment_A1}} "
        "producing the highest value ({{val_A1}} {{unit}}), which was statistically equivalent to {{treatment_A2}} "
        "({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}; "
        "LSD_0.05 = {{lsd_A}}). In contrast, the main effect of {{factor_B}} was non-significant (P > 0.05), "
        "with values ranging closely from {{val_B1}} {{unit}} under {{treatment_B1}} to {{val_B2}} {{unit}} under "
        "{{treatment_B2}} (LSD_0.05 = {{lsd_B}}). Table {{table_num}} shows that the interaction between the two "
        "factors was also non-significant (P > 0.05), indicating that the choice of {{factor_A}} was the "
        "primary factor driving differences in {{variable_name}}, with a recorded grand mean of {{grand_mean}} "
        "{{unit}} and a standard error of the mean (SEm) of {{sem_val}}."
    ),
    19: (
        "The terminal value of {{variable_name}} was significantly affected by different levels of "
        "{{factor_B}}, while {{factor_A}} and their interaction had no significant effect. Table {{table_num}} "
        "shows that the main effect of {{factor_B}} was highly significant (P \u2264 0.01), with {{treatment_B1}} "
        "producing a significantly higher value of {{val_B1}} {{unit}} compared to {{treatment_B2}} ({{val_B2}} "
        "{{unit}}; LSD_0.05 = {{lsd_B}}). In contrast, the main effect of {{factor_A}} was non-significant (P > 0.05), "
        "with values ranging closely from {{val_min_A}} {{unit}} under the lowest treatment level to {{val_max_A}} "
        "{{unit}} under {{treatment_A1}} (LSD_0.05 = {{lsd_A}}). Table {{table_num}} shows that the interaction "
        "between the two factors was also non-significant (P > 0.05), indicating that the application of "
        "{{factor_B}} was the primary factor driving differences in {{variable_name}}, with a recorded grand mean "
        "of {{grand_mean}} {{unit}} and a standard error of the mean (SEm) of {{sem_val}}."
    ),
    20: (
        "The terminal value of {{variable_name}} was not significantly affected by {{factor_A}}, "
        "{{factor_B}}, or their interactive combinations. Table {{table_num}} shows that the values remained "
        "highly uniform across all treatments (P > 0.05), with a maximum of {{val_max_A}} recorded under "
        "{{treatment_A1}} and a minimum of {{val_min_A}} recorded under {{treatment_A2}}, which fell well "
        "within the non-significant range based on the post-hoc separation (LSD_0.05 = {{lsd_A}}). Similarly, "
        "for Factor B, both {{treatment_B1}} and {{treatment_B2}} produced statistically equivalent means of "
        "{{val_B}} (LSD_0.05 = {{lsd_B}}). The SEm for both factors was recorded as {{sem_val}}, and the "
        "overall grand mean of the experimental trial was {{grand_mean}} {{unit}}. This lack of statistical "
        "divergence, coupled with a remarkably low coefficient of variation (CV = {{cv_val}}%), confirms "
        "that {{variable_name}} maintained a stable baseline that was not altered by any of the treatment "
        "combinations."
    ),
    21: (
        "The terminal evaluation of the parameter index for {{variable_name}} showed high stability "
        "and low overall variance across all treatment combinations. Table {{table_num}} shows that the "
        "Coefficients of Variation (CV%) remained very low at {{cv_val}}%, and the Standard Error of the Mean "
        "(SEm) was recorded as {{sem_val}} {{unit}}, indicating a high degree of experimental precision and "
        "baseline consistency. Under these uniform conditions, neither the main effects of {{factor_A}} nor "
        "{{factor_B}} showed any statistically significant differences (P > 0.05). The mean values "
        "remained close to the grand mean of {{grand_mean}} {{unit}}, with all treatments sharing the same "
        "post-hoc letter grouping 'a'. The non-significant interactive effect (A \u00d7 B) further "
        "confirms that this parameter was robust to external treatments, remaining stable across both factor "
        "levels."
    ),
    22: (
        "The relative percentage changes in the terminal value of {{variable_name}} were evaluated "
        "to assess the magnitude of the treatment effects. Table {{table_num}} shows that the application of "
        "{{treatment_A1}} increased {{variable_name}} by {{pct_diff_A}}% compared to the lowest treatment level "
        "({{val_max_A}} vs. {{val_min_A}} {{unit}}). Other treatments, such as {{treatment_A2}} and "
        "{{treatment_A3}}, also increased the values by {{pct_diff_A2}}% and {{pct_diff_A3}}% over the "
        "lowest treatment level, respectively. For the {{factor_B}} treatments, {{treatment_B1}} resulted in "
        "a {{pct_diff_B}}% increase compared to {{treatment_B2}} ({{val_B1}} vs. {{val_B2}} {{unit}}). These "
        "differences were highly significant for both the {{factor_A}} main effect (P \u2264 0.01) and the "
        "{{factor_B}} main effect (P \u2264 0.001), with no significant interaction observed (P > 0.05)."
    ),
    23: (
        "The interaction effect between {{factor_A}} and {{factor_B}} (A \u00d7 B) was highly "
        "significant for the terminal value of {{variable_name}} (P \u2264 0.01). Table {{table_num}} shows "
        "that this significant interaction dominated the experimental response, making the main effects "
        "secondary in explaining the variation in {{variable_name}}. Under this significant interactive "
        "response, the combination of {{treatment_A1}} and {{treatment_B1}} produced the highest overall value "
        "of {{max_val}} {{unit}}, which was significantly higher than all other treatment combinations. "
        "Conversely, the combination of the lowest treatment level of Factor A and {{treatment_B2}} "
        "produced the lowest value of {{min_val}} {{unit}}. This strong interactive response shows that the "
        "positive effect of the treatment was highly dependent on the choice of {{factor_A}} levels, with a low "
        "overall coefficient of variation (CV = {{cv_val}}%)."
    ),
    24: (
        "The terminal value of {{variable_name}} was significantly affected by both experimental "
        "factors, though their interactive effect was non-significant. Table {{table_num}} shows that the "
        "main effect of {{factor_A}} was significant (P \u2264 0.05, LSD_0.05 = {{lsd_A}}), with {{treatment_A1}} "
        "producing the highest value ({{val_A1}} {{unit}}), which was statistically equivalent to {{treatment_A2}} "
        "({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level ({{val_lowest}} {{unit}}). "
        "The standard error of the mean (SEm) for the {{factor_A}} treatment was {{sem_A}}. Similarly, "
        "the main effect of {{factor_B}} was highly significant (P \u2264 0.01, LSD_0.05 = {{lsd_B}}). "
        "Table {{table_num}} shows that {{treatment_B1}} produced a significantly higher value of {{val_B1}} "
        "{{unit}} compared to {{treatment_B2}} ({{val_B2}} {{unit}}), with a recorded SEm of {{sem_B}}. "
        "The interaction between the two factors was non-significant (P > 0.05), with the overall grand mean "
        "of the trial recorded as {{grand_mean}} {{unit}} and a low coefficient of variation (CV = {{cv_val}}%)."
    ),
    25: (
        "The final ratio of the system components at terminal harvest was significantly affected "
        "by both experimental factors, though their interactive effect was non-significant. Table {{table_num}} "
        "shows that the main effect of {{factor_A}} was significant (P \u2264 0.05), with {{treatment_A1}} "
        "producing the highest fractional ratio ({{val_A1}} {{unit}}), which was statistically equivalent to "
        "{{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level "
        "({{val_lowest}} {{unit}}; LSD_0.05 = {{lsd_A}}). The main effect of {{factor_B}} was highly "
        "significant (P \u2264 0.01). Table {{table_num}} shows that {{treatment_B1}} produced a significantly "
        "higher ratio ({{val_B1}} {{unit}}) compared to {{treatment_B2}} ({{val_B2}} {{unit}}; LSD_0.05 = "
        "{{lsd_B}}). The interaction between {{factor_A}} and {{factor_B}} was non-significant (P > 0.05), "
        "indicating that the two factors affected the component ratio independently."
    ),
    26: (
        "The final activity rate of {{variable_name}} was significantly affected by both "
        "experimental factors, with no significant interaction observed. Table {{table_num}} shows "
        "that the main effect of {{factor_A}} was highly significant (P \u2264 0.01), with {{treatment_A1}} "
        "producing the highest activity rate ({{val_A1}} {{unit}}), which was statistically on par with "
        "{{treatment_A2}} ({{val_A2}} {{unit}}) but significantly higher than the lowest treatment level "
        "({{val_lowest}} {{unit}}; LSD_0.05 = {{lsd_A}}). Similarly, the main effect of {{factor_B}} "
        "was highly significant (P \u2264 0.01). Table {{table_num}} shows that {{treatment_B1}} produced "
        "a significantly higher activity rate ({{val_B1}} {{unit}}) compared to {{treatment_B2}} "
        "({{val_B2}} {{unit}}; LSD_0.05 = {{lsd_B}}). The interaction between the two factors was "
        "non-significant (P > 0.05), with the grand mean of the experiment recorded at {{grand_mean}} "
        "{{unit}} and a low coefficient of variation (CV = {{cv_val}}%)."
    ),
    27: (
        "The final concentration of {{variable_name}} in the experimental matrix was "
        "significantly affected by different levels of {{factor_A}}, while {{factor_B}} and their "
        "interaction had no significant effect. Table {{table_num}} shows that the main effect of {{factor_A}} "
        "was highly significant (P \u2264 0.01), with {{treatment_A1}} producing the highest concentration "
        "of {{val_A1}} {{unit}}, which was significantly higher than {{treatment_A2}} ({{val_A2}} {{unit}}) "
        "and the lowest treatment level ({{val_lowest}} {{unit}}; LSD_0.05 = {{lsd_A}}). In contrast, the "
        "main effect of {{factor_B}} was non-significant (P > 0.05), with values ranging closely from {{val_B1}} "
        "{{unit}} under {{treatment_B1}} to {{val_B2}} {{unit}} under {{treatment_B2}} (LSD_0.05 = {{lsd_B}}). "
        "Table {{table_num}} shows that the interaction between the two factors was also non-significant (P > 0.05), "
        "indicating that the choice of {{factor_A}} was the primary factor driving differences in matrix concentration."
    ),
    28: (
        "The terminal value of the primary output component, {{variable_name}}, was "
        "significantly affected by both experimental factors, with no significant interaction "
        "observed. Table {{table_num}} shows that the main effect of {{factor_A}} was significant "
        "(P \u2264 0.05), with {{treatment_A1}} producing the highest value ({{val_A1}}), which was statistically "
        "equivalent to {{treatment_A2}} ({{val_A2}}) but significantly higher than the lowest treatment "
        "level ({{val_lowest}}; LSD_0.05 = {{lsd_A}}). Similarly, the main effect of {{factor_B}} was highly "
        "significant (P \u2264 0.01). Table {{table_num}} shows that {{treatment_B1}} produced a significantly "
        "higher value of {{val_B1}} compared to {{treatment_B2}} ({{val_B2}}; LSD_0.05 = {{lsd_B}}). The "
        "interaction between the two factors was non-significant (P > 0.05), indicating that both factors "
        "contributed independently to the final output performance, with a recorded grand mean of {{grand_mean}} "
        "and a low coefficient of variation (CV = {{cv_val}}%)."
    ),
    29: (
        "The terminal evaluation of the physical resistance parameter, {{variable_name}}, "
        "showed no significant differences among treatment combinations. Table {{table_num}} shows "
        "that the values remained statistically uniform across all treatments (P > 0.05), with a maximum "
        "of {{val_max_A}} recorded under {{treatment_A1}} and a minimum of {{val_min_A}} recorded under "
        "{{treatment_A2}}, which fell well within the non-significant range based on the post-hoc separation "
        "(LSD_0.05 = {{lsd_A}}). Similarly, for Factor B, both {{treatment_B1}} and {{treatment_B2}} produced "
        "statistically equivalent means of {{val_B}} (LSD_0.05 = {{lsd_B}}). The SEm for both factors "
        "was recorded as {{sem_val}}, and the overall grand mean of the experimental trial was {{grand_mean}} "
        "{{unit}}. This lack of statistical divergence, coupled with a low coefficient of variation "
        "(CV = {{cv_val}}%), confirms that this structural parameter was robust to changes in both factor levels."
    ),
    30: (
        "The final proportion of elements classified in the premium grade, designated as "
        "{{variable_name}}, was significantly affected by both experimental factors, with no "
        "significant interaction observed. Table {{table_num}} shows that the main effect of {{factor_A}} "
        "was highly significant (P \u2264 0.01), with {{treatment_A1}} producing the highest proportion "
        "of {{val_A1}}%, which was statistically equivalent to {{treatment_A2}} ({{val_A2}}%) but "
        "significantly higher than the lowest treatment level ({{val_lowest}}%; LSD_0.05 = {{lsd_A}}). "
        "Similarly, the main effect of {{factor_B}} was highly significant (P \u2264 0.01). Table "
        "{{table_num}} shows that {{treatment_B1}} produced a significantly higher proportion of {{val_B1}}% "
        "compared to {{treatment_B2}} ({{val_B2}}%; LSD_0.05 = {{lsd_B}}). The interaction between the two "
        "factors was non-significant (P > 0.05), indicating that the two factors affected the "
        "proportional grade distribution independently, with a recorded grand mean of {{grand_mean}}% "
        "and a standard error of the mean (SEm) of {{sem_val}}."
    )
}

# --- Shuffling Categories Mapping ---
SINGLE_DAY_TEMPLATES = {
    "interaction": [16, 23],
    "both_significant": [17, 22, 24, 25, 26, 28, 30],
    "factor_a_significant": [18, 27],
    "factor_b_significant": [19],
    "non_significant": [20, 21, 29]
}

TIME_SERIES_TEMPLATES = {
    "interaction": [1, 3, 12, 14],
    "upward_trend": [7, 9, 10, 13],
    "downward_trend": [2, 6, 15],
    "general": [4, 5, 8, 11]
}


class TemplatePool:
    """Manages thread-safe dynamic draws of shuffled templates within structured categories."""
    def __init__(self, templates_dict):
        self.initial_pool = {k: list(v) for k, v in templates_dict.items()}
        self.pool = {k: list(v) for k, v in templates_dict.items()}
        self.shuffle_all()

    def shuffle_all(self):
        for k in self.pool:
            random.shuffle(self.pool[k])

    def get_template_id(self, category):
        if category not in self.pool or not self.pool[category]:
            self.pool[category] = list(self.initial_pool[category])
            random.shuffle(self.pool[category])
        return self.pool[category].pop(0)


# ==============================================================================
# Hierarchical numbering for the Word report (1 / 1.1 / 1.1.1) + table numbering
# ==============================================================================
MAJOR_CATEGORY_ORDER = [
    "Vegetative Parameters",
    "Reproductive Parameters",
    "Post-harvest Parameters",
    "Stress and Health Parameters",
]

SUB_CATEGORY_ORDER = {
    "Vegetative Parameters": [
        "Germination and Establishment",
        "Plant Morphology",
        "Leaf Parameters",
        "Branch Parameters",
        "Root Parameters",
        "Biomass Parameters",
        "Growth Analysis Parameters",
        "Physiological Parameters",
        "Phenological Parameters",
    ],
    "Reproductive Parameters": [
        "Flowering Parameters",
        "Pollination Parameters",
        "Fruit Set Parameters",
        "Fruit Growth Parameters",
        "Fruit Yield Parameters",
        "Grain Yield Parameters (Cereal Crops)",
        "Seed Parameters",
        "Maturity Parameters",
    ],
    "Post-harvest Parameters": [
        "Shelf-life Parameters",
        "Chemical Quality",
        "Physical Quality",
    ],
    "Stress and Health Parameters": [
        "Disease and Pest Parameters",
    ],
}


class ReportNumberer:
    """Generates continuous 1 / 1.1 / 1.1.1 style headings and continuous Table N numbers."""
    def __init__(self):
        self.major_num = 0
        self.sub_num = 0
        self.param_num = 0
        self.table_num = 1

    def major(self, doc, title):
        self.major_num += 1
        self.sub_num = 0
        self.param_num = 0
        full_title = f"{self.major_num} {title}"
        doc.add_heading(full_title, level=1)
        return full_title

    def sub(self, doc, title):
        self.sub_num += 1
        self.param_num = 0
        full_title = f"{self.major_num}.{self.sub_num} {title}"
        doc.add_heading(full_title, level=2)
        return full_title

    def param(self, doc, title):
        self.param_num += 1
        full_title = f"{self.major_num}.{self.sub_num}.{self.param_num} {title}"
        doc.add_heading(full_title, level=3)
        return full_title

    def next_table(self):
        n = self.table_num
        self.table_num += 1
        return n


# --- P-Value Academic Notation Mapper ---
def get_p_val_notation(p_val):
    if p_val < 0.001:
        return "p < 0.001"
    elif p_val < 0.01:
        return "p < 0.01"
    elif p_val < 0.05:
        return "p < 0.05"
    else:
        return "p > 0.05"


# --- Agronomic Classification Engine ---
def classify_parameter(param):
    """Classifies parameters dynamically based on standard agricultural taxonomy."""
    param_clean = param.strip().lower()

    abbrev_map = {
        "sl": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "dl": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "plw": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "plwd": ("Post-harvest Parameters", "Shelf-life Parameters"),
        "tss": ("Post-harvest Parameters", "Chemical Quality"),
        "ta": ("Post-harvest Parameters", "Chemical Quality"),
        "ph": ("Post-harvest Parameters", "Chemical Quality"),
        "lai": ("Vegetative Parameters", "Leaf Parameters"),
        "lad": ("Vegetative Parameters", "Leaf Parameters"),
        "sla": ("Vegetative Parameters", "Leaf Parameters"),
        "slw": ("Vegetative Parameters", "Leaf Parameters"),
        "cgr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "rgr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "agr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "nar": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "lar": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "lwr": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "rue": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "hi": ("Vegetative Parameters", "Growth Analysis Parameters"),
        "rwc": ("Vegetative Parameters", "Physiological Parameters"),
        "spad": ("Vegetative Parameters", "Leaf Parameters"),
        "ndvi": ("Vegetative Parameters", "Physiological Parameters"),
        "pri": ("Vegetative Parameters", "Physiological Parameters"),
    }

    base_abbrev = re.sub(r"\d+$", "", param_clean)
    if base_abbrev in abbrev_map:
        return abbrev_map[base_abbrev]

    if any(x in param_clean for x in ["disease", "pest", "infest", "weed", "mortality", "survival",
                                       "stress tolerance", "chlorosis", "wilting"]):
        return ("Stress and Health Parameters", "Disease and Pest Parameters")

    if any(x in param_clean for x in ["plw", "physiological loss", "decay", "shelf life", "shelf-life",
                                       "spoilage", "marketability", "firmness retention"]):
        return ("Post-harvest Parameters", "Shelf-life Parameters")

    if any(x in param_clean for x in ["tss", "soluble solid", "titratable", "acidity", "ta", "ph",
                                       "vitamin c", "ascorbic", "lycopene", "carotene", "anthocyanin",
                                       "phenolic", "flavonoid", "antioxidant", "sugar", "dry matter"]):
        return ("Post-harvest Parameters", "Chemical Quality")

    if any(x in param_clean for x in ["firmness", "peel thickness", "specific gravity"]) or \
       (any(x in param_clean for x in ["fruit size", "fruit shape", "fruit color"]) and "quality" in param_clean):
        return ("Post-harvest Parameters", "Physical Quality")

    if any(x in param_clean for x in ["flower", "flowering", "initiation", "cluster"]):
        return ("Reproductive Parameters", "Flowering Parameters")

    if any(x in param_clean for x in ["pollen", "pollination", "fertilization"]):
        return ("Reproductive Parameters", "Pollination Parameters")

    if any(x in param_clean for x in ["fruit set", "fruit retention", "fruit drop"]):
        return ("Reproductive Parameters", "Fruit Set Parameters")

    if any(x in param_clean for x in ["marketable", "unmarketable", "yield", "harvest index", "harvest"]):
        return ("Reproductive Parameters", "Fruit Yield Parameters")

    if any(x in param_clean for x in ["fruit length", "fruit diameter", "fruit circumference", "fruit volume",
                                       "fruit weight", "fruit shape", "fruit color", "locule", "pericarp",
                                       "maturity index"]):
        return ("Reproductive Parameters", "Fruit Growth Parameters")

    if any(x in param_clean for x in ["spike", "panicle", "ear", "grain", "straw", "filled grains"]):
        return ("Reproductive Parameters", "Grain Yield Parameters (Cereal Crops)")

    if any(x in param_clean for x in ["seed"]):
        return ("Reproductive Parameters", "Seed Parameters")

    if any(x in param_clean for x in ["maturity", "harvest duration"]):
        return ("Reproductive Parameters", "Maturity Parameters")

    if any(x in param_clean for x in ["germination", "emergence", "establishment", "vigor index"]):
        return ("Vegetative Parameters", "Germination and Establishment")

    if any(x in param_clean for x in ["leaf", "leaves", "lai", "lad", "sla", "slw", "chlorophyll",
                                       "carotenoid", "spad", "greenness"]):
        return ("Vegetative Parameters", "Leaf Parameters")

    if any(x in param_clean for x in ["branch"]):
        return ("Vegetative Parameters", "Branch Parameters")

    if any(x in param_clean for x in ["root"]):
        return ("Vegetative Parameters", "Root Parameters")

    if any(x in param_clean for x in ["fresh weight", "dry weight", "biomass"]):
        return ("Vegetative Parameters", "Biomass Parameters")

    if any(x in param_clean for x in ["cgr", "rgr", "agr", "nar", "lar", "lwr", "rue"]):
        return ("Vegetative Parameters", "Growth Analysis Parameters")

    if any(x in param_clean for x in ["photosynthetic", "transpiration", "conductance", "water use",
                                       "fluorescence", "relative water", "rwc", "membrane stability",
                                       "electrolyte", "canopy temp", "ndvi", "pri", "fv/fm"]):
        return ("Vegetative Parameters", "Physiological Parameters")

    if any(x in param_clean for x in ["phenological", "first leaf", "vegetative maturity"]):
        return ("Vegetative Parameters", "Phenological Parameters")

    if any(x in param_clean for x in ["height", "stem", "collar", "internode", "node", "canopy",
                                       "spread", "crown"]):
        return ("Vegetative Parameters", "Plant Morphology")

    return ("Vegetative Parameters", "Plant Morphology")


# --- Word Document Table Formatting Helpers ---
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '6')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        tblBorders.append(element)
    tblPr.append(tblBorders)


def set_header_bottom_border(row_or_cells):
    if hasattr(row_or_cells, 'cells'):
        cells = row_or_cells.cells
    else:
        cells = row_or_cells
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)


# --- Dynamic Table Caption Generator ---
def generate_table_caption(table_num, factor_a, factor_b, variables_list):
    vars_txt = ", ".join(variables_list)
    if len(variables_list) > 2:
        vars_txt = f"{', '.join(variables_list[:-1])} and {variables_list[-1]}"

    captions = [
        f"Table {table_num}. Influence of {factor_a} and {factor_b} on {vars_txt}.",
        f"Table {table_num}. Response of {vars_txt} to various treatments of {factor_a} and {factor_b}.",
        f"Table {table_num}. Effect of {factor_a} and {factor_b} on {vars_txt}.",
        f"Table {table_num}. Physiological and biochemical performance of {vars_txt} under different levels of {factor_a} and {factor_b}.",
    ]
    return captions[table_num % len(captions)]


# --- Parameter Grouping Engine for Time-Series ---
def group_parameters(params):
    pattern = re.compile(r"^(.*?)[\s_\-]*(\d+)\s*(dat|das|days?|d)?$", re.IGNORECASE)
    groups = {}
    for p in params:
        match = pattern.search(p.strip())
        if match and match.group(1).strip():
            base = match.group(1).strip()
            base = re.sub(r"[\s\-_()]+$", "", base).strip().capitalize()
            day_num = int(match.group(2))
            day_str = f"Day {day_num}"
            if not base:
                base = "Parameter"
            groups.setdefault(base, []).append((p, day_num, day_str))
        else:
            base = p.strip().capitalize()
            groups.setdefault(base, []).append((p, 0, ""))

    for base in groups:
        groups[base].sort(key=lambda x: x[1])
    return groups


# --- Statistical Separation Engines ---
def get_signif_code_val(p):
    if pd.isna(p):
        return "ns"
    if p < 0.01:
        return "**"
    elif p < 0.05:
        return "*"
    else:
        return "ns"


def parse_dmrt_value(val):
    if pd.isna(val):
        return "", ""
    val_str = str(val).strip()
    match = re.match(r"^([\d.\-]+)\s*([a-zA-Z\s]+)?$", val_str)
    if match:
        num = match.group(1)
        letters = match.group(2).strip() if match.group(2) else ""
        return num, letters
    return val_str, ""


def get_cld_letters(means_dict, lsd):
    sorted_means = sorted(means_dict.items(), key=lambda x: x[1], reverse=True)
    n = len(sorted_means)
    groups = []

    for i in range(n):
        for j in range(i, n):
            is_group = True
            for k1 in range(i, j + 1):
                for k2 in range(k1, j + 1):
                    if abs(sorted_means[k1][1] - sorted_means[k2][1]) > lsd:
                        is_group = False
                        break
                if not is_group:
                    break
            if is_group:
                groups.append(set(range(i, j + 1)))

    maximal_groups = []
    for g in groups:
        is_subset = False
        for other_g in groups:
            if g != other_g and g.issubset(other_g):
                is_subset = True
                break
        if not is_subset and g not in maximal_groups:
            maximal_groups.append(g)

    maximal_groups.sort(key=lambda g: min(g))

    treatment_letters = {tname: "" for tname, _ in sorted_means}
    alphabet = string.ascii_lowercase

    for letter_idx, g in enumerate(maximal_groups):
        letter = alphabet[letter_idx % len(alphabet)]
        for idx in g:
            t_name = sorted_means[idx][0]
            treatment_letters[t_name] += letter

    return treatment_letters


def run_anova_2factor_raw(df, block_col, factor_a_col, factor_b_col, param):
    df_temp = pd.DataFrame({
        'rep': df[block_col].astype(str),
        'factor_a': df[factor_a_col].astype(str),
        'factor_b': df[factor_b_col].astype(str),
        'response': pd.to_numeric(df[param], errors='coerce')
    }).dropna(subset=['response'])

    formula = "response ~ C(rep) + C(factor_a) * C(factor_b)"
    model = ols(formula, data=df_temp).fit()
    anova_table = sm.stats.anova_lm(model, typ=1)

    df_err = anova_table.loc["Residual", "df"]
    mse = anova_table.loc["Residual", "mean_sq"]

    p_a = anova_table.loc["C(factor_a)", "PR(>F)"]
    p_b = anova_table.loc["C(factor_b)", "PR(>F)"]
    p_ab = anova_table.loc["C(factor_a):C(factor_b)", "PR(>F)"]

    grand_mean = df_temp['response'].mean()
    cv = (np.sqrt(mse) / grand_mean) * 100
    t_val = t.ppf(0.975, df_err)

    r = df_temp['rep'].nunique()
    b_levels = df_temp['factor_b'].nunique()
    a_levels = df_temp['factor_a'].nunique()

    means_a = df_temp.groupby('factor_a')['response'].mean().to_dict()
    sem_a = np.sqrt(mse / (r * b_levels))
    lsd_a = t_val * np.sqrt((2 * mse) / (r * b_levels))
    cld_a = get_cld_letters(means_a, lsd_a)

    means_b = df_temp.groupby('factor_b')['response'].mean().to_dict()
    sem_b = np.sqrt(mse / (r * a_levels))
    lsd_b = t_val * np.sqrt((2 * mse) / (r * a_levels))
    cld_b = get_cld_letters(means_b, lsd_b)

    df_temp['Combination'] = df_temp['factor_a'] + " \u00d7 " + df_temp['factor_b']
    means_comb = df_temp.groupby('Combination')['response'].mean().to_dict()
    lsd_comb = t_val * np.sqrt((2 * mse) / r)
    cld_comb = get_cld_letters(means_comb, lsd_comb)

    sig_a = get_signif_code_val(p_a)
    sig_b = get_signif_code_val(p_b)
    sig_ab = get_signif_code_val(p_ab)

    means_a_str = {}
    for lvl, val in means_a.items():
        let = cld_a[lvl] if p_a < 0.05 else ""
        means_a_str[lvl] = f"{val:.2f}{let}"

    means_b_str = {}
    for lvl, val in means_b.items():
        let = cld_b[lvl] if p_b < 0.05 else ""
        means_b_str[lvl] = f"{val:.2f}{let}"

    return {
        "means_a": means_a, "means_a_str": means_a_str, "sem_a": round(sem_a, 2), "sig_a": sig_a,
        "lsd_a": round(lsd_a, 2), "p_a": p_a,
        "means_b": means_b, "means_b_str": means_b_str, "sem_b": round(sem_b, 2), "sig_b": sig_b,
        "lsd_b": round(lsd_b, 2), "p_b": p_b,
        "means_comb": means_comb, "cld_comb": cld_comb, "p_ab": p_ab, "sig_ab": sig_ab,
        "cv": round(cv, 2), "gm": round(grand_mean, 2)
    }


# --- Summarized Table Parser Engine ---
def parse_summarized_table_to_results_2f(df_raw, idx_A, idx_B, idx_cv, idx_interaction, idx_grand,
                                          idx_A_sem, idx_A_f, idx_A_lsd, idx_B_sem, idx_B_f, idx_B_lsd,
                                          factor_a_levels, factor_b_levels, parameters):
    results_data = {}
    for param in parameters:
        col_idx = df_raw.iloc[0].tolist().index(param)

        f_val_A = str(df_raw.iloc[idx_A_f, col_idx]).strip().lower()
        f_val_B = str(df_raw.iloc[idx_B_f, col_idx]).strip().lower()
        f_val_AB = str(df_raw.iloc[idx_interaction, col_idx]).strip().lower()

        p_a = 0.01 if "**" in f_val_A else (0.04 if "*" in f_val_A else 0.5)
        p_b = 0.01 if "**" in f_val_B else (0.04 if "*" in f_val_B else 0.5)
        p_ab = 0.01 if "**" in f_val_AB else (0.04 if "*" in f_val_AB else 0.5)

        means_a = {}
        means_a_str = {}
        for i, lvl in enumerate(factor_a_levels):
            num, let = parse_dmrt_value(df_raw.iloc[idx_A + 1 + i, col_idx])
            try:
                val = float(num)
                means_a[lvl] = val
                means_a_str[lvl] = f"{val:.2f}{let}"
            except ValueError:
                means_a[lvl] = 0.0
                means_a_str[lvl] = "0.00"

        means_b = {}
        means_b_str = {}
        for i, lvl in enumerate(factor_b_levels):
            num, let = parse_dmrt_value(df_raw.iloc[idx_B + 1 + i, col_idx])
            try:
                val = float(num)
                means_b[lvl] = val
                means_b_str[lvl] = f"{val:.2f}{let}"
            except ValueError:
                means_b[lvl] = 0.0
                means_b_str[lvl] = "0.00"

        gm_val = 0.0
        try:
            gm_val = float(df_raw.iloc[idx_grand, col_idx])
        except ValueError:
            pass

        means_comb = {f"{a} \u00d7 {b}": (means_a[a] + means_b[b]) / 2 for a in factor_a_levels for b in factor_b_levels}
        cld_comb = {k: "" for k in means_comb}

        results_data[param] = {
            "means_a": means_a, "means_a_str": means_a_str, "sem_a": df_raw.iloc[idx_A_sem, col_idx],
            "sig_a": get_signif_code_val(p_a), "lsd_a": df_raw.iloc[idx_A_lsd, col_idx], "p_a": p_a,
            "means_b": means_b, "means_b_str": means_b_str, "sem_b": df_raw.iloc[idx_B_sem, col_idx],
            "sig_b": get_signif_code_val(p_b), "lsd_b": df_raw.iloc[idx_B_lsd, col_idx], "p_b": p_b,
            "means_comb": means_comb, "cld_comb": cld_comb, "p_ab": p_ab, "sig_ab": get_signif_code_val(p_ab),
            "cv": df_raw.iloc[idx_cv, col_idx], "gm": gm_val
        }
    return results_data


# --- Academic Explanation Generators with Shuffled Pools ---
def generate_two_factor_explanation_shuffled(parameter, res, factor_a_col, factor_b_col, table_label, pool):
    p_a, p_b, p_ab = res["p_a"], res["p_b"], res["p_ab"]
    placeholders = extract_single_day_facts(parameter, res, factor_a_col, factor_b_col, table_label)

    # Route dynamically to shuffled category blocks
    if p_ab < 0.05:
        category = "interaction"
    elif p_a < 0.05 and p_b < 0.05:
        category = "both_significant"
    elif p_a < 0.05 and p_b >= 0.05:
        category = "factor_a_significant"
    elif p_a >= 0.05 and p_b < 0.05:
        category = "factor_b_significant"
    else:
        category = "non_significant"

    tpl_idx = pool.get_template_id(category)
    raw_template = ACADEMIC_TEMPLATES_30[tpl_idx]
    return inject_template_placeholders(raw_template, placeholders)


def generate_trend_explanation_2f_shuffled(base_name, items, results_data, factor_a_col, factor_b_col, table_label, pool):
    first_param, _, _ = items[0]
    last_param, _, _ = items[-1]

    first_gm = results_data[first_param]["gm"]
    last_gm = results_data[last_param]["gm"]
    p_ab_last = results_data[last_param]["p_ab"]
    direction_up = last_gm >= first_gm

    placeholders = extract_time_series_facts(base_name, items, results_data, factor_a_col, factor_b_col, table_label)

    if p_ab_last < 0.05:
        category = "interaction"
    elif direction_up:
        category = "upward_trend"
    else:
        category = "downward_trend"

    # Occasionally pull a general trend to increase formatting diversity
    if random.random() < 0.25:
        category = "general"

    tpl_idx = pool.get_template_id(category)
    raw_template = ACADEMIC_TEMPLATES_30[tpl_idx]
    return inject_template_placeholders(raw_template, placeholders)


# --- Styled Excel Exporter ---
def build_styled_excel(factor_a_col, factor_b_col, params, levels_a, levels_b, results_dict):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "2-Factor RCBD Output"
    ws.sheet_view.showGridLines = True

    font_bold = Font(name="Calibri", size=11, bold=True)
    font_regular = Font(name="Calibri", size=11)
    align_left = Alignment(horizontal="left", vertical="center")
    align_center = Alignment(horizontal="center", vertical="center")
    border_thin_bottom = Border(bottom=Side(style='thin', color='000000'))
    border_medium_bottom = Border(bottom=Side(style='medium', color='000000'))

    ws.cell(row=1, column=1, value="Treatments").font = font_bold
    ws.cell(row=1, column=1).alignment = align_left

    for col_idx, param in enumerate(params, start=2):
        cell = ws.cell(row=1, column=col_idx, value=param)
        cell.font = font_bold
        cell.alignment = align_center

    a_levels_num = len(levels_a)
    b_levels_num = len(levels_b)

    row_factor_a_title = 2
    row_factor_a_levels_start = 3
    row_factor_a_levels_end = 2 + a_levels_num
    row_sem_a = 3 + a_levels_num
    row_f_a = 4 + a_levels_num
    row_lsd_a = 5 + a_levels_num

    row_factor_b_title = 6 + a_levels_num
    row_factor_b_levels_start = 7 + a_levels_num
    row_factor_b_levels_end = 6 + a_levels_num + b_levels_num
    row_sem_b = 7 + a_levels_num + b_levels_num
    row_f_b = 8 + a_levels_num + b_levels_num
    row_lsd_b = 9 + a_levels_num + b_levels_num

    row_cv = 10 + a_levels_num + b_levels_num
    row_inter = 11 + a_levels_num + b_levels_num
    row_gm = 12 + a_levels_num + b_levels_num

    ws.cell(row=row_factor_a_title, column=1, value=f"Factor A: {factor_a_col}").font = font_bold
    ws.cell(row=row_factor_a_title, column=1).alignment = align_left

    for idx, lvl in enumerate(levels_a):
        r = row_factor_a_levels_start + idx
        ws.cell(row=r, column=1, value=lvl).font = font_regular
        ws.cell(row=r, column=1).alignment = align_left

    ws.cell(row=row_sem_a, column=1, value="Sem").font = font_regular
    ws.cell(row=row_sem_a, column=1).alignment = align_left
    ws.cell(row=row_f_a, column=1, value="p-value").font = font_regular
    ws.cell(row=row_f_a, column=1).alignment = align_left
    ws.cell(row=row_lsd_a, column=1, value="LSD(0.05)").font = font_regular
    ws.cell(row=row_lsd_a, column=1).alignment = align_left

    ws.cell(row=row_factor_b_title, column=1, value=f"Factor B: {factor_b_col}").font = font_bold
    ws.cell(row=row_factor_b_title, column=1).alignment = align_left

    for idx, lvl in enumerate(levels_b):
        r = row_factor_b_levels_start + idx
        ws.cell(row=r, column=1, value=lvl).font = font_regular
        ws.cell(row=r, column=1).alignment = align_left

    ws.cell(row=row_sem_b, column=1, value="Sem").font = font_regular
    ws.cell(row=row_sem_b, column=1).alignment = align_left
    ws.cell(row=row_f_b, column=1, value="p-value").font = font_regular
    ws.cell(row=row_f_b, column=1).alignment = align_left
    ws.cell(row=row_lsd_b, column=1, value="LSD(0.05)").font = font_regular
    ws.cell(row=row_lsd_b, column=1).alignment = align_left

    ws.cell(row=row_cv, column=1, value="CV(%)").font = font_regular
    ws.cell(row=row_cv, column=1).alignment = align_left
    ws.cell(row=row_inter, column=1, value="Factor A \u00d7 Factor B").font = font_regular
    ws.cell(row=row_inter, column=1).alignment = align_left
    ws.cell(row=row_gm, column=1, value="Grand Mean").font = font_regular
    ws.cell(row=row_gm, column=1).alignment = align_left

    for col_idx, param in enumerate(params, start=2):
        p_data = results_dict[param]

        for idx, lvl in enumerate(levels_a):
            r = row_factor_a_levels_start + idx
            cell = ws.cell(row=r, column=col_idx, value=p_data["means_a_str"][lvl])
            cell.font = font_regular
            cell.alignment = align_center

        ws.cell(row=row_sem_a, column=col_idx, value=p_data["sem_a"]).font = font_regular
        ws.cell(row=row_sem_a, column=col_idx).alignment = align_center
        ws.cell(row=row_f_a, column=col_idx, value=p_data["sig_a"]).font = font_regular
        ws.cell(row=row_f_a, column=col_idx).alignment = align_center
        ws.cell(row=row_lsd_a, column=col_idx, value=p_data["lsd_a"]).font = font_regular
        ws.cell(row=row_lsd_a, column=col_idx).alignment = align_center

        for idx, lvl in enumerate(levels_b):
            r = row_factor_b_levels_start + idx
            cell = ws.cell(row=r, column=col_idx, value=p_data["means_b_str"][lvl])
            cell.font = font_regular
            cell.alignment = align_center

        ws.cell(row=row_sem_b, column=col_idx, value=p_data["sem_b"]).font = font_regular
        ws.cell(row=row_sem_b, column=col_idx).alignment = align_center
        ws.cell(row=row_f_b, column=col_idx, value=p_data["sig_b"]).font = font_regular
        ws.cell(row=row_f_b, column=col_idx).alignment = align_center
        ws.cell(row=row_lsd_b, column=col_idx, value=p_data["lsd_b"]).font = font_regular
        ws.cell(row=row_lsd_b, column=col_idx).alignment = align_center

        ws.cell(row=row_cv, column=col_idx, value=p_data["cv"]).font = font_regular
        ws.cell(row=row_cv, column=col_idx).alignment = align_center
        ws.cell(row=row_inter, column=col_idx, value=p_data["sig_ab"]).font = font_regular
        ws.cell(row=row_inter, column=col_idx).alignment = align_center
        ws.cell(row=row_gm, column=col_idx, value=p_data["gm"]).font = font_regular
        ws.cell(row=row_gm, column=col_idx).alignment = align_center

    border_rows = [row_factor_a_levels_end, row_lsd_a, row_factor_b_levels_end, row_lsd_b, row_cv, row_inter, row_gm]
    for r_idx in border_rows:
        for col in range(1, len(params) + 2):
            ws.cell(row=r_idx, column=col).border = border_medium_bottom

    for col in range(1, len(params) + 2):
        ws.cell(row=1, column=col).border = border_thin_bottom

    return wb


# --- DOCX Copy of Styled Excel Table ---
def add_excel_table_to_docx(doc, factor_a_col, factor_b_col, g_cols, levels_a, levels_b, results_data):
    num_cols = len(g_cols) + 1
    table = doc.add_table(rows=1, cols=num_cols)
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Treatments"
    set_cell_margins(hdr_cells[0])
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
    hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(10)

    for c_idx, col_name in enumerate(g_cols):
        hdr_cells[c_idx + 1].text = col_name
        set_cell_margins(hdr_cells[c_idx + 1])
        hdr_cells[c_idx + 1].paragraphs[0].alignment = 1
        hdr_cells[c_idx + 1].paragraphs[0].runs[0].font.bold = True
        hdr_cells[c_idx + 1].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_cells[c_idx + 1].paragraphs[0].runs[0].font.size = Pt(10)

    set_header_bottom_border(table.rows[0])

    def add_styled_header_row(text):
        row_cells = table.add_row().cells
        row_cells[0].text = text
        set_cell_margins(row_cells[0])
        p = row_cells[0].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Arial'
        p.runs[0].font.size = Pt(10)
        for i in range(1, num_cols):
            row_cells[i].text = ""
            set_cell_margins(row_cells[i])
        return row_cells

    def add_styled_data_row(lbl, val_dict, bold=False):
        row_cells = table.add_row().cells
        row_cells[0].text = str(lbl)
        set_cell_margins(row_cells[0])
        p = row_cells[0].paragraphs[0]
        if len(p.runs) > 0:
            p.runs[0].font.bold = bold
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(10)

        for c_idx, col_name in enumerate(g_cols):
            row_cells[c_idx + 1].text = str(val_dict.get(col_name, "N/A"))
            set_cell_margins(row_cells[c_idx + 1])
            p_val = row_cells[c_idx + 1].paragraphs[0]
            p_val.alignment = 1
            if len(p_val.runs) > 0:
                p_val.runs[0].font.bold = bold
                p_val.runs[0].font.name = 'Arial'
                p_val.runs[0].font.size = Pt(10)
        return row_cells

    add_styled_header_row(f"Factor A: {factor_a_col}")
    for lvl in sorted(levels_a):
        lvl_dict = {p: results_data[p]["means_a_str"][lvl] for p in g_cols}
        add_styled_data_row(lvl, lvl_dict)

    sem_a_dict = {p: results_data[p]["sem_a"] for p in g_cols}
    sig_a_dict = {p: results_data[p]["sig_a"] for p in g_cols}
    lsd_a_dict = {p: results_data[p]["lsd_a"] for p in g_cols}

    add_styled_data_row("SEm(\u00b1)", sem_a_dict)
    add_styled_data_row("F-value", sig_a_dict)
    r_last_a = add_styled_data_row("LSD(0.05)", lsd_a_dict)
    set_header_bottom_border(r_last_a)

    add_styled_header_row(f"Factor B: {factor_b_col}")
    for lvl in sorted(levels_b):
        lvl_dict = {p: results_data[p]["means_b_str"][lvl] for p in g_cols}
        add_styled_data_row(lvl, lvl_dict)

    sem_b_dict = {p: results_data[p]["sem_b"] for p in g_cols}
    sig_b_dict = {p: results_data[p]["sig_b"] for p in g_cols}
    lsd_b_dict = {p: results_data[p]["lsd_b"] for p in g_cols}

    add_styled_data_row("SEm(\u00b1)", sem_b_dict)
    add_styled_data_row("F-value", sig_b_dict)
    r_last_b = add_styled_data_row("LSD(0.05)", lsd_b_dict)
    set_header_bottom_border(r_last_b)

    cv_dict = {p: results_data[p]["cv"] for p in g_cols}
    sig_ab_dict = {p: results_data[p]["sig_ab"] for p in g_cols}
    gm_dict = {p: results_data[p]["gm"] for p in g_cols}

    r_cv = add_styled_data_row("CV, %", cv_dict)
    set_header_bottom_border(r_cv)

    r_inter = add_styled_data_row("Factor A \u00d7 Factor B", sig_ab_dict)
    set_header_bottom_border(r_inter)

    r_gm = add_styled_data_row("Grand mean", gm_dict)
    set_header_bottom_border(r_gm)

    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        for idx in range(1, len(g_cols) + 1):
            row.cells[idx].width = Inches(1.1)


# ==============================================================================
# Word report builder: produces the 1 / 1.1 / 1.1.1 hierarchical structure
# ==============================================================================
def build_hierarchical_report(classified_cols, factor_a_col, factor_b_col, levels_a, levels_b, results_data):
    """
    Produces a Document with clustered parameters:
    Groups up to 4 static single parameters under the same category to produce
    cohesive narrative paragraphs and multi-column tables.
    Includes active template shuffling to prevent layout fatigue.
    """
    doc = Document()
    doc.add_heading("Calculated Two-Factor Factorial RCBD Report", 0)
    numberer = ReportNumberer()

    # Thread-safe fresh shuffling pools instantiated for this report build run
    single_day_pool = TemplatePool(SINGLE_DAY_TEMPLATES)
    time_series_pool = TemplatePool(TIME_SERIES_TEMPLATES)

    ordered_majors = [m for m in MAJOR_CATEGORY_ORDER if m in classified_cols]
    ordered_majors += [m for m in classified_cols if m not in ordered_majors]

    for major in ordered_majors:
        subs_dict = classified_cols[major]
        if not any(subs_dict.values()):
            continue

        major_title = numberer.major(doc, major)
        st.write(f"## {major_title}")

        sub_order = SUB_CATEGORY_ORDER.get(major, [])
        ordered_subs = [s for s in sub_order if s in subs_dict and subs_dict[s]]
        ordered_subs += [s for s in subs_dict if s not in ordered_subs and subs_dict[s]]

        for sub in ordered_subs:
            cat_params = subs_dict[sub]
            sub_title = numberer.sub(doc, sub)
            st.write(f"### {sub_title}")

            grouped = group_parameters(cat_params)

            # Separate static parameters vs time-series parameters
            static_items = [items[0][0] for base_name, items in sorted(grouped.items()) if len(items) == 1]
            time_series_items = {base_name: items for base_name, items in sorted(grouped.items()) if len(items) > 1}

            # 1. Clustered static parameters (grouped in chunks of up to 4)
            chunk_size = 4
            for i in range(0, len(static_items), chunk_size):
                chunk = static_items[i:i + chunk_size]

                chunk_title_str = ", ".join(chunk)
                if len(chunk) > 2:
                    chunk_title_str = f"{', '.join(chunk[:-1])} and {chunk[-1]}"

                param_title = numberer.param(doc, chunk_title_str)
                st.write(f"**{param_title}**")

                table_n = numberer.next_table()
                table_label = f"Table {table_n}"

                # Dynamic, shuffled sentence transition array
                transitions = [
                    " Concurrently, regarding the performance of {param}: ",
                    " In terms of {param}, the statistical analysis indicated that: ",
                    " Evaluated parallel to other traits, {param} showed that: ",
                    " Moving onto {param}, the results revealed: ",
                    " Likewise, for the parameter {param}, the data showed: ",
                    " Additionally, an examination of {param} indicated: "
                ]
                random.shuffle(transitions)

                explanations = []
                for idx, p in enumerate(chunk):
                    p_text = generate_two_factor_explanation_shuffled(
                        p, results_data[p], factor_a_col, factor_b_col, table_label, single_day_pool
                    )
                    
                    if idx > 0:
                        transition = transitions[idx % len(transitions)].format(param=p)
                        # Clean prefix repetitions gracefully
                        p_text = re.sub(
                            r"^(The terminal value of\s+|The terminal measurement of\s+|The final ratio of the system components at terminal harvest was\s+)", 
                            "", p_text, flags=re.IGNORECASE
                        )
                        p_text = transition + p_text[0].lower() + p_text[1:]
                    explanations.append(p_text)

                combined_narrative = " ".join(explanations)
                st.write(combined_narrative)
                doc.add_paragraph(combined_narrative)

                # Render consolidated multi-column table
                add_excel_table_to_docx(doc, factor_a_col, factor_b_col, chunk, levels_a, levels_b, results_data)

                # Caption
                caption_text = generate_table_caption(table_n, factor_a_col, factor_b_col, chunk)
                p_cap = doc.add_paragraph(caption_text)
                p_cap.runs[0].font.name = 'Arial'
                p_cap.runs[0].font.size = Pt(10)
                p_cap.runs[0].font.italic = True

                st.write(f"*{table_label} rendered below*")
                doc.add_paragraph()

            # 2. Time-series parameter groups (chronological progression, up to 4 intervals)
            for base_name, items in sorted(time_series_items.items()):
                param_title = numberer.param(doc, base_name)
                st.write(f"**{param_title}**")

                table_n = numberer.next_table()
                table_label = f"Table {table_n}"
                trend_params = [it[0] for it in items]

                p_text = generate_trend_explanation_2f_shuffled(
                    base_name, items, results_data, factor_a_col, factor_b_col, table_label, time_series_pool
                )
                st.write(p_text)
                doc.add_paragraph(p_text)

                add_excel_table_to_docx(doc, factor_a_col, factor_b_col, trend_params,
                                         levels_a, levels_b, results_data)

                caption_text = generate_table_caption(table_n, factor_a_col, factor_b_col, trend_params)
                p_cap = doc.add_paragraph(caption_text)
                p_cap.runs[0].font.name = 'Arial'
                p_cap.runs[0].font.size = Pt(10)
                p_cap.runs[0].font.italic = True

                st.write(f"*{table_label} (time-series) rendered below*")
                doc.add_paragraph()

        doc.add_paragraph("-" * 60)

    return doc


def build_classified_cols(parameters):
    classified_cols = {}
    for c in parameters:
        major, sub = classify_parameter(c)
        classified_cols.setdefault(major, {}).setdefault(sub, []).append(c)
    return classified_cols


def show_category_preview(classified_cols):
    st.write("#### Automatically Categorized Parameter Divisions:")
    ordered_majors = [m for m in MAJOR_CATEGORY_ORDER if m in classified_cols]
    ordered_majors += [m for m in classified_cols if m not in ordered_majors]
    for major in ordered_majors:
        subs_dict = classified_cols[major]
        sub_order = SUB_CATEGORY_ORDER.get(major, [])
        ordered_subs = [s for s in sub_order if s in subs_dict and subs_dict[s]]
        ordered_subs += [s for s in subs_dict if s not in ordered_subs and subs_dict[s]]
        for sub in ordered_subs:
            params = subs_dict[sub]
            if params:
                st.write(f"**{major} \u2192 {sub}:** {', '.join(params)}")


# --- Web Interface Routing and Multi-Year controller ---
def show_module():
    st.markdown("### Two-Factor RCBD Analyzer")

    mode = st.radio("Choose Input Mode", ["Raw Data Mode", "Summarized Table Mode"], key="2f_mode_selector_two")
    uploaded_file = st.file_uploader("Upload Two-Factor Excel File", type=["xlsx"], key="file_uploader_2f_two")

    if uploaded_file is not None:
        if mode == "Raw Data Mode":
            run_raw_mode(uploaded_file)
        else:
            run_summary_mode_processing(uploaded_file)


def run_raw_mode(uploaded_file):
    try:
        df_raw_data = pd.read_excel(uploaded_file)
        st.write("#### Preview Raw Factorial Input Data:", df_raw_data.head())

        cols = df_raw_data.columns.tolist()

        block_col = st.selectbox("Select Block/Replication Column", cols, index=0, key="raw_2y_bk_f_mod")
        factor_a_col = st.selectbox("Select Factor A Column", cols, index=1, key="raw_2f_fa_mod")
        factor_b_col = st.selectbox("Select Factor B Column", cols, index=2, key="raw_2f_fb_mod")
        response_cols = st.multiselect("Select Response Parameters to Analyze", cols, default=cols[3:], key="raw_2f_resp_mod")

        if response_cols:
            df_raw_data[factor_a_col] = df_raw_data[factor_a_col].astype(str)
            df_raw_data[factor_b_col] = df_raw_data[factor_b_col].astype(str)
            df_raw_data[block_col] = df_raw_data[block_col].astype(str)

            classified_cols = build_classified_cols(response_cols)
            show_category_preview(classified_cols)

            if st.button("Run Two-Factor Raw Analysis", key="btn_raw_2f_calc_m"):
                results_data = {}

                levels_a = sorted(df_raw_data[factor_a_col].unique().tolist())
                levels_b = sorted(df_raw_data[factor_b_col].unique().tolist())

                for param in response_cols:
                    results_data[param] = run_anova_2factor_raw(df_raw_data, block_col, factor_a_col, factor_b_col, param)

                styled_wb = build_styled_excel(factor_a_col, factor_b_col, response_cols, levels_a, levels_b, results_data)
                excel_bio = io.BytesIO()
                styled_wb.save(excel_bio)
                excel_bio.seek(0)

                st.markdown("#### \U0001F4E5 Download Formatted Statistical Excel Results")
                st.download_button(
                    label="Download Formatted Excel Results Table",
                    data=excel_bio,
                    file_name="Result_2Factor_Output.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="btn_d_excel_styled"
                )
                st.write("---")

                st.markdown("### \U0001F4DD Analysis Results and Academic Explanations")

                doc = build_hierarchical_report(classified_cols, factor_a_col, factor_b_col, levels_a, levels_b, results_data)

                bio_doc = io.BytesIO()
                doc.save(bio_doc)
                bio_doc.seek(0)

                st.write("---")
                st.markdown("#### \U0001F4BE Save Explanations as a Word Report")
                st.download_button(
                    "Download Word Explanations Report (.docx)",
                    data=bio_doc,
                    file_name="Calculated_Factorial_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="btn_d_2f_raw_cal"
                )
    except Exception as e:
        st.error(f"Error executing raw combined Two-Factor analysis: {e}")


def run_summary_mode_processing(uploaded_file):
    try:
        df_raw = pd.read_excel(uploaded_file, header=None)

        idx_A, idx_B, idx_cv, idx_interaction, idx_grand = None, None, None, None, None
        for idx, val in enumerate(df_raw[0]):
            if pd.isna(val):
                continue
            val_str = str(val).strip().lower()
            if "factor a" in val_str:
                idx_A = idx
            elif "factor b" in val_str:
                idx_B = idx
            elif "cv" in val_str:
                idx_cv = idx
            elif any(x in val_str for x in ["factor a \u00d7 factor b", "factor a*factor b", "factor a x factor b", "interaction"]):
                idx_interaction = idx
            elif "grand mean" in val_str or "grandmean" in val_str:
                idx_grand = idx

        if any(v is None for v in [idx_A, idx_B, idx_cv, idx_interaction, idx_grand]):
            st.error("Missing structural markers (Factor A, Factor B, CV, Interaction, Grand Mean) in Column A.")
            return

        idx_A_sem, idx_A_f, idx_A_lsd = None, None, None
        for idx in range(idx_A + 1, idx_B):
            val = str(df_raw.iloc[idx, 0]).strip().lower()
            if "sem" in val:
                idx_A_sem = idx
            elif "f-value" in val or "f value" in val:
                idx_A_f = idx
            elif "lsd" in val:
                idx_A_lsd = idx

        idx_B_sem, idx_B_f, idx_B_lsd = None, None, None
        for idx in range(idx_B + 1, idx_cv):
            val = str(df_raw.iloc[idx, 0]).strip().lower()
            if "sem" in val:
                idx_B_sem = idx
            elif "f-value" in val or "f value" in val:
                idx_B_f = idx
            elif "lsd" in val:
                idx_B_lsd = idx

        factor_a_levels = [str(df_raw.iloc[i, 0]).strip() for i in range(idx_A + 1, idx_A_sem)]
        factor_b_levels = [str(df_raw.iloc[i, 0]).strip() for i in range(idx_B + 1, idx_B_sem)]

        factor_a_label = str(df_raw.iloc[idx_A, 0]).split(":")[-1].replace("(", "").replace(")", "").strip()
        factor_b_label = str(df_raw.iloc[idx_B, 0]).split(":")[-1].replace("(", "").replace(")", "").strip()

        parameters = [str(x).strip() for x in df_raw.iloc[0].tolist()[1:] if pd.notna(x)]
        st.success(f"Detected Factor A: {factor_a_label} | Factor B: {factor_b_label}")

        classified_cols = build_classified_cols(parameters)
        show_category_preview(classified_cols)

        if st.button("Generate Word Document Draft from Direct Output", key="btn_2f_sum_gen"):
            results_data = parse_summarized_table_to_results_2f(
                df_raw, idx_A, idx_B, idx_cv, idx_interaction, idx_grand,
                idx_A_sem, idx_A_f, idx_A_lsd, idx_B_sem, idx_B_f, idx_B_lsd,
                factor_a_levels, factor_b_levels, parameters
            )

            doc = build_hierarchical_report(classified_cols, factor_a_label, factor_b_label,
                                             factor_a_levels, factor_b_levels, results_data)

            bio_doc = io.BytesIO()
            doc.save(bio_doc)
            bio_doc.seek(0)

            st.write("---")
            st.markdown("#### \U0001F4BE Save Explanations & Tables as Word Document")
            st.download_button(
                "Download Word Explanations Report (.docx)",
                data=bio_doc,
                file_name="MultiYear_Summarized_Thesis_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_d_2f_sum_doc"
            )
    except Exception as e:
        st.error(f"Error parsing direct result summary table: {e}")


if __name__ == "__main__":
    show_module()
