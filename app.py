import streamlit as st
import pandas as pd
import re
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

st.set_page_config(page_title="AgriStats Report Generator", layout="wide")

st.title("🌾 AgriStats Report Generator")
st.subheader("Convert DMRT & RCBD Excel Tables into Q1 Journal-Ready Word Documents")

# --- Helper Functions for Word Document Table Formatting (APA Style) ---
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    # Apply standard APA borders (Top, Bottom, Header-bottom)
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def set_header_bottom_border(row):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            '<w:tcBorders %s>'
            '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(borders)

# --- Parsing Logic ---
def parse_dmrt_value(val):
    if pd.isna(val):
        return "", ""
    val_str = str(val).strip()
    # Match decimal or integer numbers followed by lowercase letters
    match = re.match(r"^([\d\.\-]+)\s*([a-z]+)?$", val_str)
    if match:
        num = match.group(1)
        letters = match.group(2) if match.group(2) else ""
        return num, letters
    return val_str, ""

# --- Upload Interface ---
uploaded_file = st.file_uploader("Upload your Result_Final.xlsx file", type=["xlsx"])

if uploaded_file is not None:
    try:
        # Load excel without headers initially to find the structural anchor
        xls = pd.ExcelFile(uploaded_file)
        df_raw = pd.read_excel(uploaded_file, header=None)
        
        # 1. Detect the anchor header row in Column 0
        anchor_row_idx = None
        anchor_keywords = ["genotype", "treatment", "variety", "fertilizer", "sowing", "pesticide", "spacing", "cultivar"]
        
        for idx, val in enumerate(df_raw[0]):
            if pd.notna(val) and str(val).strip().lower() in anchor_keywords:
                anchor_row_idx = idx
                break
                
        if anchor_row_idx is None:
            st.error("Could not dynamically find the Treatment/Genotype column. Please ensure Column A has an identifier in the header (e.g., 'Genotype' or 'Treatment').")
        else:
            # Re-read the dataframe starting from the detected header row
            df_cleaned = pd.read_excel(uploaded_file, skiprows=anchor_row_idx)
            
            # Clean column names
            df_cleaned.columns = [str(c).strip() for c in df_cleaned.columns]
            treatment_col = df_cleaned.columns[0]
            
            # Determine singular and plural labels dynamically
            treatment_label = treatment_col.lower()
            if treatment_label.endswith('y'):
                treatment_label_plural = treatment_label[:-1] + "ies"
            elif treatment_label.endswith('s'):
                treatment_label_plural = treatment_label
            else:
                treatment_label_plural = treatment_label + "s"
                
            st.success(f"Successfully detected '{treatment_col}' as the independent variable. Using '{treatment_label_plural}' for text generation.")
            
            # Find row indices of summary statistics
            stats_keywords = ["sem", "p-value", "lsd", "cv", "grand mean"]
            stats_rows = {}
            treatment_end_idx = len(df_cleaned)
            
            for idx, val in enumerate(df_cleaned[treatment_col]):
                if pd.notna(val):
                    val_clean = str(val).strip().lower()
                    for key in stats_keywords:
                        if key in val_clean:
                            stats_rows[key] = idx
                            # The treatments end where the first summary statistic begins
                            treatment_end_idx = min(treatment_end_idx, idx)
            
            # Separate the dataframe into Treatments and Summary Statistics
            df_treatments = df_cleaned.iloc[:treatment_end_idx].copy()
            
            # Build stats dictionary dynamically
            stats_data = {}
            for key, row_idx in stats_rows.items():
                stats_data[key] = df_cleaned.iloc[row_idx]
                
            # Render the treatments for preview
            st.write("### Data Preview (Treatments):", df_treatments)
            
            # Let user split columns into groups (Vegetative vs. Reproductive/Yield)
            all_parameters = [col for col in df_cleaned.columns if col != treatment_col]
            
            st.write("### Table Splitting (Optional)")
            col1, col2 = st.columns(2)
            with col1:
                group_1_name = st.text_input("First Table Title", "Vegetative and Morphological Traits")
                group_1_cols = st.multiselect("Select parameters for Table 1", all_parameters, default=all_parameters[:len(all_parameters)//2])
            with col2:
                group_2_name = st.text_input("Second Table Title", "Yield and Structural Traits")
                group_2_cols = st.multiselect("Select parameters for Table 2", [c for c in all_parameters if c not in group_1_cols], default=[c for c in all_parameters if c not in group_1_cols])
            
            # Generate Report Button
            if st.button("Generate Word Document Draft"):
                # Define our 20 unique template strings
                templates = [
                    "For **{parameter}**, the treatment effect was {significance} {p_value}, with {top_geno} registering the maximum value of {top_val}^{top_let}, which did not differ statistically from {other_geno} ({other_val}^{other_let}).",
                    "Regarding **{parameter}**, genotypic variations were {significance} {p_value}, led by {top_geno} ({top_val}^{top_let}) and trailed by {low_geno} ({low_val}^{low_let}).",
                    "The F-test for **{parameter}** proved to be {significance} {p_value}, highlighting {top_geno} ({top_val}^{top_let}) as the most vigorous performer, whereas {low_geno} ({low_val}^{low_let}) recorded the poorest response.",
                    "The hierarchy of performance for **{parameter}** was {significance} {p_value}; {top_geno} ({top_val}^{top_let}) occupied the statistical apex, while {low_geno} ({low_val}^{low_let}) was situated at the baseline.",
                    "Data on **{parameter}** exhibited a {significance} {p_value} trend, with {top_geno} ({top_val}^{top_let}) representing the highest magnitude and {low_geno} ({low_val}^{low_let}) marking the lowest.",
                    "With respect to **{parameter}**, the evaluated germplasm showed a {significance} {p_value} variance, partitioning into distinct clusters where {top_geno} ({top_val}^{top_let}) and {other_geno} ({other_val}^{other_let}) shared statistical parity at the top.",
                    "A distinct, {significance} {p_value} separation among treatments was noticed for **{parameter}**, where {top_geno} ({top_val}^{top_let}) and {other_geno} ({other_val}^{other_let}) fell within the same premier DMRT group.",
                    "DMRT grouping for **{parameter}** revealed a {significance} {p_value} distribution, assigning the supreme letter '{top_let}' to {top_geno} ({top_val}^{top_let}).",
                    "Statistical evaluation of **{parameter}** indicated a {significance} {p_value} treatment response, with {top_geno} ({top_val}^{top_let}) sharing statistical letters with {other_geno} ({other_val}^{other_let}).",
                    "The variation expressed in **{parameter}** was found to be {significance} {p_value}, placing {top_geno} ({top_val}^{top_let}) and {other_geno} ({other_val}^{other_let}) in the highest-performing statistical tier.",
                    "The contrast among {treatment_plural} for **{parameter}** was {significance} {p_value}, with {top_geno} ({top_val}^{top_let}) exhibiting a distinct statistical advantage over {low_geno} ({low_val}^{low_let}).",
                    "A {significance} {p_value} difference was observed in **{parameter}**, where {top_geno} ({top_val}^{top_let}) surpassed the grand mean of {grand_mean}, while {low_geno} ({low_val}^{low_let}) remained significantly below it.",
                    "The evaluated {treatment_plural} differed to a {significance} {p_value} degree for **{parameter}**, with {top_geno} ({top_val}^{top_let}) demonstrating clear superiority over {low_geno} ({low_val}^{low_let}).",
                    "In terms of **{parameter}**, the differences were {significance} {p_value}, establishing {top_geno} ({top_val}^{top_let}) as the premier {treatment_singular}, statistically outclassing {low_geno} ({low_val}^{low_let}).",
                    "The performance profile for **{parameter}** was {significance} {p_value}, as {top_geno} ({top_val}^{top_let}) outperformed {low_geno} ({low_val}^{low_let}) by a statistically significant margin.",
                    "Analysis of variance for **{parameter}** yielded a {significance} {p_value} result, where the highest value of {top_val}^{top_let} was achieved by {top_geno}.",
                    "The influence of treatment on **{parameter}** was {significance} {p_value}; the maximum manifestation was observed in {top_geno} ({top_val}^{top_let}), while {low_geno} ({low_val}^{low_let}) recorded the minimum.",
                    "The evaluated {treatment_plural} exhibited {significance} {p_value} discrepancies in **{parameter}**, with {top_geno} ({top_val}^{top_let}) showing the most favorable response.",
                    "Statistical analysis confirmed that the genotypic effect on **{parameter}** was {significance} {p_value}, with {top_geno} ({top_val}^{top_let}) emerging as the outstanding treatment.",
                    "The response pattern of {treatment_plural} for **{parameter}** was {significance} {p_value}; {top_geno} ({top_val}^{top_let}) led the rankings, followed by {other_geno} ({other_val}^{other_let})."
                ]
                
                doc = Document()
                
                # Title and introduction
                doc.add_heading("Results and Discussion Draft", 0)
                doc.add_paragraph("This document contains dynamically generated results formatted for high-tier agricultural publication submission.")
                
                # Loop through both table groups
                groups = [(group_1_name, group_1_cols, 1), (group_2_name, group_2_cols, 2)]
                
                for g_title, g_cols, table_num in groups:
                    if not g_cols:
                        continue
                        
                    doc.add_heading(f"Table {table_num}: {g_title}", level=1)
                    
                    # Generate the text explanation for this group
                    text_paragraphs = []
                    
                    for idx, param in enumerate(g_cols):
                        # Extract data for this parameter
                        p_val_raw = str(stats_data.get("p-value", {}).get(param, "0.05"))
                        
                        # Determine Significance Text & P-value notation
                        sig_type = "significant"
                        p_notation = "(p < 0.05)"
                        
                        if "NS" in p_val_raw.upper():
                            sig_type = "nonsignificant"
                            p_notation = "(p > 0.05)"
                        elif "***" in p_val_raw or "0.000" in p_val_raw:
                            sig_type = "highly significant"
                            p_notation = "(p < 0.001)"
                        elif "**" in p_val_raw:
                            sig_type = "significant"
                            p_notation = "(p < 0.01)"
                        elif "*" in p_val_raw:
                            sig_type = "significant"
                            p_notation = "(p < 0.05)"
                            
                        # Extract Means, Letters, and standard stats
                        param_series = df_treatments[param]
                        parsed_pairs = [parse_dmrt_value(x) for x in param_series]
                        
                        # Sort to find Top, Intermediate, and Bottom Genotypes
                        # Clean values for numeric sorting
                        numeric_vals = []
                        for i, (num_str, let) in enumerate(parsed_pairs):
                            try:
                                numeric_vals.append((float(num_str), let, df_treatments.iloc[i][treatment_col]))
                            except ValueError:
                                pass
                                
                        if not numeric_vals:
                            continue
                            
                        numeric_vals.sort(reverse=True, key=lambda x: x[0])
                        
                        top_val, top_let, top_geno = numeric_vals[0]
                        low_val, low_let, low_geno = numeric_vals[-1]
                        
                        # Find secondary/runner up genotype
                        other_val, other_let, other_geno = numeric_vals[1] if len(numeric_vals) > 1 else (top_val, top_let, top_geno)
                        
                        grand_mean_val = str(stats_data.get("grand mean", {}).get(param, "N/A"))
                        
                        # Select a template dynamically to diversify descriptions
                        template_index = idx % len(templates)
                        template = templates[template_index]
                        
                        desc = template.format(
                            parameter=param,
                            significance=sig_type,
                            p_value=p_notation,
                            top_geno=top_geno,
                            top_val=f"{top_val:.2f}",
                            top_let=top_let,
                            other_geno=other_geno,
                            other_val=f"{other_val:.2f}",
                            other_let=other_let,
                            low_geno=low_geno,
                            low_val=f"{low_val:.2f}",
                            low_let=low_let,
                            grand_mean=grand_mean_val,
                            treatment_singular=treatment_label,
                            treatment_plural=treatment_label_plural
                        )
                        text_paragraphs.append(desc)
                        
                    # Write explanations into document
                    for para in text_paragraphs:
                        # Simple markdown-like formatter to support Bold text in word
                        p = doc.add_paragraph()
                        parts = re.split(r'(\*\*.*?\*\*)', para)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                p.add_run(part[2:-2]).bold = True
                            else:
                                p.add_run(part)
                                
                    # Write formatted Table into Word (APA Style)
                    doc.add_paragraph(f"Table {table_num}. Mean comparison of {g_title.lower()} among evaluated {treatment_label_plural}.")
                    
                    table = doc.add_table(rows=1, cols=len(g_cols) + 1)
                    set_table_borders(table)
                    
                    # Fill Header Row
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = str(treatment_col)
                    set_cell_margins(hdr_cells[0])
                    for c_idx, col_name in enumerate(g_cols):
                        hdr_cells[c_idx + 1].text = col_name
                        set_cell_margins(hdr_cells[c_idx + 1])
                    set_header_bottom_border(table.rows[0])
                    
                    # Fill Treatment Rows
                    for r_idx, t_row in df_treatments.iterrows():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(t_row[treatment_col])
                        set_cell_margins(row_cells[0])
                        for c_idx, col_name in enumerate(g_cols):
                            row_cells[c_idx + 1].text = str(t_row[col_name])
                            set_cell_margins(row_cells[c_idx + 1])
                            
                    # Add stats rows (SEM, LSD, CV, Grand Mean)
                    for stat_key in ["sem", "lsd", "cv", "grand mean"]:
                        if stat_key in stats_data:
                            row_cells = table.add_row().cells
                            # Format label
                            label_map = {"sem": "SEm (±)", "lsd": "LSD (0.05)", "cv": "CV (%)", "grand mean": "Grand Mean"}
                            row_cells[0].text = label_map.get(stat_key, stat_key.title())
                            set_cell_margins(row_cells[0])
                            for c_idx, col_name in enumerate(g_cols):
                                row_cells[c_idx + 1].text = str(stats_data[stat_key].get(col_name, "N/A"))
                                set_cell_margins(row_cells[c_idx + 1])
                                
                    doc.add_page_break()
                
                # Save Document to stream
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                
                st.write("🎉 **Your Draft Report is ready!**")
                st.download_button(
                    label="Download Draft Report (.docx)",
                    data=bio,
                    file_name="AgriStats_Draft_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
    except Exception as e:
        st.error(f"Error parsing file: {e}")import streamlit as st
import pandas as pd
import re
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

st.set_page_config(page_title="AgriStats Report Generator", layout="wide")

st.title("🌾 AgriStats Report Generator")
st.subheader("Convert DMRT & RCBD Excel Tables into Q1 Journal-Ready Word Documents")

# --- Helper Functions for Word Document Table Formatting (APA Style) ---
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    # Apply standard APA borders (Top, Bottom, Header-bottom)
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def set_header_bottom_border(row):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            '<w:tcBorders %s>'
            '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(borders)

# --- Parsing Logic ---
def parse_dmrt_value(val):
    if pd.isna(val):
        return "", ""
    val_str = str(val).strip()
    # Match decimal or integer numbers followed by lowercase letters
    match = re.match(r"^([\d\.\-]+)\s*([a-z]+)?$", val_str)
    if match:
        num = match.group(1)
        letters = match.group(2) if match.group(2) else ""
        return num, letters
    return val_str, ""

# --- Upload Interface ---
uploaded_file = st.file_uploader("Upload your Result_Final.xlsx file", type=["xlsx"])

if uploaded_file is not None:
    try:
        # Load excel without headers initially to find the structural anchor
        xls = pd.ExcelFile(uploaded_file)
        df_raw = pd.read_excel(uploaded_file, header=None)
        
        # 1. Detect the anchor header row in Column 0
        anchor_row_idx = None
        anchor_keywords = ["genotype", "treatment", "variety", "fertilizer", "sowing", "pesticide", "spacing", "cultivar"]
        
        for idx, val in enumerate(df_raw[0]):
            if pd.notna(val) and str(val).strip().lower() in anchor_keywords:
                anchor_row_idx = idx
                break
                
        if anchor_row_idx is None:
            st.error("Could not dynamically find the Treatment/Genotype column. Please ensure Column A has an identifier in the header (e.g., 'Genotype' or 'Treatment').")
        else:
            # Re-read the dataframe starting from the detected header row
            df_cleaned = pd.read_excel(uploaded_file, skiprows=anchor_row_idx)
            
            # Clean column names
            df_cleaned.columns = [str(c).strip() for c in df_cleaned.columns]
            treatment_col = df_cleaned.columns[0]
            
            # Determine singular and plural labels dynamically
            treatment_label = treatment_col.lower()
            if treatment_label.endswith('y'):
                treatment_label_plural = treatment_label[:-1] + "ies"
            elif treatment_label.endswith('s'):
                treatment_label_plural = treatment_label
            else:
                treatment_label_plural = treatment_label + "s"
                
            st.success(f"Successfully detected '{treatment_col}' as the independent variable. Using '{treatment_label_plural}' for text generation.")
            
            # Find row indices of summary statistics
            stats_keywords = ["sem", "p-value", "lsd", "cv", "grand mean"]
            stats_rows = {}
            treatment_end_idx = len(df_cleaned)
            
            for idx, val in enumerate(df_cleaned[treatment_col]):
                if pd.notna(val):
                    val_clean = str(val).strip().lower()
                    for key in stats_keywords:
                        if key in val_clean:
                            stats_rows[key] = idx
                            # The treatments end where the first summary statistic begins
                            treatment_end_idx = min(treatment_end_idx, idx)
            
            # Separate the dataframe into Treatments and Summary Statistics
            df_treatments = df_cleaned.iloc[:treatment_end_idx].copy()
            
            # Build stats dictionary dynamically
            stats_data = {}
            for key, row_idx in stats_rows.items():
                stats_data[key] = df_cleaned.iloc[row_idx]
                
            # Render the treatments for preview
            st.write("### Data Preview (Treatments):", df_treatments)
            
            # Let user split columns into groups (Vegetative vs. Reproductive/Yield)
            all_parameters = [col for col in df_cleaned.columns if col != treatment_col]
            
            st.write("### Table Splitting (Optional)")
            col1, col2 = st.columns(2)
            with col1:
                group_1_name = st.text_input("First Table Title", "Vegetative and Morphological Traits")
                group_1_cols = st.multiselect("Select parameters for Table 1", all_parameters, default=all_parameters[:len(all_parameters)//2])
            with col2:
                group_2_name = st.text_input("Second Table Title", "Yield and Structural Traits")
                group_2_cols = st.multiselect("Select parameters for Table 2", [c for c in all_parameters if c not in group_1_cols], default=[c for c in all_parameters if c not in group_1_cols])
            
            # Generate Report Button
            if st.button("Generate Word Document Draft"):
                # Define our 20 unique template strings
                templates = [
                    "For **{parameter}**, the treatment effect was {significance} {p_value}, with {top_geno} registering the maximum value of {top_val}^{top_let}, which did not differ statistically from {other_geno} ({other_val}^{other_let}).",
                    "Regarding **{parameter}**, genotypic variations were {significance} {p_value}, led by {top_geno} ({top_val}^{top_let}) and trailed by {low_geno} ({low_val}^{low_let}).",
                    "The F-test for **{parameter}** proved to be {significance} {p_value}, highlighting {top_geno} ({top_val}^{top_let}) as the most vigorous performer, whereas {low_geno} ({low_val}^{low_let}) recorded the poorest response.",
                    "The hierarchy of performance for **{parameter}** was {significance} {p_value}; {top_geno} ({top_val}^{top_let}) occupied the statistical apex, while {low_geno} ({low_val}^{low_let}) was situated at the baseline.",
                    "Data on **{parameter}** exhibited a {significance} {p_value} trend, with {top_geno} ({top_val}^{top_let}) representing the highest magnitude and {low_geno} ({low_val}^{low_let}) marking the lowest.",
                    "With respect to **{parameter}**, the evaluated germplasm showed a {significance} {p_value} variance, partitioning into distinct clusters where {top_geno} ({top_val}^{top_let}) and {other_geno} ({other_val}^{other_let}) shared statistical parity at the top.",
                    "A distinct, {significance} {p_value} separation among treatments was noticed for **{parameter}**, where {top_geno} ({top_val}^{top_let}) and {other_geno} ({other_val}^{other_let}) fell within the same premier DMRT group.",
                    "DMRT grouping for **{parameter}** revealed a {significance} {p_value} distribution, assigning the supreme letter '{top_let}' to {top_geno} ({top_val}^{top_let}).",
                    "Statistical evaluation of **{parameter}** indicated a {significance} {p_value} treatment response, with {top_geno} ({top_val}^{top_let}) sharing statistical letters with {other_geno} ({other_val}^{other_let}).",
                    "The variation expressed in **{parameter}** was found to be {significance} {p_value}, placing {top_geno} ({top_val}^{top_let}) and {other_geno} ({other_val}^{other_let}) in the highest-performing statistical tier.",
                    "The contrast among {treatment_plural} for **{parameter}** was {significance} {p_value}, with {top_geno} ({top_val}^{top_let}) exhibiting a distinct statistical advantage over {low_geno} ({low_val}^{low_let}).",
                    "A {significance} {p_value} difference was observed in **{parameter}**, where {top_geno} ({top_val}^{top_let}) surpassed the grand mean of {grand_mean}, while {low_geno} ({low_val}^{low_let}) remained significantly below it.",
                    "The evaluated {treatment_plural} differed to a {significance} {p_value} degree for **{parameter}**, with {top_geno} ({top_val}^{top_let}) demonstrating clear superiority over {low_geno} ({low_val}^{low_let}).",
                    "In terms of **{parameter}**, the differences were {significance} {p_value}, establishing {top_geno} ({top_val}^{top_let}) as the premier {treatment_singular}, statistically outclassing {low_geno} ({low_val}^{low_let}).",
                    "The performance profile for **{parameter}** was {significance} {p_value}, as {top_geno} ({top_val}^{top_let}) outperformed {low_geno} ({low_val}^{low_let}) by a statistically significant margin.",
                    "Analysis of variance for **{parameter}** yielded a {significance} {p_value} result, where the highest value of {top_val}^{top_let} was achieved by {top_geno}.",
                    "The influence of treatment on **{parameter}** was {significance} {p_value}; the maximum manifestation was observed in {top_geno} ({top_val}^{top_let}), while {low_geno} ({low_val}^{low_let}) recorded the minimum.",
                    "The evaluated {treatment_plural} exhibited {significance} {p_value} discrepancies in **{parameter}**, with {top_geno} ({top_val}^{top_let}) showing the most favorable response.",
                    "Statistical analysis confirmed that the genotypic effect on **{parameter}** was {significance} {p_value}, with {top_geno} ({top_val}^{top_let}) emerging as the outstanding treatment.",
                    "The response pattern of {treatment_plural} for **{parameter}** was {significance} {p_value}; {top_geno} ({top_val}^{top_let}) led the rankings, followed by {other_geno} ({other_val}^{other_let})."
                ]
                
                doc = Document()
                
                # Title and introduction
                doc.add_heading("Results and Discussion Draft", 0)
                doc.add_paragraph("This document contains dynamically generated results formatted for high-tier agricultural publication submission.")
                
                # Loop through both table groups
                groups = [(group_1_name, group_1_cols, 1), (group_2_name, group_2_cols, 2)]
                
                for g_title, g_cols, table_num in groups:
                    if not g_cols:
                        continue
                        
                    doc.add_heading(f"Table {table_num}: {g_title}", level=1)
                    
                    # Generate the text explanation for this group
                    text_paragraphs = []
                    
                    for idx, param in enumerate(g_cols):
                        # Extract data for this parameter
                        p_val_raw = str(stats_data.get("p-value", {}).get(param, "0.05"))
                        
                        # Determine Significance Text & P-value notation
                        sig_type = "significant"
                        p_notation = "(p < 0.05)"
                        
                        if "NS" in p_val_raw.upper():
                            sig_type = "nonsignificant"
                            p_notation = "(p > 0.05)"
                        elif "***" in p_val_raw or "0.000" in p_val_raw:
                            sig_type = "highly significant"
                            p_notation = "(p < 0.001)"
                        elif "**" in p_val_raw:
                            sig_type = "significant"
                            p_notation = "(p < 0.01)"
                        elif "*" in p_val_raw:
                            sig_type = "significant"
                            p_notation = "(p < 0.05)"
                            
                        # Extract Means, Letters, and standard stats
                        param_series = df_treatments[param]
                        parsed_pairs = [parse_dmrt_value(x) for x in param_series]
                        
                        # Sort to find Top, Intermediate, and Bottom Genotypes
                        # Clean values for numeric sorting
                        numeric_vals = []
                        for i, (num_str, let) in enumerate(parsed_pairs):
                            try:
                                numeric_vals.append((float(num_str), let, df_treatments.iloc[i][treatment_col]))
                            except ValueError:
                                pass
                                
                        if not numeric_vals:
                            continue
                            
                        numeric_vals.sort(reverse=True, key=lambda x: x[0])
                        
                        top_val, top_let, top_geno = numeric_vals[0]
                        low_val, low_let, low_geno = numeric_vals[-1]
                        
                        # Find secondary/runner up genotype
                        other_val, other_let, other_geno = numeric_vals[1] if len(numeric_vals) > 1 else (top_val, top_let, top_geno)
                        
                        grand_mean_val = str(stats_data.get("grand mean", {}).get(param, "N/A"))
                        
                        # Select a template dynamically to diversify descriptions
                        template_index = idx % len(templates)
                        template = templates[template_index]
                        
                        desc = template.format(
                            parameter=param,
                            significance=sig_type,
                            p_value=p_notation,
                            top_geno=top_geno,
                            top_val=f"{top_val:.2f}",
                            top_let=top_let,
                            other_geno=other_geno,
                            other_val=f"{other_val:.2f}",
                            other_let=other_let,
                            low_geno=low_geno,
                            low_val=f"{low_val:.2f}",
                            low_let=low_let,
                            grand_mean=grand_mean_val,
                            treatment_singular=treatment_label,
                            treatment_plural=treatment_label_plural
                        )
                        text_paragraphs.append(desc)
                        
                    # Write explanations into document
                    for para in text_paragraphs:
                        # Simple markdown-like formatter to support Bold text in word
                        p = doc.add_paragraph()
                        parts = re.split(r'(\*\*.*?\*\*)', para)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                p.add_run(part[2:-2]).bold = True
                            else:
                                p.add_run(part)
                                
                    # Write formatted Table into Word (APA Style)
                    doc.add_paragraph(f"Table {table_num}. Mean comparison of {g_title.lower()} among evaluated {treatment_label_plural}.")
                    
                    table = doc.add_table(rows=1, cols=len(g_cols) + 1)
                    set_table_borders(table)
                    
                    # Fill Header Row
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = str(treatment_col)
                    set_cell_margins(hdr_cells[0])
                    for c_idx, col_name in enumerate(g_cols):
                        hdr_cells[c_idx + 1].text = col_name
                        set_cell_margins(hdr_cells[c_idx + 1])
                    set_header_bottom_border(table.rows[0])
                    
                    # Fill Treatment Rows
                    for r_idx, t_row in df_treatments.iterrows():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(t_row[treatment_col])
                        set_cell_margins(row_cells[0])
                        for c_idx, col_name in enumerate(g_cols):
                            row_cells[c_idx + 1].text = str(t_row[col_name])
                            set_cell_margins(row_cells[c_idx + 1])
                            
                    # Add stats rows (SEM, LSD, CV, Grand Mean)
                    for stat_key in ["sem", "lsd", "cv", "grand mean"]:
                        if stat_key in stats_data:
                            row_cells = table.add_row().cells
                            # Format label
                            label_map = {"sem": "SEm (±)", "lsd": "LSD (0.05)", "cv": "CV (%)", "grand mean": "Grand Mean"}
                            row_cells[0].text = label_map.get(stat_key, stat_key.title())
                            set_cell_margins(row_cells[0])
                            for c_idx, col_name in enumerate(g_cols):
                                row_cells[c_idx + 1].text = str(stats_data[stat_key].get(col_name, "N/A"))
                                set_cell_margins(row_cells[c_idx + 1])
                                
                    doc.add_page_break()
                
                # Save Document to stream
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                
                st.write("🎉 **Your Draft Report is ready!**")
                st.download_button(
                    label="Download Draft Report (.docx)",
                    data=bio,
                    file_name="AgriStats_Draft_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
    except Exception as e:
        st.error(f"Error parsing file: {e}")
